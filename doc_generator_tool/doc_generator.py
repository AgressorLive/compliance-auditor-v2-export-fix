"""
Document Generator — Agreement on Advertising Services
Standalone PySide6 tool: site parsing, auto-fill, validation, signature gen, DOCX export.
"""

from __future__ import annotations

import io
import os
import random
import re
import sys
from typing import Optional

_APP_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_APP_DIR)
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QSpinBox, QDoubleSpinBox,
    QTextEdit, QGroupBox, QSplitter, QStatusBar, QFileDialog,
    QDateEdit, QMessageBox, QComboBox, QCheckBox, QScrollArea,
    QFrame, QSizePolicy, QTabWidget, QDialog, QGridLayout,
)
from PySide6.QtCore import Qt, QDate, QThread, Signal, QByteArray, QBuffer, QIODevice, QSettings
from PySide6.QtGui import QFont, QTextCursor, QPixmap

# ── Optional deps ─────────────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from fpdf import FPDF as _FPDF
    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False

try:
    from openai import OpenAI as _OpenAI
    _OPENAI_AVAILABLE = True
except ImportError:
    _OPENAI_AVAILABLE = False

try:
    from PIL import Image, ImageDraw, ImageFont
    _PIL_AVAILABLE = True
except ImportError:
    _PIL_AVAILABLE = False

try:
    from core.site_profiler import analyze_site
    _PROFILER_AVAILABLE = True
except ImportError:
    _PROFILER_AVAILABLE = False

# ── Dark stylesheet ───────────────────────────────────────────────────────────
DARK_STYLESHEET = """
QWidget {
    background-color: #1e1e2e;
    color: #cdd6f4;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 13px;
}
QMainWindow { background-color: #1e1e2e; }
QScrollArea { border: none; background-color: transparent; }
QScrollArea > QWidget > QWidget { background-color: transparent; }
QGroupBox {
    border: 1px solid #45475a;
    border-radius: 6px;
    margin-top: 10px;
    padding: 8px 6px 6px 6px;
    font-weight: bold;
    color: #89b4fa;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 10px;
    padding: 0 4px;
}
QLineEdit, QSpinBox, QDoubleSpinBox, QDateEdit, QComboBox {
    background-color: #313244;
    border: 1px solid #45475a;
    border-radius: 4px;
    padding: 5px 8px;
    color: #cdd6f4;
    selection-background-color: #89b4fa;
}
QLineEdit:focus, QSpinBox:focus, QDoubleSpinBox:focus,
QDateEdit:focus, QComboBox:focus {
    border: 1px solid #89b4fa;
}
QComboBox::drop-down { border: none; }
QComboBox QAbstractItemView {
    background-color: #313244;
    border: 1px solid #45475a;
    color: #cdd6f4;
    selection-background-color: #89b4fa;
}
QTextEdit {
    background-color: #181825;
    border: 1px solid #45475a;
    border-radius: 4px;
    color: #cdd6f4;
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 12px;
}
QPushButton {
    background-color: #313244;
    border: 1px solid #45475a;
    border-radius: 6px;
    padding: 6px 14px;
    color: #cdd6f4;
    font-weight: bold;
}
QPushButton:hover { background-color: #45475a; border-color: #89b4fa; color: #89b4fa; }
QPushButton:pressed { background-color: #89b4fa; color: #1e1e2e; }
QPushButton:disabled { color: #585b70; border-color: #313244; }
QPushButton#btn_generate {
    background-color: #89b4fa; color: #1e1e2e;
    font-size: 14px; padding: 10px 24px;
}
QPushButton#btn_generate:hover { background-color: #b4befe; }
QPushButton#btn_parse { background-color: #cba6f7; color: #1e1e2e; padding: 5px 12px; }
QPushButton#btn_parse:hover { background-color: #f5c2e7; }
QPushButton#btn_autofill { background-color: #fab387; color: #1e1e2e; padding: 6px 14px; }
QPushButton#btn_autofill:hover { background-color: #fe640b; color: #fff; }
QPushButton#btn_export_docx { background-color: #a6e3a1; color: #1e1e2e; }
QPushButton#btn_export_docx:hover { background-color: #40a02b; color: #fff; }
QPushButton#btn_export_txt { background-color: #89dceb; color: #1e1e2e; }
QPushButton#btn_export_txt:hover { background-color: #04a5e5; color: #fff; }
QPushButton#btn_export_pdf { background-color: #f38ba8; color: #1e1e2e; }
QPushButton#btn_export_pdf:hover { background-color: #e64553; color: #fff; }
QPushButton#btn_export_pdf:disabled { color: #585b70; border-color: #313244; }
QPushButton#btn_validate { background-color: #f9e2af; color: #1e1e2e; }
QPushButton#btn_validate:hover { background-color: #df8e1d; color: #fff; }
QPushButton#btn_sig_gen { background-color: #313244; padding: 4px 10px; font-size: 11px; }
QPushButton#btn_sig_gen:hover { background-color: #45475a; }
QPushButton#btn_sig_upload { background-color: #313244; padding: 4px 10px; font-size: 11px; }
QPushButton#btn_sig_upload:hover { background-color: #45475a; }
QLabel#section_label { color: #89b4fa; font-weight: bold; font-size: 12px; }
QLabel#sig_preview {
    background-color: #181825;
    border: 1px dashed #45475a;
    border-radius: 4px;
    min-height: 60px;
    color: #585b70;
    font-style: italic;
}
QCheckBox { spacing: 6px; }
QCheckBox::indicator {
    width: 16px; height: 16px;
    border: 1px solid #45475a; border-radius: 3px;
    background-color: #313244;
}
QCheckBox::indicator:checked { background-color: #89b4fa; border-color: #89b4fa; }
QStatusBar { background-color: #181825; color: #6c7086; border-top: 1px solid #45475a; }
QScrollBar:vertical { background: #181825; width: 7px; border-radius: 3px; }
QScrollBar::handle:vertical { background: #45475a; border-radius: 3px; }
QScrollBar::handle:vertical:hover { background: #89b4fa; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
QSplitter::handle { background-color: #45475a; }
QTabWidget::pane { border: 1px solid #45475a; border-radius: 4px; background: #1e1e2e; }
QTabBar::tab {
    background: #313244; color: #cdd6f4;
    padding: 7px 18px; border-radius: 4px 4px 0 0; margin-right: 2px;
    font-size: 12px;
}
QTabBar::tab:selected { background: #45475a; color: #89b4fa; font-weight: bold; }
QTabBar::tab:hover { background: #45475a; }
QPushButton#btn_scan { background-color: #89b4fa; color: #1e1e2e; padding: 5px 10px; }
QPushButton#btn_scan:hover { background-color: #b4befe; }
QPushButton#btn_gen_summary { background-color: #cba6f7; color: #1e1e2e; padding: 7px 16px; font-weight: bold; }
QPushButton#btn_gen_summary:hover { background-color: #f5c2e7; }
QPushButton#btn_gen_summary:disabled { color: #585b70; background-color: #313244; }
QPushButton#btn_copy { background-color: #313244; padding: 4px 10px; font-size: 11px; }
QPushButton#btn_copy:hover { background-color: #45475a; }
QPushButton#btn_send_agreement {
    background-color: #a6e3a1; color: #1e1e2e;
    padding: 9px 18px; font-size: 13px; font-weight: bold;
}
QPushButton#btn_send_agreement:hover { background-color: #40a02b; color: #fff; }
QPushButton#btn_send_agreement:disabled { color: #585b70; background-color: #313244; }
QPushButton#btn_mark_own { background-color: #89dceb; color: #1e1e2e; padding: 6px 12px; }
QPushButton#btn_mark_own:hover { background-color: #04a5e5; color: #fff; }
QPushButton#btn_mark_partner { background-color: #fab387; color: #1e1e2e; padding: 6px 12px; }
QPushButton#btn_mark_partner:hover { background-color: #fe640b; color: #fff; }
"""

# ── Agreement template ────────────────────────────────────────────────────────
AGREEMENT_TEMPLATE = """\
AGREEMENT ON THE PROVISION OF ADVERTISING SERVICES
No. {agreementNumber}

{agreementDate}

{contractorIntro}, on the one hand, and {clientIntro}, on the other hand, collectively referred to as the "Parties," have entered into this Agreement on the Provision of Advertising Services (hereinafter referred to as the "Agreement") as follows:


1. SUBJECT OF THE AGREEMENT

1.1. Under the terms and conditions set forth herein, the Contractor undertakes to render advertising services to the Client as specified in Annex No. 1 to Agreement No. {agreementNumber} dated {agreementDate} (hereinafter referred to as the "Annex"), and the Client undertakes to accept and pay for such services in the manner and within the timeframe prescribed by this Agreement.

1.2. The Annex forms an integral part of this Agreement and defines the scope, parameters, and cost of the services to be provided.


2. RIGHTS AND OBLIGATIONS OF THE PARTIES

2.1. The Contractor shall:
2.1.1. Render the advertising services within the timeframes set out in the Annex.
2.1.2. Notify the Client promptly of any circumstances that may affect the quality or timely delivery of the services.
2.1.3. Maintain the confidentiality of any information provided by the Client in connection with this Agreement.

2.2. The Client shall:
2.2.1. Provide the Contractor with all materials, data, and access credentials required for the proper performance of the services within three (3) business days from the date of this Agreement.
2.2.2. Accept the rendered services and sign the Service Acceptance Certificate within five (5) business days from receipt thereof, or submit a reasoned written objection within the same period.
2.2.3. Pay for the services in accordance with Section 3 of this Agreement.


3. PRICE AND PAYMENT PROCEDURE

3.1. The total cost of the services rendered hereunder amounts to {currencyWord} {servicePriceUsd} (hereinafter referred to as the "Service Fee"), excluding applicable taxes unless otherwise agreed in writing.

3.2. Payment shall be effected by the Client within five (5) business days from the date of signing the Service Acceptance Certificate, unless a different payment schedule is stipulated in the Annex.

3.3. All payments under this Agreement shall be made in the currency and via the method specified by the Contractor in the relevant invoice.


4. TERM OF THE AGREEMENT

4.1. This Agreement shall enter into force upon signature by both Parties and shall remain valid for a period of {termMonths} ({termMonthsText}) calendar months from the date hereof, unless terminated earlier in accordance with the provisions of this Agreement.

4.2. Upon expiration of the term, this Agreement shall be deemed fulfilled provided that both Parties have duly performed all obligations hereunder.


5. LIABILITY OF THE PARTIES

5.1. For failure to perform or improper performance of obligations under this Agreement, the Parties shall bear liability in accordance with applicable law and the terms hereof.

5.2. In the event of late payment, the Client shall pay a penalty in the amount of 0.1% of the outstanding amount for each calendar day of delay.

5.3. The Contractor's aggregate liability under this Agreement shall not exceed the total Service Fee actually paid by the Client.


6. CONFIDENTIALITY

6.1. Each Party undertakes to keep confidential all information received from the other Party in connection with this Agreement and not to disclose such information to third parties without the prior written consent of the disclosing Party.

6.2. The obligations set forth in this Section shall survive termination or expiration of this Agreement for a period of two (2) years.


7. FORCE MAJEURE

7.1. Neither Party shall be liable for failure to perform its obligations hereunder if such failure results from circumstances beyond that Party's reasonable control, including but not limited to natural disasters, acts of government, military actions, or other force majeure events.

7.2. The Party affected by force majeure shall notify the other Party in writing within five (5) business days of the occurrence of such circumstances.


8. EARLY TERMINATION

8.1. Either Party may terminate this Agreement by providing thirty (30) calendar days' prior written notice to the other Party.

8.2. In the event of termination, the Client shall pay for all services actually rendered by the Contractor up to the effective date of termination.


9. DISPUTE RESOLUTION

9.1. All disputes arising out of or in connection with this Agreement shall be resolved through good-faith negotiations between the Parties.

9.2. If a dispute cannot be resolved through negotiation within thirty (30) calendar days from the date of written notification, it shall be submitted to the competent court in accordance with applicable law.


10. FINAL PROVISIONS

10.1. This Agreement is drawn up in two (2) original counterparts, one for each Party, each having equal legal force.

10.2. Any amendments or supplements to this Agreement shall be valid only if made in writing and duly signed by authorized representatives of both Parties.

10.3. This Agreement supersedes and cancels all prior negotiations, representations, and agreements between the Parties concerning the subject matter hereof.


SIGNATURES OF THE PARTIES

Contractor: {contractorName}
[SIG_CONTRACTOR]
[STAMP_CONTRACTOR]
Date: {agreementDate}

Client: {clientName}
[SIG_CLIENT]
Date: {agreementDate}


================================================================================

ANNEX No. 1
to Agreement on the Provision of Advertising Services No. {agreementNumber}
dated {agreementDate}

SCOPE AND COST OF ADVERTISING SERVICES

1. SERVICES TO BE PROVIDED

The Contractor shall render the following advertising services to the Client:

1.1. Development, setup, and administration of advertising campaigns across digital platforms, including Google Ads, as agreed by the Parties.
1.2. Preparation of advertising creatives, targeting parameters, and campaign analytics reporting.
1.3. Ongoing campaign optimization based on performance data throughout the service term.

2. SERVICE PARAMETERS

2.1. Service Term: {termMonths} ({termMonthsText}) calendar months from {agreementDate}.
2.2. Total Service Fee: {currencyWord} {servicePriceUsd}.
2.3. Reporting: The Contractor shall provide a performance report to the Client no later than the 5th day of each calendar month following the reporting period.

3. ACCEPTANCE OF SERVICES

3.1. Upon completion of the services (or any agreed stage thereof), the Contractor shall submit a Service Acceptance Certificate to the Client.
3.2. The Client shall sign and return the Certificate within five (5) business days, or submit written objections within the same period.
3.3. Absence of written objections within the specified period shall constitute the Client's unconditional acceptance of the services rendered.

4. SPECIAL CONDITIONS

4.1. The Client guarantees that all advertising materials provided to the Contractor comply with applicable legislation and do not infringe upon the rights of third parties.
4.2. The Contractor reserves the right to suspend services in case of non-payment by the Client exceeding ten (10) business days past the due date.


Contractor: {contractorName}
[SIG_CONTRACTOR]
[STAMP_CONTRACTOR]
Date: {agreementDate}

Client: {clientName}
[SIG_CLIENT]
Date: {agreementDate}
"""

# ── Helpers ───────────────────────────────────────────────────────────────────
_MONTHS_WORDS = {
    1: "one", 2: "two", 3: "three", 4: "four", 5: "five", 6: "six",
    7: "seven", 8: "eight", 9: "nine", 10: "ten", 11: "eleven", 12: "twelve",
    18: "eighteen", 24: "twenty-four", 36: "thirty-six",
}


def _party_intro(name: str, role: str, party_type: str) -> str:
    name = name.strip()
    if party_type == "Individual":
        return (
            f'{name}, an individual acting on their own behalf '
            f'(hereinafter referred to as the "{role}")'
        )
    return (
        f'{name} (hereinafter referred to as the "{role}"), '
        f'acting on the basis of its statutory documents'
    )

def months_to_words(n: int) -> str:
    return _MONTHS_WORDS.get(n, str(n))

_PRICE_RANGES = {
    "Small  ($500 – $2,500)":     (500,   2500),
    "Medium ($3,000 – $7,500)":   (3000,  7500),
    "Large  ($10,000 – $50,000)": (10000, 50000),
    "Custom":                      None,
}

_USED_NUMBERS: set[str] = set()

def _random_agreement_number() -> str:
    for _ in range(1000):
        n = str(random.randint(10000, 99999))
        if n not in _USED_NUMBERS:
            _USED_NUMBERS.add(n)
            return n
    return str(random.randint(10000, 99999))

def _random_price(range_key: str) -> float:
    r = _PRICE_RANGES.get(range_key)
    if not r:
        return 5000.0
    lo, hi = r
    steps = list(range(lo, hi + 1, 500))
    return float(random.choice(steps))


def _random_past_date() -> QDate:
    """Return a random past date exactly 1, 1.5, or 2 months before today
    with a random valid day within that month."""
    import calendar as _cal
    today = QDate.currentDate()
    # Pick exactly 1 month (~30), 1.5 months (~45), or 2 months (~61) back
    days_back = random.choice([30, 45, 61])
    target = today.addDays(-days_back)
    year, month = target.year(), target.month()
    max_day = _cal.monthrange(year, month)[1]
    day = random.randint(1, max_day)
    return QDate(year, month, day)


# ── ODA helpers ───────────────────────────────────────────────────────────────

def clean_url(url: str) -> str:
    """Strip query string / tracking params; return scheme+host+path."""
    url = url.strip()
    if not url:
        return ""
    # Remove everything from '?' onwards
    if "?" in url:
        url = url[:url.index("?")]
    # Ensure trailing slash for bare domains
    from urllib.parse import urlparse
    p = urlparse(url)
    if p.scheme and p.netloc and not p.path:
        url = url.rstrip("/") + "/"
    return url


def safe_filename(text: str) -> str:
    """Convert a name into a filesystem-safe filename fragment."""
    text = text.strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s]+", "_", text)
    return text


def build_service_scope_sentence(selected: list[str]) -> str:
    if not selected:
        return "advertising campaign management services"
    if len(selected) == 1:
        return selected[0].lower()
    return ", ".join(s.lower() for s in selected[:-1]) + ", and " + selected[-1].lower()


_ODA_TEMPLATE = """\
OPERATING DISPLAY NAME AUTHORIZATION STATEMENT

Date: {statementDate}

I, {adsHolderName}, confirm that the Google Ads account associated with my name \
is used to promote the advertised business name and website brand \
"{advertisedBrandName}".

{adsHolderName} is not claiming "{advertisedBrandName}" as a separate registered \
legal name of his/her/their own. "{advertisedBrandName}" is the advertised \
business name and website brand of the client.

The use of the "{advertisedBrandName}" name, website, and related advertising \
content in Google Ads is based on a direct Advertising Services Agreement between \
{adsHolderName} and {advertisedBrandName} dated {agreementDate}.

Under this agreement, {adsHolderName} provides advertising campaign management \
services, including {serviceScopeSentence}.

{advertisedBrandName} remains responsible for the advertised business information, \
website content, brand information, and the products or services presented to users. \
{adsHolderName} is responsible for advertising campaign management only.

This statement is provided to clarify the relationship between the Google Ads \
account holder name and the advertised business name used in ads and on the \
associated website.
{regulatedParagraph}
Ads Account Holder / Advertising Contractor:
{adsHolderName}

Advertised Business / Client:
{advertisedBrandName}

Associated Website:
{associatedWebsite}

Agreement Reference:
Advertising Services Agreement No. {agreementNumber}, dated {agreementDate}

Signature:
{adsHolderName}
[SIG_ODA]

Date:
{statementDate}
"""

_ODA_REGULATED_PARAGRAPH = (
    "\nThe advertised website may contain informational content related to financial, "
    "legal, health, or other sensitive topics. {adsHolderName} provides advertising "
    "campaign management services only and does not provide regulated professional "
    "services unless such role is separately confirmed by verified documentation.\n"
)


def build_oda_template(data: dict) -> str:
    regulated = ""
    if data.get("regulated_disclaimer"):
        regulated = "\n" + _ODA_REGULATED_PARAGRAPH.format(
            adsHolderName=data["ads_holder_name"]
        ) + "\n"
    else:
        regulated = "\n"
    return _ODA_TEMPLATE.format(
        statementDate=data["statement_date"],
        adsHolderName=data["ads_holder_name"],
        advertisedBrandName=data["brand_name"],
        agreementDate=data["agreement_date"],
        agreementNumber=data["agreement_number"] or "[not specified]",
        associatedWebsite=clean_url(data["website"]),
        serviceScopeSentence=build_service_scope_sentence(data["service_scope"]),
        regulatedParagraph=regulated,
    )


def validate_oda_inputs(data: dict) -> list[str]:
    errs: list[str] = []
    if not data.get("ads_holder_name"):
        errs.append("Ads account holder name is required.")
    if not data.get("brand_name"):
        errs.append("Advertised brand name is required.")
    if not data.get("website"):
        errs.append("Associated website is required.")
    if not data.get("agreement_date"):
        errs.append("Agreement date is required.")
    if not data.get("agreement_number"):
        errs.append("⚠️ Agreement number is not specified (recommended).")
    holder = (data.get("ads_holder_name") or "").strip().lower()
    brand  = (data.get("brand_name")      or "").strip().lower()
    if holder and brand and holder == brand:
        errs.append("⚠️ ODA may not be needed because Ads holder name and brand name match.")
    return errs


_ODA_REQUIRED_LABELS = [
    "Ads Account Holder / Advertising Contractor:",
    "Advertised Business / Client:",
    "Associated Website:",
    "Agreement Reference:",
    "Signature:",
    "Date:",
]
_ODA_FORBIDDEN_PHRASES = [
    "registered DBA",
    "government registration",
    "legally owns the brand",
    "trademark owner",
    "licensed financial adviser",
    "certified provider",
]


def validate_gpt_oda_output(text: str, data: dict, static_len: int) -> list[str]:
    """Return list of validation failure messages (empty = passed)."""
    issues: list[str] = []
    tl = text.lower()
    if (
        "operating display name authorization statement" not in tl
        and "advertised name authorization" not in tl
    ):
        issues.append("Missing document title.")
    holder  = data.get("ads_holder_name", "")
    brand   = data.get("brand_name", "")
    website = clean_url(data.get("website", ""))
    agr_date = data.get("agreement_date", "")
    if holder and holder not in text:
        issues.append(f"Ads holder name '{holder}' missing from GPT output.")
    if brand and brand not in text:
        issues.append(f"Brand name '{brand}' missing from GPT output.")
    if website and website not in text:
        issues.append(f"Website '{website}' missing from GPT output.")
    if agr_date and agr_date not in text:
        issues.append(f"Agreement date '{agr_date}' missing from GPT output.")
    for lbl in _ODA_REQUIRED_LABELS:
        if lbl not in text:
            issues.append(f"Required label missing: '{lbl}'")
    if static_len > 0 and len(text) < static_len * 0.70:
        issues.append(
            f"GPT output too short ({len(text)} chars vs static {static_len} chars). "
            "Suspected summarisation."
        )
    for phrase in _ODA_FORBIDDEN_PHRASES:
        if phrase.lower() in tl:
            issues.append(f"Forbidden phrase detected: '{phrase}'")
    return issues

# ── Signature generation (Pillow) ─────────────────────────────────────────────
import math as _math

_BASTLIGA_PATHS = [
    r"C:\Users\alexx\OneDrive\Desktop\Bastliga One.otf",   # dev machine
    # When bundled by PyInstaller the font is extracted next to the exe
    *(
        [str(__import__("pathlib").Path(getattr(__import__("sys"), "_MEIPASS", "")) / "Bastliga One.otf")]
        if getattr(__import__("sys"), "_MEIPASS", None) else []
    ),
]

_SCRIPT_FONTS = _BASTLIGA_PATHS + [
    r"C:\Windows\Fonts\segoesc.ttf",    # Segoe Script
    r"C:\Windows\Fonts\BRUSHSCI.TTF",   # Brush Script MT
    r"C:\Windows\Fonts\Gabriola.ttf",   # Gabriola
    r"C:\Windows\Fonts\pacifico.ttf",   # Pacifico
    r"C:\Windows\Fonts\comic.ttf",      # Comic Sans fallback
    r"C:\Windows\Fonts\calibrii.ttf",   # Calibri italic fallback
]

def _load_script_font(size: int):
    if not _PIL_AVAILABLE:
        return None
    for fp in _SCRIPT_FONTS:
        try:
            return ImageFont.truetype(fp, size=size)
        except Exception:
            continue
    try:
        return ImageFont.load_default()
    except Exception:
        return None

_STAMP_FONTS = [
    r"C:\Users\alexx\AppData\Local\Microsoft\Windows\Fonts\DieselpowerPersonalUse-axaY5.ttf",
    r"C:\Users\alexx\AppData\Local\Microsoft\Windows\Fonts\Interplanetary Crap.otf",
    r"C:\Windows\Fonts\arialbd.ttf",
    r"C:\Windows\Fonts\calibrib.ttf",
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\calibri.ttf",
    r"C:\Windows\Fonts\segoeuib.ttf",
]

def _load_stamp_font(size: int):
    if not _PIL_AVAILABLE:
        return None
    for fp in _STAMP_FONTS:
        try:
            return ImageFont.truetype(fp, size=size)
        except Exception:
            continue
    try:
        return ImageFont.load_default()
    except Exception:
        return None


def _apply_stamp_variation(src_path: str, dst_path: str) -> str:
    """Apply random hand-stamped variation: rotation, scale, opacity, tiny offset.
    Returns dst_path (or src_path on failure)."""
    if not _PIL_AVAILABLE:
        return src_path
    import random as _r
    try:
        from PIL import Image as _I
        img = _I.open(src_path).convert("RGBA")
        w, h = img.size

        # Random params
        angle   = _r.uniform(-10, 10)
        scale   = _r.uniform(0.96, 1.04)
        opacity = _r.uniform(0.82, 0.97)
        ox      = int(_r.uniform(-w * 0.03, w * 0.03))
        oy      = int(_r.uniform(-h * 0.03, h * 0.03))

        # Scale
        nw, nh = int(w * scale), int(h * scale)
        rs = getattr(_I, "Resampling", None)
        lz = getattr(rs, "LANCZOS", None) or getattr(_I, "LANCZOS", 1)
        img = img.resize((nw, nh), lz)

        # Rotate (expand so nothing clips)
        bc = getattr(getattr(_I, "Resampling", None), "BICUBIC", None) or 3
        img = img.rotate(-angle, expand=True, resample=bc)
        rw, rh = img.size

        # Apply opacity to alpha channel
        r_ch, g_ch, b_ch, a_ch = img.split()
        a_data = [int(p * opacity) for p in a_ch.getdata()]
        a_ch.putdata(a_data)
        img.putalpha(a_ch)

        # Composite onto white canvas (same size as rotated image + offset margin)
        m = max(abs(ox), abs(oy)) + 4
        canvas = _I.new("RGB", (rw + 2*m, rh + 2*m), (255, 255, 255))
        canvas.paste(img, (m + ox, m + oy), img)
        canvas = canvas.crop((m, m, m + rw, m + rh))

        canvas.save(dst_path, format="PNG")
        return dst_path
    except Exception:
        return src_path


def generate_stamp_pixmap(company_name: str, size: int = 150) -> Optional[QPixmap]:
    """Classic rubber stamp: double ring, arc text top & bottom, stars, initials."""
    if not _PIL_AVAILABLE or not company_name.strip():
        return None

    import random as _rnd
    from PIL import ImageFilter as _IFlt
    _rnd.seed(42)

    sc = 6                          # high supersample for smooth circles
    sz = size * sc
    cx = cy = sz // 2

    r_outer = int(sz * 0.470)   # slightly larger outer ring for wider white gap
    r_inner = int(sz * 0.345)   # slightly smaller inner ring for wider white gap

    # Ring stroke widths — pre-compute BEFORE r_text so we know white-gap bounds
    bw_outer = max(6, sz // 20)   # slightly thinner outer stroke = more white space
    bw_inner = max(2, sz // 70)

    # White gap = space between the inner edge of outer ring and outer edge of inner ring
    # Ellipse outline is drawn centred on the radius, so strokes extend inward/outward by half
    gap_inner = r_inner + bw_inner // 2 + 3   # just outside inner ring stroke + 3px safety
    gap_outer = r_outer - bw_outer // 2 - 3   # just inside outer ring stroke  + 3px safety
    white_gap = max(4, gap_outer - gap_inner)  # pixel height of usable white space
    r_text    = gap_inner + white_gap // 2     # geometric centre of the white gap

    band    = r_outer - r_inner   # kept for star placement

    INK   = (7, 59, 166)
    INKA  = (7, 59, 166, 255)
    TRANS = (0, 0, 0, 0)

    stamp = Image.new("RGBA", (sz, sz), TRANS)
    draw  = ImageDraw.Draw(stamp)

    # ── Rings ─────────────────────────────────────────────────────────────────
    draw.ellipse([cx - r_outer, cy - r_outer, cx + r_outer, cy + r_outer],
                 outline=INKA, width=bw_outer)
    draw.ellipse([cx - r_inner, cy - r_inner, cx + r_inner, cy + r_inner],
                 outline=INKA, width=bw_inner)

    display   = company_name.strip().upper()
    px_to_deg = 360.0 / (2 * _math.pi * r_text)

    rs     = getattr(Image, "Resampling", None)
    bicubic = getattr(rs, "BICUBIC", None) or getattr(Image, "BICUBIC", 3)

    def _str_px(s, fnt):
        try:
            return max(1, fnt.getlength(s))
        except Exception:
            try:
                bb = fnt.getbbox(s); return max(1, bb[2] - bb[0])
            except Exception:
                return len(s) * int(getattr(fnt, "size", 20) * 0.6)

    def _char_px(c, fnt):
        try:
            return max(1, fnt.getlength(c))
        except Exception:
            try:
                bb = fnt.getbbox(c); return max(1, bb[2] - bb[0])
            except Exception:
                return int(getattr(fnt, "size", 20) * 0.6)

    def _paste(layer, px, py):
        """Paste RGBA layer onto stamp with alpha_composite, clipped to canvas."""
        ix, iy = int(px), int(py)
        if ix + layer.width < 0 or iy + layer.height < 0: return
        if ix >= sz or iy >= sz: return
        # Clip source if destination goes out of bounds
        x0 = max(0, ix); y0 = max(0, iy)
        sx0 = x0 - ix;   sy0 = y0 - iy
        w = min(layer.width  - sx0, sz - x0)
        h = min(layer.height - sy0, sz - y0)
        if w <= 0 or h <= 0: return
        crop = layer.crop((sx0, sy0, sx0 + w, sy0 + h))
        stamp.alpha_composite(crop, dest=(x0, y0))

    def _draw_char(c, font, tx, ty, rot_deg):
        fsz = getattr(font, "size", 20)
        pad = fsz * 5
        ci  = Image.new("RGBA", (pad, pad), TRANS)
        cd  = ImageDraw.Draw(ci)
        try:
            cd.text((pad // 2, pad // 2), c, font=font, fill=INKA, anchor="mm")
        except TypeError:
            cd.text((pad // 4, pad // 4), c, font=font, fill=INKA)
        cr = ci.rotate(-rot_deg, expand=True, resample=bicubic)
        _paste(cr, tx - cr.width / 2, ty - cr.height / 2)

    def _arc_text(text, font, r_arc, top: bool):
        total_deg = _str_px(text, font) * px_to_deg
        if top:
            # centre at 270° (top); letters base toward centre, readable outside
            cur = 270.0 - total_deg / 2.0
            for c in text:
                cdeg = _char_px(c, font) * px_to_deg
                a    = cur + cdeg / 2.0
                ar   = _math.radians(a)
                _draw_char(c, font, cx + r_arc * _math.cos(ar),
                           cy + r_arc * _math.sin(ar), a + 90.0)
                cur += cdeg
        else:
            # centre at 90° (bottom); letters base toward centre, readable from outside bottom
            cur = 90.0 + total_deg / 2.0
            for c in text:
                cdeg = _char_px(c, font) * px_to_deg
                a    = cur - cdeg / 2.0
                ar   = _math.radians(a)
                _draw_char(c, font, cx + r_arc * _math.cos(ar),
                           cy + r_arc * _math.sin(ar), a - 90.0)
                cur -= cdeg

    # ── Font sizing: 60% of white gap, verified against real rendered cap height ─
    safe_pad  = max(3, int(white_gap * 0.15))    # 15% margin from each ring edge
    max_cap_h = white_gap - 2 * safe_pad         # hard ceiling for glyph height

    font_size = min(int(white_gap * 0.60), max(8, int(sz * 0.070)))
    font      = _load_stamp_font(font_size)

    # Measure REAL cap height via getbbox("M") — uppercase caps define the visual height
    def _cap_height(fnt):
        try:
            bb = fnt.getbbox("M")
            return max(1, bb[3] - bb[1])
        except Exception:
            return int(getattr(fnt, "size", 20) * 0.72)

    # Iteratively shrink until cap fits within max_cap_h
    cap_h = _cap_height(font)
    while cap_h > max_cap_h and font_size > 6:
        font_size = max(6, int(font_size * max_cap_h / cap_h) - 1)
        font  = _load_stamp_font(font_size)
        cap_h = _cap_height(font)

    # Shrink if arc span exceeds 195°
    total_arc = _str_px(display, font) * px_to_deg
    if total_arc > 195:
        font_size = max(6, int(font_size * 195 / total_arc))
        font  = _load_stamp_font(font_size)
        cap_h = _cap_height(font)

    # ── Compute per-arc anchor radii from ACTUAL letter metrics ──────────────
    # With anchor="mm", glyph bbox[1] is the top offset from anchor (negative = above),
    # bbox[3] is the bottom offset (positive = below).
    # For the TOP arc: letter top points outward (toward outer ring).
    #   Safe outer edge:  r_anchor - bb[1]  <= gap_outer - safe_pad
    #   → r_anchor_top = gap_outer - safe_pad + bb[1]   (bb[1] < 0, so this moves inward)
    # For the BOTTOM arc: letter top also points outward (toward outer ring, letters flipped).
    #   Same logic applies symmetrically.
    # We take the midpoint so padding is equal on both sides:
    #   r_anchor = ((gap_outer - safe_pad + bb[1]) + (gap_inner + safe_pad - bb[3])) / 2
    try:
        bb_m = font.getbbox("M", anchor="mm")  # relative to "mm" anchor
        # r where letter top (outward) just clears the outer ring
        r_from_outer = gap_outer - safe_pad + bb_m[1]   # bb_m[1] < 0 → moves r inward ✓
        # r where letter bottom (inward) just clears the inner ring
        r_from_inner = gap_inner + safe_pad - bb_m[3]   # bb_m[3] > 0 → moves r outward ✓
        # Average → equal visual breathing room on both sides
        r_arc_base   = (r_from_outer + r_from_inner) / 2
        # Top and bottom arcs have the same geometry; bias only if font has internal leading
        cap_bias     = (bb_m[1] + bb_m[3]) / 2          # >0 = bbox skewed downward
        r_text_top   = r_arc_base - cap_bias
        r_text_bot   = r_arc_base + cap_bias
    except Exception:
        r_text_top = r_text_bot = r_text

    _arc_text(display, font, r_text_top, top=True)
    _arc_text(display, font, r_text_bot, top=False)

    # ── Stars between rings at 3 & 9 o'clock ──────────────────────────────────
    star_sz   = max(6, int(sz * 0.065))
    star_font = _load_stamp_font(star_sz)
    r_star    = r_inner + (r_outer - r_inner) // 2
    for ang in (0, 180):
        ar  = _math.radians(ang)
        sx  = cx + r_star * _math.cos(ar)
        sy  = cy + r_star * _math.sin(ar)
        pad = star_sz * 3
        si  = Image.new("RGBA", (pad, pad), TRANS)
        sd  = ImageDraw.Draw(si)
        try:
            sd.text((pad // 2, pad // 2), "★", font=star_font, fill=INKA, anchor="mm")
            _paste(si, sx - pad / 2, sy - pad / 2)
        except Exception:
            pass

    # ── Centre: big initials + horizontal lines ────────────────────────────────
    abbr = "".join(w[0] for w in company_name.strip().upper().split() if w)[:4]
    if abbr:
        cf_sz = max(10, int(sz * 0.160))
        cf    = _load_stamp_font(cf_sz)
        try:
            draw.text((cx, cy), abbr, font=cf, fill=INKA, anchor="mm")
        except TypeError:
            bb = draw.textbbox((0, 0), abbr, font=cf)
            draw.text((cx - (bb[2]-bb[0])//2, cy - (bb[3]-bb[1])//2 - bb[1]),
                      abbr, font=cf, fill=INKA)

    lw   = max(3, sz // 70)
    lpad = int(r_inner * 0.25)
    for dy in (-int(sz * 0.002), int(sz * 0.002)):    # two close parallel lines
        draw.line([(cx - r_inner + lpad, cy + dy), (cx + r_inner - lpad, cy + dy)],
                  fill=INKA, width=lw)

    # ── Ink-bleed texture: darken outer ring edge, lighten inner ─────────────
    # Simulate ink absorption by multiplying alpha with radial gradient
    pixels = list(stamp.getdata())
    worn   = []
    for idx, (r, g, b, a) in enumerate(pixels):
        if a == 0:
            worn.append((r, g, b, 0))
            continue
        px = idx % sz - cx
        py = idx // sz - cy
        d  = _math.sqrt(px*px + py*py)
        # Slightly more opaque near outer edge, slightly transparent in mid-ring
        edge_boost = 1.0
        if r_inner < d < r_outer:
            # inside ring band — boost
            edge_boost = 1.05
        # random micro-dropout to simulate worn rubber
        dropout = _rnd.random()
        if dropout < 0.012:    # 1.2% random missing ink
            a = int(a * _rnd.uniform(0.0, 0.45))
        elif dropout < 0.035:  # additional 2.3% partial wear
            a = int(a * _rnd.uniform(0.55, 0.80))
        else:
            a = min(255, int(a * edge_boost))
        worn.append((r, g, b, a))
    stamp.putdata(worn)

    # Light blur to smooth the jagged dropout edges → natural ink look
    stamp = stamp.filter(_IFlt.GaussianBlur(radius=max(0.5, sz / 700)))

    # ── Composite on white ────────────────────────────────────────────────────
    bg = Image.new("RGB", (sz, sz), (255, 255, 255))
    bg.paste(stamp, mask=stamp.split()[3])

    rs2  = getattr(Image, "Resampling", None)
    lczs = getattr(rs2, "LANCZOS", None) or getattr(Image, "LANCZOS", 1)
    out  = bg.resize((size, size), lczs)

    buf = io.BytesIO()
    out.save(buf, format="PNG")
    buf.seek(0)
    pixmap = QPixmap()
    pixmap.loadFromData(QByteArray(buf.read()))
    return pixmap




    sc = 4
    sz = size * sc
    cx = cy = sz // 2

    r_outer = int(sz * 0.455)
    r_inner = int(sz * 0.360)
    r_text  = int(sz * 0.410)

    BLUE  = (7, 59, 166)
    BLUEA = (7, 59, 166, 240)
    TRANS = (0, 0, 0, 0)
    WHITE = (255, 255, 255, 255)

    # ── Main canvas (RGBA, transparent) ──────────────────────────────────────
    stamp = Image.new("RGBA", (sz, sz), TRANS)
    draw  = ImageDraw.Draw(stamp)

    bw_outer = max(6, sz // 20)
    bw_inner = max(2, sz // 65)
    draw.ellipse([cx - r_outer, cy - r_outer, cx + r_outer, cy + r_outer],
                 outline=BLUEA, width=bw_outer)
    draw.ellipse([cx - r_inner, cy - r_inner, cx + r_inner, cy + r_inner],
                 outline=BLUEA, width=bw_inner)

    display   = company_name.strip().upper()
    px_to_deg = 360.0 / (2 * _math.pi * r_text)

    def _str_px(s, fnt):
        try:
            return max(1, fnt.getlength(s))
        except Exception:
            try:
                bb = fnt.getbbox(s)
                return max(1, bb[2] - bb[0])
            except Exception:
                return len(s) * int(getattr(fnt, "size", 20) * 0.6)

    def _char_px(c, fnt):
        try:
            return max(1, fnt.getlength(c))
        except Exception:
            try:
                bb = fnt.getbbox(c)
                return max(1, bb[2] - bb[0])
            except Exception:
                return int(getattr(fnt, "size", 20) * 0.6)

    rs     = getattr(Image, "Resampling", None)
    resamp = getattr(rs, "BICUBIC", None) or getattr(Image, "BICUBIC", 3)

    def _paste_char(c, font, tx, ty, rot, color=BLUEA):
        fsz = getattr(font, "size", 20)
        pad = fsz * 4
        c_img = Image.new("RGBA", (pad, pad), TRANS)
        c_drw = ImageDraw.Draw(c_img)
        try:
            c_drw.text((pad // 2, pad // 2), c, font=font, fill=color, anchor="mm")
        except TypeError:
            c_drw.text((pad // 4, pad // 4), c, font=font, fill=color)
        c_rot = c_img.rotate(-rot, expand=True, resample=resamp)
        stamp.alpha_composite(c_rot, dest=(
            max(0, int(tx - c_rot.width / 2)),
            max(0, int(ty - c_rot.height / 2))
        ))

    def _draw_arc_text(text, font, r_arc, top: bool):
        total_deg = _str_px(text, font) * px_to_deg
        if top:
            start_a = 270.0 - total_deg / 2.0
            cur_a   = start_a
            for c in text:
                char_deg = _char_px(c, font) * px_to_deg
                a_deg = cur_a + char_deg / 2.0
                a_rad = _math.radians(a_deg)
                _paste_char(c, font,
                            cx + r_arc * _math.cos(a_rad),
                            cy + r_arc * _math.sin(a_rad),
                            a_deg + 90.0)
                cur_a += char_deg
        else:
            start_a = 90.0 - total_deg / 2.0
            cur_a   = start_a
            for c in text:
                char_deg = _char_px(c, font) * px_to_deg
                a_deg = cur_a + char_deg / 2.0
                a_rad = _math.radians(a_deg)
                _paste_char(c, font,
                            cx + r_arc * _math.cos(a_rad),
                            cy + r_arc * _math.sin(a_rad),
                            a_deg - 90.0)
                cur_a += char_deg

    # ── Fonts ─────────────────────────────────────────────────────────────────
    font_size = max(10, int(sz * 0.072))
    font      = _load_stamp_font(font_size)
    total_deg = _str_px(display, font) * px_to_deg
    if total_deg > 200:
        font_size = max(8, int(font_size * 200 / total_deg))
        font      = _load_stamp_font(font_size)

    _draw_arc_text(display, font, r_text, top=True)
    _draw_arc_text(display, font, r_text, top=False)

    # ── Stars between rings at 3 & 9 o'clock ──────────────────────────────────
    star_sz   = max(6, int(sz * 0.068))
    star_font = _load_stamp_font(star_sz)
    r_star    = r_inner + (r_outer - r_inner) // 2
    for angle_deg in (0, 180):
        a_rad = _math.radians(angle_deg)
        sx = int(cx + r_star * _math.cos(a_rad))
        sy = int(cy + r_star * _math.sin(a_rad))
        pad = star_sz * 3
        s_img = Image.new("RGBA", (pad, pad), TRANS)
        s_drw = ImageDraw.Draw(s_img)
        try:
            s_drw.text((pad // 2, pad // 2), "★", font=star_font, fill=BLUEA, anchor="mm")
            stamp.alpha_composite(s_img, dest=(sx - pad // 2, sy - pad // 2))
        except Exception:
            pass

    # ── Diagonal banner across centre ─────────────────────────────────────────
    abbr = "".join(w[0] for w in company_name.strip().upper().split() if w)[:4]
    banner_w = int(r_inner * 1.72)
    banner_h = int(sz * 0.155)
    banner_img = Image.new("RGBA", (banner_w, banner_h), TRANS)
    b_drw = ImageDraw.Draw(banner_img)
    brad  = banner_h // 4
    b_drw.rounded_rectangle([0, 0, banner_w - 1, banner_h - 1],
                             radius=brad, fill=BLUEA)

    # Stars inside banner (left and right of text)
    if abbr:
        abr_sz  = max(10, int(banner_h * 0.60))
        abr_fnt = _load_stamp_font(abr_sz)
        # abbreviation centered
        try:
            b_drw.text((banner_w // 2, banner_h // 2), abbr,
                       font=abr_fnt, fill=WHITE, anchor="mm")
        except TypeError:
            b_drw.text((banner_w // 4, banner_h // 4), abbr,
                       font=abr_fnt, fill=WHITE)

        # small stars inside banner left & right
        bstar_sz  = max(4, int(banner_h * 0.35))
        bstar_fnt = _load_stamp_font(bstar_sz)
        try:
            bpad = bstar_sz * 3
            for bx in (int(banner_w * 0.15), int(banner_w * 0.85)):
                bs_img = Image.new("RGBA", (bpad, bpad), TRANS)
                bs_drw = ImageDraw.Draw(bs_img)
                bs_drw.text((bpad // 2, bpad // 2), "★",
                            font=bstar_fnt, fill=WHITE, anchor="mm")
                banner_img.alpha_composite(bs_img,
                    dest=(bx - bpad // 2, banner_h // 2 - bpad // 2))
        except Exception:
            pass

    # Rotate banner ~-5° (diagonal look), paste onto stamp centred
    brot = banner_img.rotate(5, expand=True, resample=resamp)
    bx_off = cx - brot.width  // 2
    by_off = cy - brot.height // 2
    stamp.alpha_composite(brot, dest=(max(0, bx_off), max(0, by_off)))

    # ── Stars above & below banner (inside inner ring) ────────────────────────
    inner_star_sz   = max(6, int(sz * 0.075))
    inner_star_font = _load_stamp_font(inner_star_sz)
    for iy_off in (-int(r_inner * 0.52), int(r_inner * 0.52)):
        ipad = inner_star_sz * 3
        si = Image.new("RGBA", (ipad, ipad), TRANS)
        sd = ImageDraw.Draw(si)
        try:
            sd.text((ipad // 2, ipad // 2), "★",
                    font=inner_star_font, fill=BLUEA, anchor="mm")
            stamp.alpha_composite(si, dest=(cx - ipad // 2, cy + iy_off - ipad // 2))
        except Exception:
            pass

    # ── Grunge: subtle white scratches + tiny dots inside circle only ─────────
    grunge = Image.new("RGBA", (sz, sz), TRANS)
    g_drw  = ImageDraw.Draw(grunge)
    n_dots = int(sz * sz * 0.00035)   # ~0.035% — just a hint of wear
    for _ in range(n_dots):
        x = _rnd.randint(0, sz - 1)
        y = _rnd.randint(0, sz - 1)
        if (x - cx) ** 2 + (y - cy) ** 2 > r_outer ** 2:
            continue
        r = _rnd.randint(1, max(1, sz // 130))
        alpha = _rnd.randint(80, 160)
        g_drw.ellipse([x - r, y - r, x + r, y + r], fill=(255, 255, 255, alpha))
    # A few short white scratches
    for _ in range(int(sz * 0.025)):
        x1 = _rnd.randint(cx - r_outer, cx + r_outer)
        y1 = _rnd.randint(cy - r_outer, cy + r_outer)
        if (x1 - cx) ** 2 + (y1 - cy) ** 2 > r_outer ** 2:
            continue
        ang  = _rnd.uniform(0, _math.pi)
        llen = _rnd.randint(sz // 70, sz // 30)
        x2   = int(x1 + llen * _math.cos(ang))
        y2   = int(y1 + llen * _math.sin(ang))
        alpha = _rnd.randint(60, 140)
        g_drw.line([(x1, y1), (x2, y2)], fill=(255, 255, 255, alpha), width=1)
    stamp.alpha_composite(grunge)

    # ── Composite onto white background ──────────────────────────────────────
    bg = Image.new("RGB", (sz, sz), (255, 255, 255))
    bg.paste(stamp, mask=stamp.split()[3])

    rs2  = getattr(Image, "Resampling", None)
    lczs = getattr(rs2, "LANCZOS", None) or getattr(Image, "LANCZOS", 1)
    out  = bg.resize((size, size), lczs)

    buf = io.BytesIO()
    out.save(buf, format="PNG")
    buf.seek(0)
    pixmap = QPixmap()
    pixmap.loadFromData(QByteArray(buf.read()))
    return pixmap



    sc = 4
    sz = size * sc
    cx = cy = sz // 2

    r_outer  = int(sz * 0.455)
    r_inner  = int(sz * 0.360)
    r_text   = int(sz * 0.410)

    BLUE  = (7, 59, 166, 235)
    TRANS = (255, 255, 255, 0)

    # ── Build grunge mask (random blobs that eat into ink) ────────────────────
    from PIL import ImageFilter as _IFlt
    grunge = Image.new("L", (sz, sz), 255)
    g_drw  = ImageDraw.Draw(grunge)
    n_holes = int(sz * sz * 0.018)   # ~1.8% of pixels are hole seeds
    for _ in range(n_holes):
        x = _rnd.randint(0, sz - 1)
        y = _rnd.randint(0, sz - 1)
        r = _rnd.randint(1, max(2, sz // 60))
        g_drw.ellipse([x - r, y - r, x + r, y + r], fill=_rnd.randint(0, 80))
    # Blur so holes blend naturally
    grunge = grunge.filter(_IFlt.GaussianBlur(radius=max(1, sz // 120)))
    # Threshold: pixels below 180 become transparent in ink
    grunge = grunge.point(lambda p: 255 if p > 160 else 0)

    # ── Stamp layer (RGBA transparent) ───────────────────────────────────────
    stamp = Image.new("RGBA", (sz, sz), (0, 0, 0, 0))
    draw  = ImageDraw.Draw(stamp)

    bw_outer = max(6, sz // 20)
    bw_inner = max(2, sz // 60)
    draw.ellipse([cx - r_outer, cy - r_outer, cx + r_outer, cy + r_outer],
                 outline=BLUE, width=bw_outer)
    draw.ellipse([cx - r_inner, cy - r_inner, cx + r_inner, cy + r_inner],
                 outline=BLUE, width=bw_inner)

    display   = company_name.strip().upper()
    px_to_deg = 360.0 / (2 * _math.pi * r_text)

    def _str_px(s, fnt):
        try:
            return max(1, fnt.getlength(s))
        except Exception:
            try:
                bb = fnt.getbbox(s)
                return max(1, bb[2] - bb[0])
            except Exception:
                return len(s) * int(getattr(fnt, "size", 20) * 0.6)

    def _char_px(c, fnt):
        try:
            return max(1, fnt.getlength(c))
        except Exception:
            try:
                bb = fnt.getbbox(c)
                return max(1, bb[2] - bb[0])
            except Exception:
                return int(getattr(fnt, "size", 20) * 0.6)

    rs     = getattr(Image, "Resampling", None)
    resamp = getattr(rs, "BICUBIC", None) or getattr(Image, "BICUBIC", 3)

    def _paste_char(c, font, tx, ty, rot):
        fsz = getattr(font, "size", 20)
        pad = fsz * 4
        c_img = Image.new("RGBA", (pad, pad), TRANS)
        c_drw = ImageDraw.Draw(c_img)
        try:
            c_drw.text((pad // 2, pad // 2), c, font=font, fill=BLUE, anchor="mm")
        except TypeError:
            c_drw.text((pad // 4, pad // 4), c, font=font, fill=BLUE)
        c_rot = c_img.rotate(-rot, expand=True, resample=resamp)
        stamp.paste(c_rot, (int(tx - c_rot.width / 2), int(ty - c_rot.height / 2)), c_rot)

    def _draw_arc_text(text, font, r_arc, top: bool):
        total_deg = _str_px(text, font) * px_to_deg
        if top:
            start_a = 270.0 - total_deg / 2.0
            cur_a   = start_a
            for c in text:
                char_deg = _char_px(c, font) * px_to_deg
                a_deg = cur_a + char_deg / 2.0
                a_rad = _math.radians(a_deg)
                _paste_char(c, font, cx + r_arc * _math.cos(a_rad),
                            cy + r_arc * _math.sin(a_rad), a_deg + 90.0)
                cur_a += char_deg
        else:
            start_a = 90.0 - total_deg / 2.0
            cur_a   = start_a
            for c in text:
                char_deg = _char_px(c, font) * px_to_deg
                a_deg = cur_a + char_deg / 2.0
                a_rad = _math.radians(a_deg)
                _paste_char(c, font, cx + r_arc * _math.cos(a_rad),
                            cy + r_arc * _math.sin(a_rad), a_deg - 90.0)
                cur_a += char_deg

    # ── Font ──────────────────────────────────────────────────────────────────
    font_size = max(10, int(sz * 0.072))
    font      = _load_stamp_font(font_size)
    total_deg = _str_px(display, font) * px_to_deg
    if total_deg > 200:
        font_size = max(8, int(font_size * 200 / total_deg))
        font      = _load_stamp_font(font_size)

    _draw_arc_text(display, font, r_text, top=True)
    _draw_arc_text(display, font, r_text, top=False)

    # ── Stars between rings at 3 & 9 o'clock ─────────────────────────────────
    star_sz   = max(6, int(sz * 0.070))
    star_font = _load_stamp_font(star_sz)
    r_star    = r_inner + (r_outer - r_inner) // 2
    for angle_deg in (0, 180):
        a_rad = _math.radians(angle_deg)
        sx = int(cx + r_star * _math.cos(a_rad))
        sy = int(cy + r_star * _math.sin(a_rad))
        pad = star_sz * 3
        s_img = Image.new("RGBA", (pad, pad), TRANS)
        s_drw = ImageDraw.Draw(s_img)
        try:
            s_drw.text((pad // 2, pad // 2), "★", font=star_font, fill=BLUE, anchor="mm")
            stamp.paste(s_img, (sx - pad // 2, sy - pad // 2), s_img)
        except Exception:
            pass

    # ── Stars inside inner circle at 3 & 9 o'clock ───────────────────────────
    inner_star_sz   = max(6, int(sz * 0.090))
    inner_star_font = _load_stamp_font(inner_star_sz)
    r_inner_star    = int(r_inner * 0.72)
    for angle_deg in (0, 180):
        a_rad = _math.radians(angle_deg)
        sx = int(cx + r_inner_star * _math.cos(a_rad))
        sy = int(cy + r_inner_star * _math.sin(a_rad))
        pad = inner_star_sz * 3
        s_img = Image.new("RGBA", (pad, pad), TRANS)
        s_drw = ImageDraw.Draw(s_img)
        try:
            s_drw.text((pad // 2, pad // 2), "★", font=inner_star_font, fill=BLUE, anchor="mm")
            stamp.paste(s_img, (sx - pad // 2, sy - pad // 2), s_img)
        except Exception:
            pass

    # ── Centre abbreviation ───────────────────────────────────────────────────
    abbr = "".join(w[0] for w in company_name.strip().upper().split() if w)[:4]
    if abbr:
        cf_size = max(10, int(sz * 0.155))
        cf      = _load_stamp_font(cf_size)
        try:
            draw.text((cx, cy), abbr, font=cf, fill=BLUE, anchor="mm")
        except TypeError:
            bb = draw.textbbox((0, 0), abbr, font=cf)
            draw.text((cx - (bb[2]-bb[0])//2, cy - (bb[3]-bb[1])//2 - bb[1]), abbr, font=cf, fill=BLUE)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    lw   = max(2, sz // 80)
    lpad = int(r_inner * 0.22)
    draw.line([(cx - r_inner + lpad, cy), (cx + r_inner - lpad, cy)], fill=BLUE, width=lw)

    # ── Apply grunge: mask out random holes from the ink layer ────────────────
    # Split stamp alpha, multiply by grunge mask, put back
    r_ch, g_ch, b_ch, a_ch = stamp.split()
    a_worn = Image.new("L", (sz, sz))
    # Multiply alpha by grunge mask
    import struct as _st
    a_pixels = list(a_ch.getdata())
    g_pixels = list(grunge.getdata())
    worn = [int(a * g / 255) for a, g in zip(a_pixels, g_pixels)]
def _bezier_pts(p0, p1, p2, p3, steps: int = 40):
    """Return list of (x, y) points along a cubic Bézier curve."""
    pts = []
    for i in range(steps + 1):
        t = i / steps
        mt = 1 - t
        x = mt**3*p0[0] + 3*mt**2*t*p1[0] + 3*mt*t**2*p2[0] + t**3*p3[0]
        y = mt**3*p0[1] + 3*mt**2*t*p1[1] + 3*mt*t**2*p2[1] + t**3*p3[1]
        pts.append((x, y))
    return pts


def _draw_curve(draw, pts, color, width: int = 2):
    """Draw a smooth polyline through a list of (x, y) points."""
    for i in range(len(pts) - 1):
        draw.line([pts[i], pts[i + 1]], fill=color, width=width)


def _draw_ink_stroke(draw, pts, color, base_width: int = 3):
    """
    Draw a stroke with simulated pen-pressure variation:
    thick in the middle, tapers at ends.
    """
    n = len(pts)
    for i in range(n - 1):
        t = i / max(n - 2, 1)
        # bell-curve weight: thick centre, thin ends
        w = max(1, int(base_width * (1.0 - abs(2 * t - 1) ** 1.5)))
        draw.line([pts[i], pts[i + 1]], fill=color, width=w)


def generate_signature_pixmap(name: str, width: int = 340,
                               height: int = 90) -> Optional[QPixmap]:
    """
    Render a pen-on-paper style signature: full cursive name sitting on a
    horizontal baseline, with a small tail flourish — similar to a notarised
    document signature.
    """
    if not _PIL_AVAILABLE or not name.strip():
        return None

    # ── Resampling helpers (Pillow ≥ 9 uses Image.Resampling enum) ────────────
    def _get_resample(attr):
        rs = getattr(Image, "Resampling", None)
        if rs and hasattr(rs, attr):
            return getattr(rs, attr)
        return getattr(Image, attr, 3)

    _bicubic = _get_resample("BICUBIC")
    _lanczos  = _get_resample("LANCZOS")

    # ── Work on 3× canvas for anti-aliasing, then downsample ─────────────────
    scale = 3
    cw, ch = width * scale, height * scale
    INK       = ( 7,  59, 166, 255)   # #073ba6 — ink blue
    INK_MID   = ( 7,  59, 166, 190)
    INK_THIN  = ( 7,  59, 166, 120)

    img  = Image.new("RGBA", (cw, ch), (255, 255, 255, 255))  # white bg — visible in DOCX/PDF
    draw = ImageDraw.Draw(img)

    # ── Font: try to load a natural-looking script ────────────────────────────
    # Prefer fonts that render as flowing cursive, not print-style
    font_size = int(ch * 0.38)
    font = _load_script_font(font_size)

    # ── Display text: full name (capitalise each word) ────────────────────────
    display = " ".join(w.capitalize() for w in name.strip().split())

    # ── Measure ───────────────────────────────────────────────────────────────
    try:
        bbox  = draw.textbbox((0, 0), display, font=font)
        tw    = bbox[2] - bbox[0]
        th    = bbox[3] - bbox[1]
        ty_off = bbox[1]
    except Exception:
        tw, th, ty_off = font_size * len(display) // 2, font_size, 0

    # Scale down if text wider than ~85 % of canvas
    max_tw = int(cw * 0.85)
    if tw > max_tw and tw > 0:
        ratio      = max_tw / tw
        font_size  = int(font_size * ratio)
        font       = _load_script_font(font_size)
        try:
            bbox   = draw.textbbox((0, 0), display, font=font)
            tw     = bbox[2] - bbox[0]; th = bbox[3] - bbox[1]; ty_off = bbox[1]
        except Exception:
            tw, th, ty_off = font_size * len(display) // 2, font_size, 0

    # ── Baseline sits at 72 % of canvas height ────────────────────────────────
    base_y = int(ch * 0.72)
    # Text is positioned so its visual bottom aligns with base_y
    tx = (cw - tw) // 2
    ty = base_y - th - ty_off

    # ── Draw baseline (full width, like a form line) ──────────────────────────
    line_x0 = int(cw * 0.04)
    line_x1 = int(cw * 0.96)
    line_w   = max(2, int(scale * 0.9))
    draw.line([(line_x0, base_y), (line_x1, base_y)], fill=(7, 59, 166, 200), width=line_w)

    # ── Subtle ink-shadow (depth effect) ─────────────────────────────────────
    draw.text((tx + 1, ty + 2), display, font=font, fill=(7, 59, 166, 35))
    # Main text pass
    draw.text((tx, ty), display, font=font, fill=INK)

    # ── Small end-flourish: a short curling tail to the right ─────────────────
    fl_x0 = tx + tw
    fl_y0 = base_y
    tail = _bezier_pts(
        (fl_x0,                    fl_y0),
        (fl_x0 + int(12 * scale),  fl_y0 - int(9 * scale)),
        (fl_x0 + int(22 * scale),  fl_y0 - int(5 * scale)),
        (fl_x0 + int(16 * scale),  fl_y0 + int(3 * scale)),
        steps=30
    )
    _draw_ink_stroke(draw, tail, INK_MID, base_width=max(1, int(scale * 0.8)))

    # tiny exit stroke going right-down
    exit_pts = _bezier_pts(
        tail[-1],
        (tail[-1][0] + int(6 * scale),  tail[-1][1] + int(2 * scale)),
        (tail[-1][0] + int(14 * scale), tail[-1][1] + int(1 * scale)),
        (tail[-1][0] + int(20 * scale), tail[-1][1] - int(1 * scale)),
        steps=20
    )
    _draw_ink_stroke(draw, exit_pts, INK_THIN, base_width=max(1, int(scale * 0.6)))

    # ── Very slight CCW tilt (−1 … −3°) ──────────────────────────────────────
    angle = random.uniform(-3, -1)
    rotated = img.rotate(angle, resample=_bicubic, expand=False)

    # ── Downsample to final size ──────────────────────────────────────────────
    final = rotated.resize((width, height), _lanczos)

    # ── Convert to QPixmap ────────────────────────────────────────────────────
    buf = io.BytesIO()
    # Convert to RGB (white bg) before saving — ensures DOCX/PDF always show the sig
    final_rgb = Image.new("RGB", final.size, (255, 255, 255))
    final_rgb.paste(final, mask=final.split()[3] if final.mode == "RGBA" else None)
    final_rgb.save(buf, format="PNG")
    buf.seek(0)
    pixmap = QPixmap()
    pixmap.loadFromData(QByteArray(buf.read()))
    return pixmap

def _pixmap_to_png_bytes(pixmap: QPixmap) -> bytes:
    ba = QByteArray()
    buf = QBuffer(ba)
    buf.open(QIODevice.WriteOnly)
    pixmap.save(buf, "PNG")
    buf.close()
    return bytes(ba)

class NoScrollComboBox(QComboBox):
    """QComboBox без реакції на колесо миші — запобігає випадковим змінам."""
    def wheelEvent(self, event):
        event.ignore()


class NoScrollSpinBox(QSpinBox):
    """QSpinBox без колеса миші."""
    def wheelEvent(self, event):
        event.ignore()


class NoScrollDoubleSpinBox(QDoubleSpinBox):
    """QDoubleSpinBox без колеса миші."""
    def wheelEvent(self, event):
        event.ignore()


# ── Signature Widget ──────────────────────────────────────────────────────────
class SignatureWidget(QWidget):
    changed = Signal()

    def __init__(self, party_label: str, parent=None):
        super().__init__(parent)
        self._party = party_label
        self._pixmap: Optional[QPixmap] = None
        self._name: str = ""
        self._build()

    def _build(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(4)

        lbl = QLabel(f"{self._party} signature:")
        lbl.setObjectName("section_label")
        lay.addWidget(lbl)

        self.preview = QLabel("[ no signature — generate or upload ]")
        self.preview.setObjectName("sig_preview")
        self.preview.setAlignment(Qt.AlignCenter)
        self.preview.setFixedHeight(68)
        self.preview.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.preview.setWordWrap(True)
        lay.addWidget(self.preview)

        row = QHBoxLayout()
        row.setSpacing(4)

        self.btn_gen = QPushButton("✍️  Generate")
        self.btn_gen.setObjectName("btn_sig_gen")
        self.btn_gen.setToolTip("Generate handwriting-style signature from name field")
        self.btn_gen.clicked.connect(self._on_generate)
        if not _PIL_AVAILABLE:
            self.btn_gen.setEnabled(False)
            self.btn_gen.setToolTip("pip install Pillow")
        row.addWidget(self.btn_gen)

        self.btn_upload = QPushButton("📁  Upload")
        self.btn_upload.setObjectName("btn_sig_upload")
        self.btn_upload.setToolTip("Upload real signature image (PNG/JPG)")
        self.btn_upload.clicked.connect(self._on_upload)
        row.addWidget(self.btn_upload)

        self.btn_clear = QPushButton("✕")
        self.btn_clear.setObjectName("btn_sig_gen")
        self.btn_clear.setFixedWidth(28)
        self.btn_clear.clicked.connect(self._on_clear)
        row.addWidget(self.btn_clear)

        lay.addLayout(row)

    def set_name(self, name: str):
        self._name = name

    def _show(self, px: QPixmap):
        self._pixmap = px
        self.preview.setPixmap(
            px.scaled(self.preview.width() - 4, self.preview.height() - 4,
                      Qt.KeepAspectRatio, Qt.SmoothTransformation)
        )
        self.preview.setText("")
        self.changed.emit()

    def _on_generate(self):
        name = self._name.strip()
        if not name:
            QMessageBox.warning(self, "Name Required",
                                f"Fill in the {self._party} name first.")
            return
        px = generate_signature_pixmap(name)
        if px and not px.isNull():
            self._show(px)
        else:
            QMessageBox.warning(self, "Error", "Signature generation failed. Is Pillow installed?")

    def _on_upload(self):
        path, _ = QFileDialog.getOpenFileName(
            self, f"Upload {self._party} Signature",
            "", "Images (*.png *.jpg *.jpeg *.bmp)"
        )
        if path:
            px = QPixmap(path)
            if not px.isNull():
                self._show(px)

    def _on_clear(self):
        self._pixmap = None
        self.preview.clear()
        self.preview.setText("[ no signature — generate or upload ]")
        self.changed.emit()

    def get_pixmap(self) -> Optional[QPixmap]:
        return self._pixmap

    def has_signature(self) -> bool:
        return self._pixmap is not None and not self._pixmap.isNull()


# ── Stamp Widget ─────────────────────────────────────────────────────────────
class StampWidget(QWidget):
    changed = Signal()

    def __init__(self, party_label: str, parent=None):
        super().__init__(parent)
        self._party   = party_label
        self._pixmap: Optional[QPixmap] = None
        self._name: str = ""
        self._build()

    def _build(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(4)

        lbl = QLabel(f"{self._party} stamp / seal:")
        lbl.setObjectName("section_label")
        lay.addWidget(lbl)

        self.preview = QLabel("[ no stamp — generate or upload ]")
        self.preview.setObjectName("sig_preview")
        self.preview.setAlignment(Qt.AlignCenter)
        self.preview.setFixedHeight(90)
        self.preview.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.preview.setWordWrap(True)
        lay.addWidget(self.preview)

        row = QHBoxLayout()
        row.setSpacing(4)

        self.btn_gen = QPushButton("🔵  Generate")
        self.btn_gen.setObjectName("btn_sig_gen")
        self.btn_gen.setToolTip("Generate a round company stamp from the name field")
        self.btn_gen.clicked.connect(self._on_generate)
        if not _PIL_AVAILABLE:
            self.btn_gen.setEnabled(False)
            self.btn_gen.setToolTip("pip install Pillow")
        row.addWidget(self.btn_gen)

        self.btn_upload = QPushButton("📁  Upload")
        self.btn_upload.setObjectName("btn_sig_upload")
        self.btn_upload.clicked.connect(self._on_upload)
        row.addWidget(self.btn_upload)

        self.btn_clear = QPushButton("✕")
        self.btn_clear.setObjectName("btn_sig_gen")
        self.btn_clear.setFixedWidth(28)
        self.btn_clear.clicked.connect(self._on_clear)
        row.addWidget(self.btn_clear)
        lay.addLayout(row)

    def set_name(self, name: str):
        self._name = name

    def _show(self, px: QPixmap):
        self._pixmap = px
        self.preview.setPixmap(
            px.scaled(self.preview.height() - 4, self.preview.height() - 4,
                      Qt.KeepAspectRatio, Qt.SmoothTransformation)
        )
        self.preview.setText("")
        self.changed.emit()

    def _on_generate(self):
        name = self._name.strip()
        if not name:
            QMessageBox.warning(self, "Name Required",
                                f"Fill in the {self._party} name first.")
            return
        px = generate_stamp_pixmap(name)
        if px and not px.isNull():
            self._show(px)
        else:
            QMessageBox.warning(self, "Error", "Stamp generation failed. Is Pillow installed?")

    def _on_upload(self):
        path, _ = QFileDialog.getOpenFileName(
            self, f"Upload {self._party} Stamp",
            "", "Images (*.png *.jpg *.jpeg *.bmp)"
        )
        if path:
            px = QPixmap(path)
            if not px.isNull():
                self._show(px)

    def _on_clear(self):
        self._pixmap = None
        self.preview.clear()
        self.preview.setText("[ no stamp — generate or upload ]")
        self.changed.emit()

    def get_pixmap(self) -> Optional[QPixmap]:
        return self._pixmap

    def has_stamp(self) -> bool:
        return self._pixmap is not None and not self._pixmap.isNull()


# ── Logo Widget ─────────────────────────────────────────────────────────────────
class LogoWidget(QWidget):
    changed = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self._path: Optional[str] = None
        self._pixmap: Optional[QPixmap] = None
        self._build()

    def _build(self):
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(6)

        self.preview = QLabel()
        self.preview.setFixedSize(60, 40)
        self.preview.setAlignment(Qt.AlignCenter)
        self.preview.setStyleSheet(
            "border: 1px dashed #585b70; border-radius: 4px; background: #181825;"
        )
        self.preview.setText("🖼️")
        lay.addWidget(self.preview)

        btns = QVBoxLayout()
        btns.setSpacing(4)
        self.btn_upload = QPushButton("📁  Upload Logo")
        self.btn_upload.setObjectName("btn_sig_upload")
        self.btn_upload.clicked.connect(self._on_upload)
        btns.addWidget(self.btn_upload)
        self.btn_clear = QPushButton("✕  Clear")
        self.btn_clear.setObjectName("btn_sig_gen")
        self.btn_clear.clicked.connect(self._on_clear)
        btns.addWidget(self.btn_clear)
        lay.addLayout(btns)

    def _on_upload(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Upload Client Logo", "", "Images (*.png *.jpg *.jpeg *.bmp *.svg)"
        )
        if path:
            px = QPixmap(path)
            if not px.isNull():
                self._path = path
                self._pixmap = px
                self.preview.setPixmap(
                    px.scaled(56, 36, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                )
                self.changed.emit()

    def _on_clear(self):
        self._path = None
        self._pixmap = None
        self.preview.clear()
        self.preview.setText("🖼️")
        self.changed.emit()

    def get_path(self) -> Optional[str]:
        return self._path

    def get_pixmap(self) -> Optional[QPixmap]:
        return self._pixmap

    def has_logo(self) -> bool:
        return self._path is not None and os.path.exists(self._path)


# ── Workers ───────────────────────────────────────────────────────────────────
class SiteParseWorker(QThread):
    finished = Signal(dict)
    error = Signal(str)

    def __init__(self, url: str, parent=None):
        super().__init__(parent)
        self._url = url

    def run(self):
        if not _PROFILER_AVAILABLE:
            self.error.emit("core.site_profiler not available")
            return
        try:
            self.finished.emit(analyze_site(self._url))
        except Exception as e:
            self.error.emit(str(e))


class OpenAIWorker(QThread):
    delta = Signal(str)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, api_key: str, prompt: str, parent=None):
        super().__init__(parent)
        self._key = api_key
        self._prompt = prompt

    def run(self):
        if not _OPENAI_AVAILABLE:
            self.error.emit("openai package not installed")
            return
        try:
            client = _OpenAI(api_key=self._key)
            full = ""
            with client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": self._prompt}],
                stream=True, max_tokens=4096,
            ) as stream:
                for chunk in stream:
                    d = chunk.choices[0].delta.content or ""
                    if d:
                        full += d
                        self.delta.emit(d)
            self.finished.emit(full)
        except Exception as e:
            self.error.emit(str(e))


# ─────────────────────────────────────────────────────────────────────────────
# WEBSITE BUSINESS SUMMARY PARSER — prompt, workers, helpers, widget
# ─────────────────────────────────────────────────────────────────────────────

SUMMARY_MASTER_PROMPT = """\
You are a Google Ads Business Operations verification assistant.

Your task is to generate a factual commercial activity summary based only on the
parsed website data provided. Use the parser data as the source of truth.
Do not invent missing facts. Do not overclaim ownership, licenses, partnerships,
trademarks, phone numbers, addresses, or legal registration.

The summary must help explain what the website is about, what business or service
is presented, who the target audience is, how users interact with the website,
who appears responsible for the content and service delivery, what trust signals
are visible, and what information is missing.

Avoid technical advertising terms. Avoid mentioning tracking, redirects, cloaking,
traffic routing, prelands, offers, or affiliate infrastructure.
Write in clear business English.

PARSER DATA:
{parser_data}

Generate EXACTLY the following sections:

## 1. Website Identity
Website: {final_url}
Detected brand: {brand_name_detected}
Detected company/legal entity: {company_name_detected}

Note any brand/domain mismatch briefly if applicable.
If company name is not visible write: [NEED INPUT: company name]
If legal entity is not visible write: [NEED INPUT: legal entity]

## 2. Business Type
Describe the website business category using visible content only.

## 3. What the Website Provides
Summarize visible services, products, or informational content in 2-4 sentences.

## 4. Target Audience
Describe who the website appears to target using neutral wording only.

## 5. User Interaction Model
Explain how users interact: contact form / registration / booking / quote request /
email or phone / reading only.
If unclear write: [NEED INPUT: confirm how users interact with the business]

## 6. Product or Service Delivery
Explain whether the website provides services directly or mainly collects inquiries.
If unclear write: [NEED INPUT: confirm who provides the service]

## 7. Content Ownership / Management
Based on visible data explain who appears to manage website content.
If not visible write: [NEED INPUT: confirm who creates and manages website content]

## 8. Trust Signals Detected
List only elements actually detected from: Privacy Policy, Terms, Contact page,
Email, Phone, Address, About page, Registration/company info,
License/certification references.

## 9. Missing or Weak Trust Signals
List missing elements factually.
Example: "Phone number was not detected." / "Physical address was not detected."

## 10. Commercial Activity Summary
Write a polished 3-5 sentence summary suitable for Google Ads Business Operations
verification. Use [NEED INPUT] placeholders for missing facts.
Template:
"[Company/brand] operates a website that presents [service/product/content].
The website is used to provide users with information about [topic] and allows
them to [contact/register/request/book]. The target audience is users interested
in [detected topic]. Based on the visible website content the business activity is
focused on [summary]."

## 11. Suggested Form Answers

A. Please describe the company's type of business.
[draft answer]

B. Please describe the company's business model.
[draft answer]

C. Who are your customers / target audience?
[draft answer]

D. Please describe how your company interacts with your target audience.
[draft answer]

E. Describe how the target audience receives the goods or services promoted in your ads.
[draft answer]

F. Who creates the content of your website?
[draft answer]

G. Who would a customer hold responsible for non-fulfillment?
[draft answer]

H. How does your company protect customer personal information?
[draft answer]

I. Other comments or information.
[draft answer]

STRICT RULES:
- Do not invent facts.
- If company name not visible: [NEED INPUT: company name]
- If legal entity not visible: [NEED INPUT: legal entity]
- If contact data missing say it is missing.
- If privacy policy not detected say it is not detected.
- Do not claim licenses, certifications, partnerships, trademarks unless detected.
- Do not describe redirects, tracking, cloaking, affiliate infrastructure.
- Focus on user-facing website content only.
- Use neutral business language.
- Keep suitable for Google Ads verification.
"""


class FullSiteParseWorker(QThread):
    """Scrapes a URL and returns all 25 structured fields for the business summary."""
    finished = Signal(dict)
    error = Signal(str)

    def __init__(self, url: str, parent=None):
        super().__init__(parent)
        self._url = url

    def run(self):
        try:
            import requests as _req
            from bs4 import BeautifulSoup as _BS
            import urllib.parse as _up
            import re as _re2
            import json as _json

            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0.0.0 Safari/537.36"
                )
            }
            session = _req.Session()
            resp = session.get(self._url, headers=headers, timeout=20,
                               allow_redirects=True)

            redirect_chain = [str(r.url) for r in resp.history] + [str(resp.url)]
            final_url = str(resp.url)
            parsed_url = _up.urlparse(final_url)
            domain = parsed_url.netloc

            soup = _BS(resp.text, "html.parser")

            # ── Basic metadata ──────────────────────────────────────────────
            title_tag = soup.find("title")
            site_title = title_tag.get_text(strip=True) if title_tag else ""

            meta_desc_tag = soup.find("meta", attrs={"name": "description"})
            meta_description = (
                (meta_desc_tag.get("content") or "").strip()
                if meta_desc_tag else ""
            )

            h1_tag = soup.find("h1")
            h1 = h1_tag.get_text(strip=True) if h1_tag else ""

            h2_headings = [h.get_text(strip=True) for h in soup.find_all("h2")][:8]

            # ── Navigation ──────────────────────────────────────────────────
            nav_tag = (
                soup.find("nav") or
                soup.find(attrs={"role": "navigation"}) or
                soup.find(id=_re2.compile(r"nav|menu", _re2.I)) or
                soup.find(class_=_re2.compile(r"nav|menu", _re2.I))
            )
            navigation_items = []
            if nav_tag:
                navigation_items = [
                    a.get_text(strip=True)
                    for a in nav_tag.find_all("a")
                    if a.get_text(strip=True)
                ][:12]

            # ── Visible text ────────────────────────────────────────────────
            soup_text = _BS(resp.text, "html.parser")
            for tag in soup_text(["script", "style", "noscript", "meta", "link"]):
                tag.decompose()
            visible_text_summary = " ".join(
                soup_text.get_text(" ", strip=True).split()
            )[:1500]

            # ── CTAs ────────────────────────────────────────────────────────
            cta_kw = [
                "contact us", "get started", "request a quote", "sign up",
                "register", "book", "buy now", "order", "subscribe", "download",
                "try free", "free trial", "start now", "learn more", "get a quote",
                "schedule", "apply now", "get in touch",
            ]
            cta_buttons = []
            for el in soup.find_all(["button", "a"]):
                txt = el.get_text(strip=True)
                if txt and any(kw in txt.lower() for kw in cta_kw) and len(txt) < 80:
                    cta_buttons.append(txt)
            call_to_action_buttons = list(dict.fromkeys(cta_buttons))[:8]

            # ── Trust page detection ────────────────────────────────────────
            all_hrefs = [
                (a.get("href") or "").lower()
                for a in soup.find_all("a", href=True)
            ]
            all_link_texts = [
                a.get_text(strip=True).lower()
                for a in soup.find_all("a")
            ]
            combined = all_hrefs + all_link_texts

            contact_page_detected = any("contact" in s for s in combined)
            privacy_policy_detected = any("privacy" in s for s in combined)
            terms_page_detected = any(
                t in s for s in combined
                for t in ["terms", "/tos", "conditions"]
            )

            # ── Contact info ────────────────────────────────────────────────
            raw_html = resp.text
            emails = _re2.findall(
                r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', raw_html
            )
            emails = [
                e for e in emails
                if not any(x in e for x in
                           ["woff", ".png", ".jpg", ".gif", "example.com",
                            "schema.org", "sentry", "w3.org"])
            ]
            contact_email = emails[0] if emails else ""

            tel_links = [
                a.get("href", "")
                for a in soup.find_all("a", href=_re2.compile(r"^tel:", _re2.I))
            ]
            if tel_links:
                phone = tel_links[0].replace("tel:", "").strip()
            else:
                ph = _re2.findall(
                    r'(?:\+?\d[\d\s\-\(\)\.]{6,18}\d)', visible_text_summary
                )
                phone = ph[0].strip() if ph else ""

            addr_tag = (
                soup.find(attrs={"itemprop": "address"}) or
                soup.find("address")
            )
            address = addr_tag.get_text(" ", strip=True)[:200] if addr_tag else ""

            # ── Brand / company ─────────────────────────────────────────────
            og_site = soup.find("meta", property="og:site_name")
            brand_name_detected = (
                (og_site.get("content") or "").strip()
                if og_site else
                site_title.split("|")[0].split(" - ")[0].strip()
            )

            company_kw = [
                "LLC", "Ltd.", "Ltd", "Limited", "Inc.", "Inc", "Corp.", "Corp",
                "GmbH", "SRL", "UAB", "OÜ", "JSC", "PJSC", "PLC", "S.A.", "SA",
            ]
            company_name_detected = ""
            for kw in company_kw:
                m = _re2.search(
                    r'(?:[\w\s&,\.]{2,50}?)' + _re2.escape(kw) + r'\b',
                    raw_html
                )
                if m:
                    candidate = m.group(0).strip()
                    if 4 < len(candidate) < 100:
                        company_name_detected = candidate
                        break

            # ── Products / niche ────────────────────────────────────────────
            product_or_service_detected = "; ".join(
                filter(None, [h1] + h2_headings[:3])
            )[:300]

            niche_map = {
                "e-commerce / retail": [
                    "shop", "store", "cart", "checkout", "product", "price"],
                "saas / software": [
                    "software", "platform", "dashboard", "api", "integration", "saas"],
                "services / agency": [
                    "service", "consulting", "agency", "solution", "expert"],
                "health / medical": [
                    "health", "medical", "clinic", "doctor", "wellness", "pharmacy"],
                "finance": [
                    "finance", "loan", "credit", "investment", "insurance", "trading"],
                "education": [
                    "course", "learn", "education", "training", "certificate", "academy"],
                "real estate": [
                    "property", "real estate", "rent", "lease", "apartment", "listing"],
                "travel / hospitality": [
                    "hotel", "booking", "tour", "travel", "flight", "resort"],
                "news / media / content": [
                    "news", "article", "blog", "media", "content", "editorial"],
            }
            vt_lower = visible_text_summary.lower()
            niche_detected = "general / informational"
            best_score = 0
            for niche, kws in niche_map.items():
                score = sum(1 for kw in kws if kw in vt_lower)
                if score > best_score:
                    best_score = score
                    niche_detected = niche

            # ── Language / geo ──────────────────────────────────────────────
            html_root = soup.find("html")
            language = (html_root.get("lang") or "") if html_root else ""
            if not language:
                meta_ct = soup.find(
                    "meta", attrs={"http-equiv": "Content-Language"}
                )
                language = (meta_ct.get("content") or "") if meta_ct else ""
            language = language or "not detected"

            geo_hints = []
            tld = domain.split(".")[-1].upper()
            if tld not in ("COM", "NET", "ORG", "IO", "CO", "AI", "APP", "DEV"):
                geo_hints.append(f".{tld.lower()} TLD")
            geo_meta = soup.find(
                "meta", attrs={"name": _re2.compile(r"geo\.", _re2.I)}
            )
            if geo_meta:
                geo_hints.append(geo_meta.get("content", ""))
            country_or_geo_hint = ", ".join(filter(None, geo_hints)) or "not detected"

            # ── External links ──────────────────────────────────────────────
            ext_links = [
                a["href"] for a in soup.find_all("a", href=True)
                if a["href"].startswith("http") and domain not in a["href"]
            ]
            external_links = list(dict.fromkeys(ext_links))[:10]

            # ── Forms / login ───────────────────────────────────────────────
            forms_detected = len(soup.find_all("form")) > 0
            login_kw = [
                "login", "sign in", "log in", "register", "sign up",
                "create account", "my account", "member area",
            ]
            login_or_registration_detected = any(kw in vt_lower for kw in login_kw)

            # ── Run LandingAnalyzer for deep policy checks ──────────────────
            analyzer_score = 0
            analyzer_findings = []
            analyzer_business_info = {}
            analyzer_appeal_risk = {}
            robots_txt_status = "not checked"
            noindex_detected = False
            subpages_count = 0
            try:
                try:
                    from landing_analyzer.core.analyzer import LandingAnalyzer as _LA
                except ImportError:
                    from core.analyzer import LandingAnalyzer as _LA
                _analyzer = _LA(timeout=20)
                _result = _analyzer.analyze(self._url)
                analyzer_score = _result.score
                analyzer_business_info = _result.business_info or {}
                analyzer_appeal_risk = _result.appeal_risk or {}
                subpages_count = len(getattr(_analyzer, '_subpages', {}))
                for f in _result.findings:
                    entry = {
                        "severity": f.severity.value,
                        "category": f.category.value,
                        "title": f.title,
                        "detail": f.detail[:300],
                        "recommendation_ua": f.recommendation_ua,
                    }
                    analyzer_findings.append(entry)
                    if "robots.txt" in f.title:
                        robots_txt_status = f.title
                    if "noindex" in f.title.lower():
                        noindex_detected = True
            except Exception:
                pass

            self.finished.emit({
                "final_url": final_url,
                "domain": domain,
                "site_title": site_title,
                "meta_description": meta_description,
                "h1": h1,
                "h2_headings": h2_headings,
                "navigation_items": navigation_items,
                "visible_text_summary": visible_text_summary,
                "call_to_action_buttons": call_to_action_buttons,
                "contact_page_detected": contact_page_detected,
                "privacy_policy_detected": privacy_policy_detected,
                "terms_page_detected": terms_page_detected,
                "contact_email": contact_email,
                "phone": phone,
                "address": address,
                "brand_name_detected": brand_name_detected,
                "company_name_detected": company_name_detected,
                "product_or_service_detected": product_or_service_detected,
                "niche_detected": niche_detected,
                "language": language,
                "country_or_geo_hint": country_or_geo_hint,
                "redirect_chain": redirect_chain,
                "external_links": external_links,
                "forms_detected": forms_detected,
                "login_or_registration_detected": login_or_registration_detected,
                # ── LandingAnalyzer enrichment ──
                "analyzer_score": analyzer_score,
                "analyzer_findings": analyzer_findings,
                "analyzer_business_info": analyzer_business_info,
                "analyzer_appeal_risk": analyzer_appeal_risk,
                "robots_txt_status": robots_txt_status,
                "noindex_detected": noindex_detected,
                "subpages_crawled": subpages_count,
            })

        except Exception as exc:
            self.error.emit(str(exc))


class SummaryAIWorker(QThread):
    """Calls GPT-4o to generate the business summary from parsed site data."""
    delta = Signal(str)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, api_key: str, parsed_data: dict, parent=None):
        super().__init__(parent)
        self._key = api_key
        self._data = parsed_data

    def run(self):
        if not _OPENAI_AVAILABLE:
            self.error.emit("openai package not installed — pip install openai")
            return
        try:
            import json as _json
            pd = self._data
            # Exclude verbose visible_text from the JSON block; add separately
            slim = {k: v for k, v in pd.items() if k != "visible_text_summary"}
            parser_data_str = (
                _json.dumps(slim, indent=2, ensure_ascii=False)
                + f'\n\n"visible_text_summary":\n{pd.get("visible_text_summary", "")}'
            )
            prompt = SUMMARY_MASTER_PROMPT.format(
                parser_data=parser_data_str,
                final_url=pd.get("final_url", ""),
                brand_name_detected=pd.get("brand_name_detected") or "[not detected]",
                company_name_detected=(
                    pd.get("company_name_detected") or "[NEED INPUT: company name]"
                ),
            )
            client = _OpenAI(api_key=self._key)
            full = ""
            with client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                stream=True,
                max_tokens=3500,
            ) as stream:
                for chunk in stream:
                    d = chunk.choices[0].delta.content or ""
                    if d:
                        full += d
                        self.delta.emit(d)
            self.finished.emit(full)
        except Exception as exc:
            self.error.emit(str(exc))


def _build_risk_report(data: dict) -> str:
    """Generate a risk / missing inputs report from parsed data (no AI needed)."""
    lines = ["═" * 64, "  RISK & MISSING INPUTS REPORT", "═" * 64, ""]

    # ── Missing trust signals ────────────────────────────────────────────────
    missing = []
    if not data.get("company_name_detected"):
        missing.append("❌  Company / legal entity name — NOT DETECTED")
    if not data.get("contact_email"):
        missing.append("❌  Contact email — NOT DETECTED")
    if not data.get("phone"):
        missing.append("❌  Phone number — NOT DETECTED")
    if not data.get("address"):
        missing.append("❌  Physical address — NOT DETECTED")
    if not data.get("privacy_policy_detected"):
        missing.append("❌  Privacy Policy page — NOT DETECTED")
    if not data.get("terms_page_detected"):
        missing.append("❌  Terms of Service page — NOT DETECTED")
    if not data.get("contact_page_detected"):
        missing.append("❌  Contact page — NOT DETECTED")

    if missing:
        lines.append("MISSING TRUST SIGNALS:")
        lines.extend(missing)
    else:
        lines.append("✅  All basic trust signals detected.")
    lines.append("")

    # ── Brand / domain mismatch ──────────────────────────────────────────────
    brand = data.get("brand_name_detected", "")
    domain = data.get("domain", "")
    if brand and domain:
        brand_slug = re.sub(r"[^a-z0-9]", "", brand.lower())
        domain_clean = re.sub(r"[^a-z0-9]", "", domain.lower().split(".")[0])
        if brand_slug and domain_clean:
            if brand_slug not in domain_clean and domain_clean not in brand_slug:
                lines.append(
                    f"⚠️  BRAND / DOMAIN MISMATCH DETECTED\n"
                    f"    Brand detected: «{brand}»\n"
                    f"    Domain:         «{domain}»\n"
                    f"    → Confirm whether this is intentional or a separate entity."
                )
                lines.append("")

    # ── Redirect chain ───────────────────────────────────────────────────────
    chain = data.get("redirect_chain", [])
    if len(chain) > 2:
        lines.append(f"⚠️  REDIRECT CHAIN — {len(chain) - 1} hop(s):")
        for step in chain:
            lines.append(f"    → {step}")
        lines.append("")

    # ── Login / registration ─────────────────────────────────────────────────
    if data.get("login_or_registration_detected"):
        lines.append(
            "ℹ️  LOGIN / REGISTRATION DETECTED\n"
            "    → Confirm whether user accounts are required for service delivery."
        )
        lines.append("")

    # ── Forms ────────────────────────────────────────────────────────────────
    if data.get("forms_detected"):
        lines.append(
            "ℹ️  FORMS DETECTED\n"
            "    → Confirm purpose: lead capture / registration / purchase / support."
        )
        lines.append("")

    # ── Action items ─────────────────────────────────────────────────────────
    lines.append("─" * 64)
    lines.append("ACTION ITEMS — confirm before Google Ads verification submission:")
    lines.append("")
    actions = []
    if not data.get("company_name_detected"):
        actions.append("• Provide full legal company name and entity type")
    if not data.get("contact_email"):
        actions.append("• Add or confirm a visible contact email address")
    if not data.get("phone"):
        actions.append("• Add or confirm a visible business phone number")
    if not data.get("address"):
        actions.append("• Add or confirm a visible physical/business address")
    if not data.get("privacy_policy_detected"):
        actions.append("• Add a Privacy Policy page before submitting verification")
    if not data.get("terms_page_detected"):
        actions.append("• Consider adding Terms of Service / Terms of Use page")
    if not actions:
        actions.append("• No critical actions required — review summary for accuracy")
    lines.extend(actions)
    lines.append("")
    lines.append("═" * 64)
    return "\n".join(lines)


# ── Website Business Summary Widget ──────────────────────────────────────────
class WebsiteSummaryWidget(QWidget):
    """
    Standalone module: Website Business Summary Parser.
    Scans a URL, collects 25 structured data fields, generates a
    Google Ads commercial activity summary via GPT-4o.
    Emits send_to_agreement(dict) to fill the Agreement Generator.
    """
    send_to_agreement = Signal(dict)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._parsed_data: Optional[dict] = None
        self._scan_worker: Optional[QThread] = None
        self._ai_worker: Optional[QThread] = None
        self._ownership: str = ""   # "own" | "partner" | ""
        self._build_ui()

    # ── Build UI ──────────────────────────────────────────────────────────────
    def _build_ui(self):
        outer = QHBoxLayout(self)
        outer.setContentsMargins(10, 10, 10, 10)
        outer.setSpacing(8)

        splitter = QSplitter(Qt.Horizontal)
        outer.addWidget(splitter)

        # ── Left control panel ────────────────────────────────────────────────
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFixedWidth(375)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(4, 4, 8, 4)
        ll.setSpacing(8)
        scroll.setWidget(left)

        # Title
        title_lbl = QLabel("🌐  Website Business Summary Parser")
        title_lbl.setFont(QFont("Segoe UI", 13, QFont.Bold))
        title_lbl.setStyleSheet("color: #89b4fa; padding-bottom: 2px;")
        title_lbl.setWordWrap(True)
        ll.addWidget(title_lbl)

        sub_lbl = QLabel(
            "Generate structured commercial activity summaries\n"
            "for Google Ads Business Operations verification.\n\n"
            "This module does not create contracts. It only\n"
            "prepares factual business summaries from parsed data."
        )
        sub_lbl.setStyleSheet("color: #6c7086; font-size: 11px;")
        sub_lbl.setWordWrap(True)
        ll.addWidget(sub_lbl)

        # ── Scan group ────────────────────────────────────────────────────────
        grp_scan = QGroupBox("🔍  Scan Website")
        grp_scan_lay = QVBoxLayout(grp_scan)
        grp_scan_lay.setSpacing(5)

        row_url = QHBoxLayout()
        self.inp_scan_url = QLineEdit()
        self.inp_scan_url.setPlaceholderText("https://example.com")
        row_url.addWidget(self.inp_scan_url)
        self.btn_scan = QPushButton("🔍 Scan")
        self.btn_scan.setObjectName("btn_scan")
        self.btn_scan.setFixedWidth(72)
        self.btn_scan.clicked.connect(self._on_scan)
        row_url.addWidget(self.btn_scan)
        grp_scan_lay.addLayout(row_url)

        self.lbl_scan_status = QLabel("Enter a URL and click Scan.")
        self.lbl_scan_status.setStyleSheet("color: #6c7086; font-size: 11px;")
        self.lbl_scan_status.setWordWrap(True)
        grp_scan_lay.addWidget(self.lbl_scan_status)
        ll.addWidget(grp_scan)

        # ── AI Summary group ──────────────────────────────────────────────────
        grp_ai = QGroupBox("🤖  AI Summary (GPT-4o)")
        grp_ai_lay = QVBoxLayout(grp_ai)
        grp_ai_lay.setSpacing(5)

        ai_note = QLabel(
            "Provide an OpenAI API key to generate the full\n"
            "business summary and form answers via GPT-4o."
        )
        ai_note.setStyleSheet("color: #6c7086; font-size: 11px;")
        ai_note.setWordWrap(True)
        grp_ai_lay.addWidget(ai_note)

        self.inp_sum_api_key = QLineEdit()
        self.inp_sum_api_key.setPlaceholderText("sk-...")
        self.inp_sum_api_key.setEchoMode(QLineEdit.Password)
        grp_ai_lay.addWidget(self.inp_sum_api_key)

        self.btn_gen_summary = QPushButton("🤖  Generate AI Summary")
        self.btn_gen_summary.setObjectName("btn_gen_summary")
        self.btn_gen_summary.setEnabled(False)
        self.btn_gen_summary.clicked.connect(self._on_gen_summary)
        if not _OPENAI_AVAILABLE:
            self.btn_gen_summary.setToolTip("pip install openai")
        grp_ai_lay.addWidget(self.btn_gen_summary)
        ll.addWidget(grp_ai)

        # ── Classification group ──────────────────────────────────────────────
        grp_cls = QGroupBox("🏷️  Website Classification")
        grp_cls_lay = QVBoxLayout(grp_cls)
        grp_cls_lay.setSpacing(6)

        cls_note = QLabel(
            "Classify the website before sending to the\nAgreement Generator."
        )
        cls_note.setStyleSheet("color: #6c7086; font-size: 11px;")
        cls_note.setWordWrap(True)
        grp_cls_lay.addWidget(cls_note)

        self.btn_mark_own = QPushButton("🏢  Mark as Own Website")
        self.btn_mark_own.setObjectName("btn_mark_own")
        self.btn_mark_own.clicked.connect(lambda: self._set_ownership("own"))
        grp_cls_lay.addWidget(self.btn_mark_own)

        self.btn_mark_partner = QPushButton("🤝  Mark as Partner Website")
        self.btn_mark_partner.setObjectName("btn_mark_partner")
        self.btn_mark_partner.clicked.connect(lambda: self._set_ownership("partner"))
        grp_cls_lay.addWidget(self.btn_mark_partner)

        self.lbl_ownership = QLabel("[ not classified ]")
        self.lbl_ownership.setStyleSheet("font-size: 11px; color: #585b70;")
        self.lbl_ownership.setAlignment(Qt.AlignCenter)
        grp_cls_lay.addWidget(self.lbl_ownership)
        ll.addWidget(grp_cls)

        # ── Send to Agreement ─────────────────────────────────────────────────
        sep1 = QFrame()
        sep1.setFrameShape(QFrame.HLine)
        sep1.setStyleSheet("color: #45475a;")
        ll.addWidget(sep1)

        send_note = QLabel(
            "Sends detected brand to the Agreement Generator\n"
            "Contractor field. Does not auto-generate a contract."
        )
        send_note.setStyleSheet("color: #6c7086; font-size: 11px;")
        send_note.setWordWrap(True)
        ll.addWidget(send_note)

        self.btn_send = QPushButton("📤  Send to Agreement Generator")
        self.btn_send.setObjectName("btn_send_agreement")
        self.btn_send.setEnabled(False)
        self.btn_send.clicked.connect(self._on_send_to_agreement)
        ll.addWidget(self.btn_send)

        ll.addStretch()

        # ── Right output tabs ─────────────────────────────────────────────────
        right_container = QWidget()
        rt_lay = QVBoxLayout(right_container)
        rt_lay.setContentsMargins(0, 0, 0, 0)
        rt_lay.setSpacing(0)

        self._out_tabs = QTabWidget()
        rt_lay.addWidget(self._out_tabs)

        # Tab 0 — Raw Parser Data
        raw_tab = QWidget()
        raw_lay = QVBoxLayout(raw_tab)
        raw_lay.setContentsMargins(6, 6, 6, 6)
        raw_lay.setSpacing(4)
        raw_hdr = QHBoxLayout()
        raw_hdr.addWidget(self._section_lbl("📊  Raw Parser Data (25 fields)"))
        raw_hdr.addStretch()
        btn_copy_raw = QPushButton("📋 Copy")
        btn_copy_raw.setObjectName("btn_copy")
        btn_copy_raw.setFixedWidth(72)
        btn_copy_raw.clicked.connect(lambda: self._copy_tab(0))
        raw_hdr.addWidget(btn_copy_raw)
        raw_lay.addLayout(raw_hdr)
        self.txt_raw = QTextEdit()
        self.txt_raw.setReadOnly(True)
        self.txt_raw.setFont(QFont("Consolas", 10))
        self.txt_raw.setPlaceholderText("Raw parsed data will appear here after scanning…")
        raw_lay.addWidget(self.txt_raw)
        self._out_tabs.addTab(raw_tab, "📊 Raw Data")

        # Tab 1 — Business Summary (sections 1-10)
        sum_tab = QWidget()
        sum_lay = QVBoxLayout(sum_tab)
        sum_lay.setContentsMargins(6, 6, 6, 6)
        sum_lay.setSpacing(4)
        sum_hdr = QHBoxLayout()
        sum_hdr.addWidget(self._section_lbl(
            "📝  Website Business Summary  (editable)"))
        sum_hdr.addStretch()
        self.lbl_sum_badge = QLabel("")
        self.lbl_sum_badge.setStyleSheet("font-size: 11px; color: #a6e3a1;")
        sum_hdr.addWidget(self.lbl_sum_badge)
        btn_copy_sum = QPushButton("📋 Copy")
        btn_copy_sum.setObjectName("btn_copy")
        btn_copy_sum.setFixedWidth(72)
        btn_copy_sum.clicked.connect(lambda: self._copy_tab(1))
        sum_hdr.addWidget(btn_copy_sum)
        sum_lay.addLayout(sum_hdr)
        self.txt_summary = QTextEdit()
        self.txt_summary.setReadOnly(False)
        self.txt_summary.setFont(QFont("Segoe UI", 11))
        self.txt_summary.setPlaceholderText(
            "Business summary (sections 1–10) will appear here after AI generation.\n\n"
            "This field is editable — review and adjust before use."
        )
        sum_lay.addWidget(self.txt_summary)
        self._out_tabs.addTab(sum_tab, "📝 Business Summary")

        # Tab 2 — Form Answers (section 11, A-I)
        form_tab = QWidget()
        form_lay = QVBoxLayout(form_tab)
        form_lay.setContentsMargins(6, 6, 6, 6)
        form_lay.setSpacing(4)
        form_hdr = QHBoxLayout()
        form_hdr.addWidget(self._section_lbl(
            "📋  Suggested Form Answers A–I  (editable)"))
        form_hdr.addStretch()
        btn_copy_form = QPushButton("📋 Copy")
        btn_copy_form.setObjectName("btn_copy")
        btn_copy_form.setFixedWidth(72)
        btn_copy_form.clicked.connect(lambda: self._copy_tab(2))
        form_hdr.addWidget(btn_copy_form)
        form_lay.addLayout(form_hdr)
        self.txt_form = QTextEdit()
        self.txt_form.setReadOnly(False)
        self.txt_form.setFont(QFont("Segoe UI", 11))
        self.txt_form.setPlaceholderText(
            "Suggested answers for Google Ads Business Operations form fields A–I\n"
            "will appear here after AI generation.\n\n"
            "This field is editable — review and adjust before use."
        )
        form_lay.addWidget(self.txt_form)
        self._out_tabs.addTab(form_tab, "📋 Form Answers")

        # Tab 3 — Risk & Missing Inputs
        risk_tab = QWidget()
        risk_lay = QVBoxLayout(risk_tab)
        risk_lay.setContentsMargins(6, 6, 6, 6)
        risk_lay.setSpacing(4)
        risk_hdr = QHBoxLayout()
        risk_hdr.addWidget(self._section_lbl("⚠️  Risk & Missing Inputs"))
        risk_hdr.addStretch()
        btn_copy_risk = QPushButton("📋 Copy")
        btn_copy_risk.setObjectName("btn_copy")
        btn_copy_risk.setFixedWidth(72)
        btn_copy_risk.clicked.connect(lambda: self._copy_tab(3))
        risk_hdr.addWidget(btn_copy_risk)
        risk_lay.addLayout(risk_hdr)
        self.txt_risk = QTextEdit()
        self.txt_risk.setReadOnly(True)
        self.txt_risk.setFont(QFont("Consolas", 10))
        self.txt_risk.setPlaceholderText(
            "Automated risk analysis and missing inputs report\n"
            "will appear here after scanning…"
        )
        risk_lay.addWidget(self.txt_risk)
        self._out_tabs.addTab(risk_tab, "⚠️ Risk & Missing")

        splitter.addWidget(scroll)
        splitter.addWidget(right_container)
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        splitter.setSizes([375, 905])

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _section_lbl(self, text: str) -> QLabel:
        l = QLabel(text)
        l.setObjectName("section_label")
        return l

    def _copy_tab(self, tab_index: int):
        txt_widgets = [self.txt_raw, self.txt_summary, self.txt_form, self.txt_risk]
        content = txt_widgets[tab_index].toPlainText()
        if content.strip():
            QApplication.clipboard().setText(content)
        else:
            QMessageBox.information(self, "Nothing to Copy", "This tab is currently empty.")

    def _set_ownership(self, mode: str):
        self._ownership = mode
        if mode == "own":
            self.lbl_ownership.setText("🏢  Classified: Own Website")
            self.lbl_ownership.setStyleSheet(
                "font-size: 11px; color: #89dceb; font-weight: bold;"
            )
            self.btn_mark_own.setStyleSheet(
                "background-color: #89dceb; color: #1e1e2e; "
                "border: 2px solid #04a5e5;"
            )
            self.btn_mark_partner.setStyleSheet("")
        else:
            self.lbl_ownership.setText("🤝  Classified: Partner Website")
            self.lbl_ownership.setStyleSheet(
                "font-size: 11px; color: #fab387; font-weight: bold;"
            )
            self.btn_mark_partner.setStyleSheet(
                "background-color: #fab387; color: #1e1e2e; "
                "border: 2px solid #fe640b;"
            )
            self.btn_mark_own.setStyleSheet("")

    # ── Scan ──────────────────────────────────────────────────────────────────
    def _on_scan(self):
        url = self.inp_scan_url.text().strip()
        if not url:
            QMessageBox.warning(self, "URL Required", "Enter a website URL to scan.")
            return
        if not url.startswith("http"):
            url = "https://" + url
            self.inp_scan_url.setText(url)

        self.btn_scan.setEnabled(False)
        self.btn_gen_summary.setEnabled(False)
        self.btn_send.setEnabled(False)
        self.lbl_scan_status.setText("⏳ Scanning…")
        self.txt_raw.clear()
        self.txt_risk.clear()
        self.txt_summary.clear()
        self.txt_form.clear()
        self.lbl_sum_badge.setText("")

        self._scan_worker = FullSiteParseWorker(url, self)
        self._scan_worker.finished.connect(self._on_scan_done)
        self._scan_worker.error.connect(self._on_scan_error)
        self._scan_worker.start()

    def _on_scan_done(self, data: dict):
        self._parsed_data = data
        self.btn_scan.setEnabled(True)
        self.btn_send.setEnabled(True)
        if _OPENAI_AVAILABLE:
            self.btn_gen_summary.setEnabled(True)

        domain = data.get("domain", "")
        brand = data.get("brand_name_detected", "")
        score = data.get("analyzer_score", 0)
        score_icon = "🟢" if score >= 70 else ("🟡" if score >= 40 else "🔴")
        noindex = data.get("noindex_detected", False)
        robots = data.get("robots_txt_status", "not checked")
        self.lbl_scan_status.setText(
            f"✅ Scanned: {domain}   {score_icon} Trust score: {score}/100\n"
            f"Brand detected: {brand or '[not detected]'}\n"
            f"Privacy Policy: {'✅' if data.get('privacy_policy_detected') else '❌'}  "
            f"Contact page: {'✅' if data.get('contact_page_detected') else '❌'}  "
            f"Terms: {'✅' if data.get('terms_page_detected') else '❌'}  "
            f"Noindex: {'🔴 YES' if noindex else '✅ No'}\n"
            f"robots.txt: {robots}  │  Subpages crawled: {data.get('subpages_crawled', 0)}"
        )

        # ── Populate Raw Data tab ─────────────────────────────────────────────
        field_order = [
            "final_url", "domain", "site_title", "meta_description",
            "h1", "h2_headings", "navigation_items", "call_to_action_buttons",
            "contact_page_detected", "privacy_policy_detected", "terms_page_detected",
            "contact_email", "phone", "address",
            "brand_name_detected", "company_name_detected",
            "product_or_service_detected", "niche_detected",
            "language", "country_or_geo_hint",
            "redirect_chain", "external_links",
            "forms_detected", "login_or_registration_detected",
            "visible_text_summary",
        ]
        raw_lines = []
        for field in field_order:
            val = data.get(field, "")
            if isinstance(val, bool):
                icon = "✅" if val else "❌"
                raw_lines.append(f"{icon}  {field}: {val}")
            elif isinstance(val, list):
                raw_lines.append(f"{'─' * 52}")
                raw_lines.append(f"  {field}:")
                if val:
                    for item in val:
                        raw_lines.append(f"    • {item}")
                else:
                    raw_lines.append("    (none)")
            else:
                raw_lines.append(f"{'─' * 52}")
                raw_lines.append(f"  {field}:")
                raw_lines.append(f"  {val or '(not detected)'}")

        # ── LandingAnalyzer deep analysis section ─────────────────────────────
        _sev_icon = {"CRITICAL": "🔴", "HIGH": "🟡", "MEDIUM": "⚠️", "LOW": "ℹ️", "OK": "✅"}
        score = data.get("analyzer_score", 0)
        raw_lines.append("")
        raw_lines.append("═" * 52)
        raw_lines.append(f"  🔍 LANDING ANALYZER — Trust Score: {score}/100")
        raw_lines.append(f"  Subpages crawled: {data.get('subpages_crawled', 0)}")
        raw_lines.append(f"  robots.txt: {data.get('robots_txt_status', 'not checked')}")
        raw_lines.append(f"  noindex: {'🔴 YES' if data.get('noindex_detected') else '✅ No'}")
        raw_lines.append("─" * 52)
        raw_lines.append("  FINDINGS:")
        for f in data.get("analyzer_findings", []):
            sev = f.get("severity", "")
            icon = _sev_icon.get(sev, "")
            raw_lines.append(f"  {icon} [{sev}] {f.get('title', '')}")
            if f.get("recommendation_ua"):
                raw_lines.append(f"      → {f['recommendation_ua']}")
        bi = data.get("analyzer_business_info", {})
        if bi:
            raw_lines.append("─" * 52)
            raw_lines.append("  BUSINESS INFO (auto-extracted):")
            for k, v in bi.items():
                if v:
                    raw_lines.append(f"    {k}: {v}")

        # ── Appeal Risk Engine v2 block ────────────────────────────────────
        ar = data.get("analyzer_appeal_risk", {})
        if ar:
            raw_lines.append("─" * 52)
            raw_lines.append("  ⚡ APPEAL RISK ENGINE v2")

            niche = ar.get("niche", {})
            if niche:
                raw_lines.append(f"    🏷  Niche: {niche.get('service_label', '—')} "
                                 f"(conf: {niche.get('confidence', 0)}%, "
                                 f"method: {niche.get('method', '—')})")

            strat = ar.get("appeal_strategy", {})
            if strat:
                can = "✅ Yes" if strat.get("can_submit_now") else "❌ Fix issues first"
                raw_lines.append(f"    📋 Strategy: {strat.get('policy', '—')} / {strat.get('recommended_angle', '—')}")
                raw_lines.append(f"    🚦 Risk level: {strat.get('risk_level', '—')}   Can submit now: {can}")
                raw_lines.append(f"    💬 {strat.get('suggested_statement', '')}")

            fixes = ar.get("pre_submit_fixes", [])
            if fixes:
                raw_lines.append("    🔧 PRE-SUBMIT FIXES:")
                _fix_icons = {"critical": "🔴", "high": "🟡", "medium": "⚠️", "low": "ℹ️"}
                for fx in fixes:
                    fi = _fix_icons.get(fx.get("severity", ""), "•")
                    raw_lines.append(f"      {fi} {fx.get('item', '')}")

            claims = ar.get("unsupported_claims", {})
            if claims.get("items"):
                raw_lines.append(f"    ⚠️  Unsupported claims (risk: {claims.get('risk', '—')}):")
                for cl in claims["items"][:6]:
                    raw_lines.append(f"      [{cl['severity'].upper()}] '{cl['text']}' → {cl['recommendation']}")

            brands = ar.get("third_party_brand_risk", {})
            if brands.get("severity") not in ("ok", None, ""):
                raw_lines.append(f"    🏴 Brand risk: {brands.get('severity', '—')}")
                if brands.get("high_risk_brands"):
                    raw_lines.append(f"      High-risk: {', '.join(brands['high_risk_brands'])}")
                if brands.get("generic_fake_brands"):
                    raw_lines.append(f"      Generic/fake: {', '.join(brands['generic_fake_brands'])}")

            facts = ar.get("appeal_safe_facts", [])
            if facts:
                raw_lines.append("    ✅ APPEAL-SAFE FACTS:")
                for fact in facts:
                    raw_lines.append(f"      • {fact}")

            do_not = ar.get("appeal_do_not_mention", [])
            if do_not:
                raw_lines.append("    🚫 DO NOT MENTION IN APPEAL:")
                for item in do_not:
                    raw_lines.append(f"      • {item}")

        raw_lines.append("═" * 52)

        self.txt_raw.setPlainText("\n".join(raw_lines))

        # ── Populate Risk tab ─────────────────────────────────────────────────
        self.txt_risk.setPlainText(_build_risk_report(data))

        self._out_tabs.setCurrentIndex(0)

    def _on_scan_error(self, err: str):
        self.btn_scan.setEnabled(True)
        self.lbl_scan_status.setText(f"❌ Error: {err}")
        QMessageBox.critical(self, "Scan Error", err)

    # ── AI Summary ────────────────────────────────────────────────────────────
    def _on_gen_summary(self):
        if not self._parsed_data:
            QMessageBox.warning(self, "No Data", "Scan a website first.")
            return
        api_key = self.inp_sum_api_key.text().strip()
        if not api_key:
            QMessageBox.warning(
                self, "API Key Required",
                "Enter your OpenAI API key to generate the AI summary."
            )
            return

        self.btn_gen_summary.setEnabled(False)
        self.txt_summary.clear()
        self.txt_form.clear()
        self.lbl_sum_badge.setText("⏳ Generating…")
        self._out_tabs.setCurrentIndex(1)

        self._ai_worker = SummaryAIWorker(api_key, self._parsed_data, self)
        self._ai_worker.delta.connect(self._on_summary_delta)
        self._ai_worker.finished.connect(self._on_summary_done)
        self._ai_worker.error.connect(self._on_summary_error)
        self._ai_worker.start()

    def _on_summary_delta(self, chunk: str):
        # Stream directly into the summary tab during generation
        self.txt_summary.moveCursor(QTextCursor.End)
        self.txt_summary.insertPlainText(chunk)
        self.txt_summary.moveCursor(QTextCursor.End)

    def _on_summary_done(self, full: str):
        self.btn_gen_summary.setEnabled(True)
        self.lbl_sum_badge.setText("✅ Done")
        # Split at "## 11." — sections 1-10 → Summary tab, section 11 → Form tab
        split_marker = "## 11."
        if split_marker in full:
            summary_part = full[: full.index(split_marker)].strip()
            form_part = full[full.index(split_marker) :].strip()
        else:
            summary_part = full.strip()
            form_part = ""
        self.txt_summary.setPlainText(summary_part)
        self.txt_form.setPlainText(form_part)
        if form_part:
            self._out_tabs.setCurrentIndex(2)

    def _on_summary_error(self, err: str):
        self.btn_gen_summary.setEnabled(True)
        self.lbl_sum_badge.setText("❌ Error")
        QMessageBox.critical(self, "AI Summary Error", err)

    # ── Send to Agreement Generator ───────────────────────────────────────────
    def _on_send_to_agreement(self):
        if not self._parsed_data:
            QMessageBox.warning(self, "No Data", "Scan a website first.")
            return
        pd = self._parsed_data
        brand = (
            pd.get("brand_name_detected") or
            pd.get("company_name_detected") or
            pd.get("site_title", "").split("|")[0].strip() or
            ""
        )
        self.send_to_agreement.emit({
            "brand": brand,
            "domain": pd.get("domain", ""),
            "ownership": self._ownership,
        })


# ── Validation ────────────────────────────────────────────────────────────────
def validate_document(text: str, number: str, date: str,
                      contractor: str, client: str, price: str) -> list[str]:
    issues: list[str] = []
    if not text.strip():
        return ["Document is empty."]

    # 1. Agreement number — expect ≥ 3 occurrences (title, clause 1.1, Annex title)
    num_count = len(re.findall(re.escape(f"No. {number}"), text))
    if num_count < 3:
        issues.append(
            f"⚠️ Agreement number «{number}» found {num_count}× (expected ≥ 3: "
            f"title, clause 1.1, Annex title)."
        )

    # 2. Date — expect ≥ 6 (header, clause 1.1, ×2 sig agreement, Annex title, Annex sig ×2)
    date_count = text.count(date)
    if date_count < 6:
        issues.append(
            f"⚠️ Date «{date}» found {date_count}× (expected ≥ 6 across all blocks)."
        )

    # 3. Contractor in signature blocks
    contractor_sig_lines = [
        m.start() for m in re.finditer(r"Contractor:\s*" + re.escape(contractor), text)
    ]
    if len(contractor_sig_lines) < 2:
        issues.append(
            f"⚠️ Contractor «{contractor}» found in only {len(contractor_sig_lines)} "
            f"signature block(s) (expected 2)."
        )

    # 4. Client in signature blocks
    client_sig_lines = [
        m.start() for m in re.finditer(r"Client:\s*" + re.escape(client), text)
    ]
    if len(client_sig_lines) < 2:
        issues.append(
            f"⚠️ Client «{client}» found in only {len(client_sig_lines)} "
            f"signature block(s) (expected 2)."
        )

    # 5. Price consistency — must appear in clause 3.1 AND Annex 2.2
    price_plain = price.replace(",", "")
    price_hits = max(text.count(price), text.count(price_plain))
    if price_hits < 2:
        issues.append(
            f"⚠️ Price «{price}» found {price_hits}× (expected ≥ 2: clause 3.1 + Annex 2.2)."
        )

    # 6. Empty field guards
    for field, val in [("Agreement Number", number), ("Date", date),
                       ("Contractor", contractor), ("Client", client), ("Price", price)]:
        if not val.strip():
            issues.append(f"❌ Field «{field}» is empty.")

    # 7. Swap detection — contractor name must not appear in Client column of sig
    swap_re = re.compile(r"Contractor:\s*(.{1,80}?)\s{3,}Client:\s*(.{1,80})")
    for m in swap_re.finditer(text):
        left, right = m.group(1).strip(), m.group(2).strip()
        if contractor and client:
            if client in left and contractor not in left:
                issues.append("❌ Contractor and Client appear SWAPPED in a signature line!")
                break

    return issues


# ── ODA Widget ────────────────────────────────────────────────────────────────

class OdaWidget(QWidget):
    """Reserve Docs → ODA / Advertised Name Authorization panel."""

    _SERVICE_OPTIONS = [
        "Google Ads campaign setup",
        "Campaign management",
        "Ad content preparation",
        "Targeting setup",
        "Reporting",
        "Campaign optimization",
    ]
    _RESP_OPTIONS = [
        "Website content",
        "Business information shown on website",
        "Product / service fulfillment",
        "Customer support",
        "Brand / displayed name",
    ]
    _RESP_DEFAULTS = {"Website content", "Business information shown on website",
                      "Product / service fulfillment", "Brand / displayed name"}

    def __init__(self, parent=None):
        super().__init__(parent)
        root = QHBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        # ── Left: form ────────────────────────────────────────────────────────
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFixedWidth(400)
        form_w = QWidget()
        ll = QVBoxLayout(form_w)
        ll.setSpacing(8)
        ll.setContentsMargins(10, 10, 10, 10)
        scroll.setWidget(form_w)
        root.addWidget(scroll)

        def lbl(t):
            l = QLabel(t)
            l.setObjectName("section_label")
            return l

        def field(layout, label, placeholder=""):
            layout.addWidget(lbl(label))
            w = QLineEdit()
            w.setPlaceholderText(placeholder)
            layout.addWidget(w)
            return w

        # ── Parties ───────────────────────────────────────────────────────────
        grp_p = QGroupBox("👤  Parties")
        pl = QVBoxLayout(grp_p)
        pl.setSpacing(5)
        self.inp_ads_holder = field(pl, "Ads account holder name", "e.g. PANORMITIS TSOULFAS")
        pl.addWidget(lbl("Ads account holder type"))
        self.cmb_holder_type = QComboBox()
        self.cmb_holder_type.addItems(["Individual", "Organization"])
        pl.addWidget(self.cmb_holder_type)
        self.inp_brand = field(pl, "Advertised brand / displayed business name",
                               "e.g. Equity Money Matters")
        self.inp_website = field(pl, "Associated website", "https://example.com/")
        ll.addWidget(grp_p)

        # ── Agreement ref ─────────────────────────────────────────────────────
        grp_a = QGroupBox("📋  Agreement Reference")
        al = QVBoxLayout(grp_a)
        al.setSpacing(5)
        pl.addWidget(lbl("Relationship type"))
        self.cmb_rel_type = QComboBox()
        self.cmb_rel_type.addItems([
            "Advertising services contractor",
            "Marketing agency with direct agreement",
            "Authorized campaign manager",
            "Other",
        ])
        pl.addWidget(self.cmb_rel_type)
        self.inp_agr_number = field(al, "Agreement number", "e.g. 23510")
        al.addWidget(lbl("Agreement date"))
        self.inp_agr_date = QDateEdit(_random_past_date())
        self.inp_agr_date.setCalendarPopup(True)
        self.inp_agr_date.setDisplayFormat("MMMM d, yyyy")
        al.addWidget(self.inp_agr_date)
        al.addWidget(lbl("Statement date"))
        self.inp_stmt_date = QDateEdit(_random_past_date())
        self.inp_stmt_date.setCalendarPopup(True)
        self.inp_stmt_date.setDisplayFormat("MMMM d, yyyy")
        al.addWidget(self.inp_stmt_date)
        ll.addWidget(grp_a)

        # ── Service scope ─────────────────────────────────────────────────────
        grp_s = QGroupBox("🛠️  Service Scope (select all that apply)")
        sl = QVBoxLayout(grp_s)
        sl.setSpacing(3)
        self._svc_checks: list[QCheckBox] = []
        for opt in self._SERVICE_OPTIONS:
            cb = QCheckBox(opt)
            cb.setChecked(True)
            sl.addWidget(cb)
            self._svc_checks.append(cb)
        ll.addWidget(grp_s)

        # ── Brand responsibilities ─────────────────────────────────────────────
        grp_r = QGroupBox("🏢  Advertised Business Responsibilities")
        rl = QVBoxLayout(grp_r)
        rl.setSpacing(3)
        self._resp_checks: list[QCheckBox] = []
        for opt in self._RESP_OPTIONS:
            cb = QCheckBox(opt)
            cb.setChecked(opt in self._RESP_DEFAULTS)
            rl.addWidget(cb)
            self._resp_checks.append(cb)
        ll.addWidget(grp_r)

        # ── Options ───────────────────────────────────────────────────────────
        grp_o = QGroupBox("⚙️  Options")
        ol = QVBoxLayout(grp_o)
        ol.setSpacing(5)
        self.chk_regulated = QCheckBox("Add regulated-services disclaimer")
        self.chk_regulated.setChecked(False)
        ol.addWidget(self.chk_regulated)
        self.chk_gpt = QCheckBox("Use GPT API to polish wording")
        self.chk_gpt.setChecked(False)
        ol.addWidget(self.chk_gpt)
        ll.addWidget(grp_o)

        # ── Auto-fill ──────────────────────────────────────────────────────
        grp_af = QGroupBox("🎲  Auto-fill")
        afl = QVBoxLayout(grp_af)
        afl.setSpacing(4)
        self.chk_auto_num = QCheckBox("Random 5-digit agreement number (unique)")
        self.chk_auto_num.setChecked(True)
        afl.addWidget(self.chk_auto_num)
        self.chk_auto_date = QCheckBox("Random past date (1 / 1.5 / 2 months before today)")
        self.chk_auto_date.setChecked(True)
        afl.addWidget(self.chk_auto_date)
        ll.addWidget(grp_af)
        # Lock date fields when auto-date is on
        def _oda_toggle_date(checked):
            self.inp_agr_date.setEnabled(not checked)
            self.inp_stmt_date.setEnabled(not checked)
        self.chk_auto_date.toggled.connect(_oda_toggle_date)
        _oda_toggle_date(True)

        # ── Signature ───────────────────────────────────────────────────────
        grp_sig = QGroupBox("✍️  Signature")
        sig_l = QVBoxLayout(grp_sig)
        sig_l.setSpacing(6)
        if not _PIL_AVAILABLE:
            sig_warn = QLabel("⚠️ Pillow not installed — generation disabled.\npip install Pillow")
            sig_warn.setStyleSheet("color: #f38ba8; font-size: 11px;")
            sig_warn.setWordWrap(True)
            sig_l.addWidget(sig_warn)
        self.sig_oda = SignatureWidget("Ads Account Holder")
        sig_l.addWidget(self.sig_oda)
        ll.addWidget(grp_sig)

        # ── Generate button ───────────────────────────────────────────────────
        self.btn_generate = QPushButton("⚡  Generate ODA Document")
        self.btn_generate.setObjectName("btn_generate")
        self.btn_generate.clicked.connect(self._on_generate)
        ll.addWidget(self.btn_generate)

        # ── Export buttons ────────────────────────────────────────────────────
        grp_exp = QGroupBox("💾  Export")
        el = QVBoxLayout(grp_exp)
        el.setSpacing(5)
        self.btn_txt  = QPushButton("💾  Save as .TXT")
        self.btn_txt.setObjectName("btn_export_txt")
        self.btn_txt.clicked.connect(self._export_txt)
        el.addWidget(self.btn_txt)
        self.btn_docx = QPushButton("📄  Save as .DOCX")
        self.btn_docx.setObjectName("btn_export_docx")
        self.btn_docx.clicked.connect(self._export_docx)
        el.addWidget(self.btn_docx)
        self.btn_pdf  = QPushButton("🖨️  Save as .PDF")
        self.btn_pdf.setObjectName("btn_export_pdf")
        self.btn_pdf.clicked.connect(self._export_pdf)
        el.addWidget(self.btn_pdf)
        ll.addWidget(grp_exp)
        ll.addStretch()

        # ── Right: preview ────────────────────────────────────────────────────
        rw = QWidget()
        rl2 = QVBoxLayout(rw)
        rl2.setContentsMargins(6, 0, 0, 0)
        badge_row = QHBoxLayout()
        self.lbl_badge = QLabel("")
        self.lbl_badge.setStyleSheet("font-weight: bold; font-size: 12px;")
        badge_row.addWidget(self.lbl_badge)
        badge_row.addStretch()
        rl2.addLayout(badge_row)
        self.preview = QTextEdit()
        self.preview.setReadOnly(False)
        self.preview.setFont(QFont("Courier New", 10))
        self.preview.setPlaceholderText("Generated ODA document will appear here…")
        rl2.addWidget(self.preview)
        root.addWidget(rw, stretch=1)

        self._api_key_getter: Optional["callable"] = None  # injected by parent
        # Sync holder name → signature widget
        self.inp_ads_holder.textChanged.connect(self.sig_oda.set_name)

    # ── Collect data ──────────────────────────────────────────────────────────
    def _collect_data(self) -> Optional[dict]:
        data = dict(
            ads_holder_name = self.inp_ads_holder.text().strip(),
            holder_type     = self.cmb_holder_type.currentText(),
            brand_name      = self.inp_brand.text().strip(),
            website         = self.inp_website.text().strip(),
            relationship    = self.cmb_rel_type.currentText(),
            agreement_number= self.inp_agr_number.text().strip(),
            agreement_date  = self.inp_agr_date.date().toString("MMMM d, yyyy"),
            statement_date  = self.inp_stmt_date.date().toString("MMMM d, yyyy"),
            service_scope   = [cb.text() for cb in self._svc_checks if cb.isChecked()],
            responsibilities= [cb.text() for cb in self._resp_checks if cb.isChecked()],
            regulated_disclaimer = self.chk_regulated.isChecked(),
        )
        issues = validate_oda_inputs(data)
        # Errors (not warnings) block generation
        hard = [e for e in issues if not e.startswith("⚠️")]
        if hard:
            QMessageBox.warning(self, "Validation Error", "\n".join(hard))
            return None
        soft = [e for e in issues if e.startswith("⚠️")]
        if soft:
            QMessageBox.information(self, "Notice", "\n".join(soft))
        return data

    def _default_filename(self, ext: str) -> str:
        holder = safe_filename(self.inp_ads_holder.text().strip() or "Holder")
        brand  = safe_filename(self.inp_brand.text().strip()        or "Brand")
        return f"Advertised_Name_Authorization_{holder}_{brand}.{ext}"

    # ── Generate ──────────────────────────────────────────────────────────────
    def _on_generate(self):
        # Randomise before collecting
        if self.chk_auto_num.isChecked():
            self.inp_agr_number.setText(_random_agreement_number())
        if self.chk_auto_date.isChecked():
            d = _random_past_date()
            self.inp_agr_date.setDate(d)
            self.inp_stmt_date.setDate(d)
        data = self._collect_data()
        if not data:
            return
        text = build_oda_template(data)
        if self.chk_gpt.isChecked():
            api_key = ""
            if self._api_key_getter:
                api_key = self._api_key_getter()
            if not api_key or not _OPENAI_AVAILABLE:
                QMessageBox.warning(self, "GPT unavailable",
                    "No API key found or openai package not installed.\n"
                    "Document generated using static template instead.")
                self.preview.setPlainText(text)
                self.lbl_badge.setText("📄 Static template")
                self.lbl_badge.setStyleSheet("color: #cdd6f4; font-weight: bold;")
                return
            self._run_gpt(data, text, api_key)
        else:
            self.preview.setPlainText(text)
            self.lbl_badge.setText("📄 Static template")
            self.lbl_badge.setStyleSheet("color: #a6e3a1; font-weight: bold;")

    def _build_gpt_prompt(self, static_doc: str) -> str:
        return (
            "You are editing an Operating Display Name Authorization Statement "
            "for Google Ads verification.\n\n"
            "Your task is to LIGHTLY POLISH the wording of the document below while "
            "preserving the full document structure exactly.\n\n"
            "STRICT RULES:\n"
            "1. Do not remove any section.\n"
            "2. Do not remove any heading.\n"
            "3. Do not remove the document title.\n"
            "4. Do not remove the signature block.\n"
            "5. Do not remove these labels:\n"
            "   - Ads Account Holder / Advertising Contractor:\n"
            "   - Advertised Business / Client:\n"
            "   - Associated Website:\n"
            "   - Agreement Reference:\n"
            "   - Signature:\n"
            "   - Date:\n"
            "6. Do not change names, dates, URLs, agreement numbers, or brand names.\n"
            "7. Do not invent facts.\n"
            "8. Do not claim business registration, DBA registration, trademark "
            "ownership, licenses, certification, or legal ownership unless explicitly "
            "written in the source document.\n"
            "9. Do not say the Ads account holder owns the advertised brand unless the "
            "source document says so.\n"
            "10. Do not summarize. Do not shorten aggressively.\n"
            "11. Do not return markdown.\n"
            "12. Return only the final polished document text.\n\n"
            "Keep the meaning:\n"
            "- Ads account holder is an advertising contractor / campaign manager.\n"
            "- Advertised brand is the client / advertised business.\n"
            "- Relationship is based on an Advertising Services Agreement.\n"
            "- Ads account holder is responsible for advertising campaign management only.\n"
            "- Advertised brand remains responsible for website content, business "
            "information, and product/service fulfillment.\n"
            "- The statement explains the name mismatch between Ads account holder and "
            "advertised brand.\n\n"
            "SOURCE DOCUMENT TO POLISH:\n"
            "\"\"\"\n"
            f"{static_doc}\n"
            "\"\"\"\n\n"
            "Return the full polished document."
        )

    def _run_gpt(self, data: dict, fallback_text: str, api_key: str):
        self.btn_generate.setEnabled(False)
        self.preview.clear()
        self.lbl_badge.setText("🤖 GPT polishing…")
        self.lbl_badge.setStyleSheet("color: #89b4fa; font-weight: bold;")
        # Store context for _on_gpt_done
        self._oda_gpt_fallback = fallback_text
        self._oda_gpt_data = data
        prompt = self._build_gpt_prompt(fallback_text)
        self._oda_worker = OpenAIWorker(api_key, prompt, self)
        # Collect silently; do NOT stream into preview until validation passes
        self._oda_worker.finished.connect(self._on_gpt_done)
        self._oda_worker.error.connect(self._on_gpt_error)
        self._oda_worker.start()

    def _on_gpt_error(self, err: str):
        QMessageBox.warning(self, "GPT Error",
            f"GPT error: {err}\nFalling back to static template.")
        self.preview.setPlainText(self._oda_gpt_fallback)
        self.btn_generate.setEnabled(True)
        self.lbl_badge.setText("📄 Static template")
        self.lbl_badge.setStyleSheet("color: #cdd6f4; font-weight: bold;")

    def _on_gpt_done(self, full: str):
        self.btn_generate.setEnabled(True)
        fallback = getattr(self, "_oda_gpt_fallback", "")
        data     = getattr(self, "_oda_gpt_data", {})
        issues   = validate_gpt_oda_output(full, data, len(fallback))
        if issues:
            QMessageBox.warning(
                self, "⚠️ GPT output failed validation",
                "GPT polish output failed validation — static template used instead.\n\n"
                + "\n".join(f"• {e}" for e in issues)
            )
            self.preview.setPlainText(fallback)
            self.lbl_badge.setText("📄 Static template fallback")
            self.lbl_badge.setStyleSheet("color: #fab387; font-weight: bold;")
        else:
            self.preview.setPlainText(full)
            self.lbl_badge.setText("🤖 GPT polished")
            self.lbl_badge.setStyleSheet("color: #a6e3a1; font-weight: bold;")

    # ── Export helpers ────────────────────────────────────────────────────────
    def _current_text(self) -> Optional[str]:
        t = self.preview.toPlainText().strip()
        if not t:
            QMessageBox.warning(self, "Nothing to export", "Generate a document first.")
            return None
        return t

    def _export_txt(self):
        text = self._current_text()
        if not text:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Save TXT", self._default_filename("txt"), "Text files (*.txt)")
        if not path:
            return
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)

    def _export_docx(self):
        if not _DOCX_AVAILABLE:
            QMessageBox.warning(self, "DOCX unavailable",
                "python-docx is not installed.\npip install python-docx")
            return
        text = self._current_text()
        if not text:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Save DOCX", self._default_filename("docx"),
            "Word documents (*.docx)")
        if not path:
            return
        doc = DocxDocument()
        # Save signature to temp file if available
        import tempfile as _tmp
        _sig_tmp = None
        if self.sig_oda.has_signature():
            _sig_tmp = _tmp.NamedTemporaryFile(suffix=".png", delete=False)
            self.sig_oda.get_pixmap().save(_sig_tmp.name, "PNG")
            _sig_tmp.close()
        for line in text.splitlines():
            if line.strip() == "[SIG_ODA]":
                if _sig_tmp:
                    try:
                        doc.add_picture(_sig_tmp.name, width=Inches(2.0))
                    except Exception:
                        doc.add_paragraph("[signature]")
                else:
                    doc.add_paragraph("_" * 30)
                continue
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)
            if line.strip().isupper() and len(line.strip()) > 4:
                if p.runs:
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(12)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
        if _sig_tmp:
            try:
                os.unlink(_sig_tmp.name)
            except Exception:
                pass
        doc.save(path)

    def _export_pdf(self):
        if not _FPDF_AVAILABLE:
            QMessageBox.warning(self, "PDF unavailable",
                "fpdf2 is not installed.\npip install fpdf2")
            return
        text = self._current_text()
        if not text:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF", self._default_filename("pdf"),
            "PDF files (*.pdf)")
        if not path:
            return
        import re as _re
        pdf = _FPDF()
        pdf.set_margins(25, 20, 25)
        _WIN_FONTS = r"C:\Windows\Fonts"
        _arial = os.path.join(_WIN_FONTS, "arial.ttf")
        _arialb = os.path.join(_WIN_FONTS, "arialbd.ttf")
        _use_arial = os.path.exists(_arial)
        if _use_arial:
            pdf.add_font("Arial", "",  _arial)
            pdf.add_font("Arial", "B", _arialb if os.path.exists(_arialb) else _arial)
            _FN = "Arial"
        else:
            _FN = "Helvetica"
        def _safe(s):
            return s if _use_arial else s.encode("latin-1", errors="replace").decode("latin-1")
        FONT_N = (_FN, "",  10)
        FONT_B = (_FN, "B", 10)
        FONT_T = (_FN, "B", 12)
        # Prepare signature temp file
        import tempfile as _tmp2
        _sig_tmp2 = None
        if self.sig_oda.has_signature():
            _sig_tmp2 = _tmp2.NamedTemporaryFile(suffix=".png", delete=False)
            self.sig_oda.get_pixmap().save(_sig_tmp2.name, "PNG")
            _sig_tmp2.close()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        W = pdf.epw

        def _mcell(txt: str, font=None, h: float = 5, align: str = "L"):
            """Write a multi_cell line, always resetting X to l_margin first."""
            pdf.set_x(pdf.l_margin)
            pdf.set_font(*(font or FONT_N))
            pdf.multi_cell(W, h, _safe(txt), align=align)
            pdf.set_x(pdf.l_margin)

        for i, line in enumerate(text.splitlines()):
            stripped = line.strip()

            # Signature image marker
            if stripped == "[SIG_ODA]":
                pdf.set_x(pdf.l_margin)
                if _sig_tmp2:
                    try:
                        pdf.image(_sig_tmp2.name, x=pdf.l_margin, w=50)
                        pdf.set_x(pdf.l_margin)
                    except Exception:
                        _mcell("_" * 30)
                else:
                    _mcell("_" * 30)
                continue

            # Blank line
            if stripped == "":
                pdf.ln(3)
                pdf.set_x(pdf.l_margin)
                continue

            # Document title (first line, all-caps)
            if i == 0 and stripped.isupper():
                _mcell(stripped, font=FONT_T, h=6, align="C")
                continue

            # Section headings — all-caps words, short (like "SIGNATURES")
            if stripped.isupper() and len(stripped) > 4:
                pdf.ln(3)
                _mcell(stripped, font=FONT_B)
                continue

            # Label lines ending with ":" → bold, acts as field header
            if stripped.endswith(":") and len(stripped) < 80:
                pdf.ln(1)
                _mcell(stripped, font=FONT_B)
                continue

            # Default: normal wrapped text
            _mcell(stripped)
        pdf.output(path)
        if _sig_tmp2:
            try:
                os.unlink(_sig_tmp2.name)
            except Exception:
                pass


# ── DBA Statement Widget ─────────────────────────────────────────────────────

_TRACKING_RE = re.compile(
    r"[?&](utm_[a-z_]+|gclid|gclsrc|dclid|gbraid|wbraid|fbclid|msclkid|twclid|ttclid"
    r"|mc_eid|mc_cid|_ga|_gl|ref|referrer|source)=[^&]*",
    re.IGNORECASE,
)
_GENERIC_RE = re.compile(
    r"^(home|homepage|welcome|index|untitled|new tab|main page|default|page not found|404|error)$",
    re.IGNORECASE,
)
_TITLE_SPLIT_RE = re.compile(r"\s*[\|\-\u2013\u2014\xb7\u2022]\s*")


def _clean_tracking_url(url: str) -> str:
    """Strip known tracking parameters from a URL."""
    if not url:
        return url
    from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
    try:
        p = urlparse(url.strip())
        _TRACKING_KEYS = {
            "utm_source", "utm_medium", "utm_campaign", "utm_term", "utm_content",
            "gclid", "gclsrc", "dclid", "gbraid", "wbraid",
            "fbclid", "msclkid", "twclid", "ttclid",
            "mc_eid", "mc_cid", "ref", "referrer", "source", "_ga", "_gl",
        }
        qs = parse_qs(p.query, keep_blank_values=False)
        clean_qs = {k: v for k, v in qs.items()
                    if k.lower() not in _TRACKING_KEYS
                    and not re.match(r"^(gl_|_)", k)}
        query = urlencode(clean_qs, doseq=True)
        result = urlunparse((p.scheme or "https", p.netloc,
                             p.path or "/", p.params, query, ""))
        if not p.path or p.path == "/":
            result = result.rstrip("/") + "/"
        return result
    except Exception:
        return url


def _clean_site_title(title: str) -> str:
    """
    Extract the brand part from a site title.
    "Fortune Financial | Smart money starts here" → "Fortune Financial"
    "Expert Motor Works - Trust the professionals" → "Expert Motor Works"
    """
    if not title:
        return ""
    parts = _TITLE_SPLIT_RE.split(title.strip())
    # Take first segment; skip if it looks generic or is a URL
    cand = parts[0].strip() if parts else title.strip()
    if not cand or _GENERIC_RE.search(cand) or cand.startswith("http") or len(cand) < 2:
        return ""
    return cand


def _clean_domain_to_name(domain: str) -> str:
    """opendirect.site → Opendirect   |   expert-motor-works.com → Expert Motor Works"""
    d = re.sub(r"^www\.", "", domain.strip())
    d = re.sub(r"\.[a-z]{2,10}$", "", d)          # strip TLD
    d = re.sub(r"[-_.]", " ", d).strip().title()
    return d


def _extract_dba_from_profile(profile: dict, legal_name: str = "") -> dict:
    """
    Extract DBA fields from analyze_site() / SiteParseWorker output.
    Actual keys: visible_brand, og_site_name, title, final_url,
                 headings (list[{level,text}]), summary_text, keywords.
    Legal name is NEVER taken from the parser.
    """
    from urllib.parse import urlparse as _up
    warnings: list[str] = []

    # ── Derive domain from final_url ──────────────────────────────────────
    raw_url = (profile.get("final_url") or profile.get("url") or "").strip()
    try:
        _parsed = _up(raw_url)
        _domain_raw = _parsed.netloc or ""
    except Exception:
        _domain_raw = ""

    # ── Operating name — strict priority ─────────────────────────────────
    # 1. visible_brand  (header logo alt / og:site_name / title-left / short H1)
    # 2. og_site_name   (explicit og:site_name)
    # 3. clean(title)   — first segment before | or -
    # 4. clean(domain)  — last fallback
    operating_name = ""
    op_source = ""

    vb = (profile.get("visible_brand") or "").strip()
    if vb and not _GENERIC_RE.search(vb) and not vb.startswith("http"):
        operating_name, op_source = vb, "visible_brand (header/logo/og:site_name)"

    if not operating_name:
        og = (profile.get("og_site_name") or "").strip()
        if og and not _GENERIC_RE.search(og):
            operating_name, op_source = og, "og_site_name"

    if not operating_name:
        title_clean = _clean_site_title(profile.get("title") or "")
        if title_clean:
            operating_name, op_source = title_clean, "title (cleaned)"

    if not operating_name:
        domain_name = _clean_domain_to_name(_domain_raw)
        if domain_name:
            operating_name, op_source = domain_name, "domain (fallback)"
            warnings.append("Operating name derived from domain — low confidence. Verify manually.")

    # ── Associated website ────────────────────────────────────────────────
    had_tracking = bool(raw_url and _TRACKING_RE.search(raw_url))
    if raw_url:
        associated_website = _clean_tracking_url(raw_url)
        if had_tracking:
            warnings.append(f"Tracking parameters stripped from URL → {associated_website}")
    else:
        associated_website = (f"https://{_domain_raw}/" if _domain_raw else "")

    # ── Business activity summary ─────────────────────────────────────────
    headings = profile.get("headings") or []  # list of {level, text}
    h1_text  = next((h["text"] for h in headings if h.get("level") == "h1"), "")
    h2_texts = [h["text"] for h in headings
                if h.get("level") == "h2" and len(h.get("text", "")) > 3
                and not _GENERIC_RE.search(h.get("text", ""))]
    summary  = (profile.get("summary_text") or "").strip()
    keywords = profile.get("keywords") or []  # list of str

    parts: list[str] = []
    # Top keywords as activity context
    if keywords:
        kw_str = ", ".join(str(k) for k in keywords[:8])
        if kw_str:
            parts.append(kw_str)
    # H1 if short and meaningful
    if h1_text and len(h1_text) < 80 and not _GENERIC_RE.search(h1_text):
        if h1_text.lower() not in (p.lower() for p in parts):
            parts.append(h1_text)
    # up to 2 H2s
    for h in h2_texts[:2]:
        if len(h) < 70 and h.lower() not in (p.lower() for p in parts):
            parts.append(h)
    # fallback: first sentence of summary text
    if not parts and summary:
        first = re.split(r"[.!?]", summary)[0].strip()
        if len(first) > 10:
            parts.append(first)

    activity_summary = ", ".join(parts)[:400]

    # ── Validation warnings ───────────────────────────────────────────────
    if not legal_name.strip():
        warnings.append("Legal Name is empty. Enter it manually (from Ads account).")
    if not operating_name:
        warnings.append("Could not detect operating name. Enter manually.")
    elif _GENERIC_RE.search(operating_name):
        warnings.append(f"Detected name '{operating_name}' looks generic. Verify or override.")
    if legal_name and operating_name and legal_name.strip().lower() == operating_name.strip().lower():
        warnings.append("Legal name equals operating name — use Same Entity Statement instead.")
    if not associated_website:
        warnings.append("Could not detect website URL from parser.")
    if not activity_summary:
        warnings.append("Business activity summary is empty — add a description manually.")

    return {
        "operating_name":            operating_name,
        "operating_name_source":     op_source,
        "associated_website":        associated_website,
        "business_activity_summary": activity_summary,
        "warnings":                  warnings,
    }


_DBA_TEMPLATE = """\
STATEMENT OF USE OF OPERATING BUSINESS NAME

Date: {date}

{legal_name} confirms that it uses the operating / displayed business name "{operating_name}" for advertising and promotional activity related to the website:

{associated_website}

The name "{operating_name}" is used as the public-facing business name on the website, in advertising materials, and in Google Ads campaigns.
{activity_line}
This statement is provided to clarify the relationship between the legal / Google Ads account holder name and the displayed business name shown in ads and on the associated website.

This statement does not claim that "{operating_name}" is a separate registered legal entity or officially registered DBA unless such registration is separately confirmed by official documentation.


Legal Name / Google Ads Account Holder:
{legal_name}

Operating / Displayed Business Name:
{operating_name}

Associated Website:
{associated_website}
{cid_line}{ref_line}

Signature:
[SIG_DBA]

Date:
{date}
"""


class _DbaGenWorker(QThread):
    finished = Signal(str)
    error    = Signal(str)

    def __init__(self, data: dict, api_key: str, parent=None):
        super().__init__(parent)
        self._data    = data
        self._api_key = api_key

    def run(self):
        try:
            if not _OPENAI_AVAILABLE or not self._api_key:
                raise RuntimeError("no_gpt")
            from openai import OpenAI as _OAI
            client = _OAI(api_key=self._api_key)
            prompt = (
                "You are a legal document specialist for Google Ads verification.\n"
                "Polish the following Statement of Use of Operating Business Name.\n"
                "Rules:\n"
                "- Keep it factual, calm, professional.\n"
                "- Do NOT claim official DBA/government registration.\n"
                "- Use wording: 'operating business name', 'displayed business name', "
                "'public-facing business name'.\n"
                "- Do NOT invent data.\n"
                "- Keep all placeholder markers like [SIG_DBA] exactly.\n"
                "- Return ONLY the document text, no commentary.\n\n"
                "DOCUMENT:\n\"\"\"\n"
                f"{self._data['static_text']}\n\"\"\""
            )
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
            )
            content = resp.choices[0].message.content or ""
            self.finished.emit(content.strip())
        except RuntimeError as e:
            if str(e) == "no_gpt":
                self.finished.emit(self._data["static_text"])
            else:
                self.error.emit(str(e))
        except Exception as e:
            self.error.emit(str(e))


class DbaStatementWidget(QWidget):
    """DBA / Operating Business Name Statement generator tab."""

    def __init__(self, parent=None):
        super().__init__(parent)
        from typing import Callable
        self._api_key_getter: Callable[[], str] = lambda: ""
        self._gen_worker: Optional[QThread] = None

        root = QHBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        splitter = QSplitter()
        splitter.setOrientation(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)
        root.addWidget(splitter)

        # ── Left: form (scrollable) ───────────────────────────────────────
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setMinimumWidth(320)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_w = QWidget()
        ll = QVBoxLayout(form_w)
        ll.setSpacing(8)
        ll.setContentsMargins(10, 10, 10, 10)
        scroll.setWidget(form_w)
        splitter.addWidget(scroll)

        def lbl(t: str) -> QLabel:
            l = QLabel(t)
            l.setObjectName("section_label")
            return l

        def field(label: str, placeholder: str = "") -> QLineEdit:
            ll.addWidget(lbl(label))
            w = QLineEdit()
            w.setPlaceholderText(placeholder)
            ll.addWidget(w)
            return w

        def textarea(label: str, placeholder: str = "", h: int = 70) -> QTextEdit:
            ll.addWidget(lbl(label))
            w = QTextEdit()
            w.setPlaceholderText(placeholder)
            w.setMaximumHeight(h)
            ll.addWidget(w)
            return w

        # ── Site URL scan ─────────────────────────────────────────────────
        grp_parser = QGroupBox("① Auto-fill from Website")
        gl = QVBoxLayout(grp_parser)

        url_row = QHBoxLayout()
        self.inp_dba_url = QLineEdit()
        self.inp_dba_url.setPlaceholderText("Enter domain or URL, e.g. expertmotorworks.com")
        self.btn_dba_scan = QPushButton("🔍  Parse Site")
        self.btn_dba_scan.setFixedWidth(120)
        self.btn_dba_scan.clicked.connect(self._on_dba_scan)
        url_row.addWidget(self.inp_dba_url)
        url_row.addWidget(self.btn_dba_scan)
        gl.addLayout(url_row)

        self.lbl_extract_status = QLabel("")
        self.lbl_extract_status.setWordWrap(True)
        self.lbl_extract_status.setStyleSheet("color:#fab387; font-size:11px;")
        gl.addWidget(self.lbl_extract_status)

        # collapsible JSON fallback
        details_lbl = QLabel("<a href='#'>▸ Or paste raw parser JSON manually</a>")
        details_lbl.setTextInteractionFlags(Qt.TextBrowserInteraction)
        details_lbl.setOpenExternalLinks(False)
        self.inp_parser_json = QTextEdit()
        self.inp_parser_json.setPlaceholderText(
            "Paste profile_builder v2 JSON here for auto-fill.\n"
            "Fields used: identity.brand_final, input.final_url, "
            "site_profile, classification.niche"
        )
        self.inp_parser_json.setMaximumHeight(110)
        self.inp_parser_json.setVisible(False)
        btn_extract = QPushButton("🔍  Extract from JSON")
        btn_extract.setVisible(False)
        btn_extract.clicked.connect(self._on_extract)

        def _toggle_json_area(_checked=False):
            vis = not self.inp_parser_json.isVisible()
            self.inp_parser_json.setVisible(vis)
            btn_extract.setVisible(vis)
            details_lbl.setText(
                "<a href='#'>▾ Or paste raw parser JSON manually</a>" if vis
                else "<a href='#'>▸ Or paste raw parser JSON manually</a>"
            )

        details_lbl.linkActivated.connect(_toggle_json_area)
        gl.addWidget(details_lbl)
        gl.addWidget(self.inp_parser_json)
        gl.addWidget(btn_extract)
        ll.addWidget(grp_parser)

        # ── Core fields ───────────────────────────────────────────────────
        grp_core = QGroupBox("② Statement Fields")
        cl = QVBoxLayout(grp_core)

        def cfield(label: str, placeholder: str = "") -> QLineEdit:
            cl.addWidget(lbl(label))
            w = QLineEdit()
            w.setPlaceholderText(placeholder)
            cl.addWidget(w)
            return w

        def ctextarea(label: str, placeholder: str = "", h: int = 65) -> QTextEdit:
            cl.addWidget(lbl(label))
            w = QTextEdit()
            w.setPlaceholderText(placeholder)
            w.setMaximumHeight(h)
            cl.addWidget(w)
            return w

        self.inp_legal_name     = cfield("Legal Name / Ads Account Holder *",
                                         "e.g. RDP  (NEVER from parser — Ads profile only)")
        self.inp_operating_name = cfield("Operating / Displayed Business Name *",
                                         "e.g. Expert Motor Works  (auto-filled from parser)")
        self.inp_website        = cfield("Associated Website *",
                                         "e.g. https://pandanbeta.info/")
        self.inp_activity       = ctextarea("Business Activity Summary",
                                            "e.g. automotive services, engine tuning, diagnostics"
                                            " (auto-filled from parser)")
        self.inp_cid            = cfield("Google Ads CID", "xxx-xxx-xxxx (optional)")
        self.inp_reference      = cfield("Reference / Case #", "optional")

        cl.addWidget(lbl("Statement Date"))
        self.inp_date = QDateEdit(QDate.currentDate())
        self.inp_date.setCalendarPopup(True)
        self.inp_date.setDisplayFormat("MMMM d, yyyy")
        cl.addWidget(self.inp_date)

        self.chk_disclaimer = QCheckBox(
            "Add disclaimer: operating name is not a separate registered legal entity"
        )
        self.chk_disclaimer.setChecked(True)
        cl.addWidget(self.chk_disclaimer)
        ll.addWidget(grp_core)

        # ── Same-name warning ─────────────────────────────────────────────
        self.lbl_same_warn = QLabel("")
        self.lbl_same_warn.setWordWrap(True)
        self.lbl_same_warn.setStyleSheet("color:#f38ba8; font-weight:bold;")
        ll.addWidget(self.lbl_same_warn)
        self.inp_legal_name.textChanged.connect(self._check_names)
        self.inp_operating_name.textChanged.connect(self._check_names)

        # ── Signature ─────────────────────────────────────────────────────
        grp_sig = QGroupBox("③ Signature")
        sl = QVBoxLayout(grp_sig)
        self.sig_dba = SignatureWidget(party_label="DBA Statement")
        sl.addWidget(self.sig_dba)
        ll.addWidget(grp_sig)

        # keep signature name in sync with the legal name field
        self.inp_legal_name.textChanged.connect(
            lambda t: self.sig_dba.set_name(t.strip())
        )

        # ── Validation ────────────────────────────────────────────────────
        grp_val = QGroupBox("④ Validation")
        vl = QVBoxLayout(grp_val)
        self.txt_validation = QTextEdit()
        self.txt_validation.setReadOnly(True)
        self.txt_validation.setMaximumHeight(80)
        vl.addWidget(self.txt_validation)
        btn_validate = QPushButton("🔍  Validate")
        btn_validate.setObjectName("btn_validate")
        btn_validate.clicked.connect(self._on_validate)
        vl.addWidget(btn_validate)
        ll.addWidget(grp_val)

        # ── Generate ──────────────────────────────────────────────────────
        self.btn_generate = QPushButton("📄  Generate DBA Statement")
        self.btn_generate.setObjectName("btn_generate")
        self.btn_generate.clicked.connect(self._on_generate)
        ll.addWidget(self.btn_generate)

        # ── Export buttons ────────────────────────────────────────────────
        exp_row = QHBoxLayout()
        btn_txt  = QPushButton("💾 TXT");  btn_txt.setObjectName("btn_export_txt")
        btn_docx = QPushButton("📝 DOCX"); btn_docx.setObjectName("btn_export_docx")
        btn_pdf  = QPushButton("📄 PDF");  btn_pdf.setObjectName("btn_export_pdf")
        btn_txt.clicked.connect(self._export_txt)
        btn_docx.clicked.connect(self._export_docx)
        btn_pdf.clicked.connect(self._export_pdf)
        exp_row.addWidget(btn_txt)
        exp_row.addWidget(btn_docx)
        exp_row.addWidget(btn_pdf)
        ll.addLayout(exp_row)
        ll.addStretch()

        # ── Right: preview ────────────────────────────────────────────────
        rw = QWidget()
        rw.setMinimumWidth(260)
        rl = QVBoxLayout(rw)
        rl.setContentsMargins(6, 6, 6, 6)
        self.lbl_badge = QLabel("📋 DBA Statement preview")
        self.lbl_badge.setObjectName("section_label")
        rl.addWidget(self.lbl_badge)
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        self.preview.setFont(QFont("Consolas", 11))
        rl.addWidget(self.preview)
        splitter.addWidget(rw)
        splitter.setStretchFactor(0, 2)
        splitter.setStretchFactor(1, 3)

    # ── Logic ─────────────────────────────────────────────────────────────

    def _check_names(self):
        legal = self.inp_legal_name.text().strip()
        op    = self.inp_operating_name.text().strip()
        if legal and op and legal.lower() == op.lower():
            self.lbl_same_warn.setText(
                "⚠  Legal name equals operating name — use Same Entity Statement instead."
            )
        else:
            self.lbl_same_warn.setText("")

    # ── URL-based site scan ───────────────────────────────────────────────
    def _on_dba_scan(self):
        url = self.inp_dba_url.text().strip()
        if not url:
            self.lbl_extract_status.setText("⚠  Enter a domain or URL first.")
            return
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        if not _PROFILER_AVAILABLE:
            self.lbl_extract_status.setText("❌  Site profiler not available in this environment.")
            return
        self.btn_dba_scan.setEnabled(False)
        self.lbl_extract_status.setText("🔄  Scanning site…")
        self._dba_scan_worker = SiteParseWorker(url, self)
        self._dba_scan_worker.finished.connect(self._on_dba_scan_done)
        self._dba_scan_worker.error.connect(self._on_dba_scan_error)
        self._dba_scan_worker.start()

    def _on_dba_scan_done(self, profile: dict):
        self.btn_dba_scan.setEnabled(True)
        legal  = self.inp_legal_name.text().strip()
        result = _extract_dba_from_profile(profile, legal)
        if result["operating_name"]:
            self.inp_operating_name.setText(result["operating_name"])
        if result["associated_website"]:
            self.inp_website.setText(result["associated_website"])
        if result["business_activity_summary"]:
            self.inp_activity.setPlainText(result["business_activity_summary"])
        lines  = [f"✔ Operating name: {result['operating_name']!r}  [{result['operating_name_source']}]"]
        lines += [f"✔ Website: {result['associated_website']!r}"]
        if result["warnings"]:
            lines += [""] + [f"⚠  {w}" for w in result["warnings"]]
        self.lbl_extract_status.setStyleSheet("color:#a6e3a1; font-size:11px;")
        self.lbl_extract_status.setText("\n".join(lines))

    def _on_dba_scan_error(self, msg: str):
        self.btn_dba_scan.setEnabled(True)
        self.lbl_extract_status.setStyleSheet("color:#f38ba8; font-size:11px;")
        self.lbl_extract_status.setText(f"❌  Scan failed: {msg}")

    def _on_extract(self):
        import json as _json
        raw = self.inp_parser_json.toPlainText().strip()
        if not raw:
            self.lbl_extract_status.setText("⚠  Parser JSON is empty.")
            return
        try:
            profile = _json.loads(raw)
        except Exception as e:
            self.lbl_extract_status.setText(f"❌  Invalid JSON: {e}")
            return
        legal  = self.inp_legal_name.text().strip()
        result = _extract_dba_from_profile(profile, legal)
        if result["operating_name"]:
            self.inp_operating_name.setText(result["operating_name"])
        if result["associated_website"]:
            self.inp_website.setText(result["associated_website"])
        if result["business_activity_summary"]:
            self.inp_activity.setPlainText(result["business_activity_summary"])
        lines  = [f"✔ Operating name: {result['operating_name']!r}  [{result['operating_name_source']}]"]
        lines += [f"✔ Website: {result['associated_website']!r}"]
        if result["warnings"]:
            lines += [""] + [f"⚠  {w}" for w in result["warnings"]]
        self.lbl_extract_status.setText("\n".join(lines))

    def _collect(self) -> dict:
        return {
            "legal_name":     self.inp_legal_name.text().strip(),
            "operating_name": self.inp_operating_name.text().strip(),
            "website":        self.inp_website.text().strip(),
            "activity":       self.inp_activity.toPlainText().strip(),
            "cid":            self.inp_cid.text().strip(),
            "reference":      self.inp_reference.text().strip(),
            "date":           self.inp_date.date().toString("MMMM d, yyyy"),
            "disclaimer":     self.chk_disclaimer.isChecked(),
        }

    def _validate_data(self, d: dict) -> list[str]:
        warns: list[str] = []
        if not d["legal_name"]:
            warns.append("Legal Name is required.")
        if not d["operating_name"]:
            warns.append("Operating / Displayed Business Name is required.")
        if not d["website"]:
            warns.append("Associated Website is required.")
        if d["legal_name"] and d["operating_name"] and \
                d["legal_name"].lower() == d["operating_name"].lower():
            warns.append("Legal name equals operating name — use Same Entity Statement.")
        if d["website"] and _TRACKING_RE.search(d["website"]):
            warns.append("Website URL contains tracking parameters. Clean before export.")
        if d["operating_name"] and _GENERIC_RE.search(d["operating_name"]):
            warns.append(f"Operating name '{d['operating_name']}' looks generic.")
        if not d["activity"]:
            warns.append("Business activity summary is empty — consider adding a short description.")
        return warns

    def _on_validate(self):
        warns = self._validate_data(self._collect())
        if warns:
            self.txt_validation.setPlainText("\n".join(f"⚠  {w}" for w in warns))
        else:
            self.txt_validation.setPlainText("✅  All checks passed.")

    def _build_static_text(self, d: dict) -> str:
        activity_line = (
            f"\nThe website presents information about {d['activity']}.\n"
            if d["activity"] else "\n"
        )
        cid_line = f"\nGoogle Ads CID: {d['cid']}" if d["cid"] else ""
        ref_line = f"\nReference: {d['reference']}"  if d["reference"] else ""
        return _DBA_TEMPLATE.format(
            legal_name     = d["legal_name"],
            operating_name = d["operating_name"],
            associated_website = d["website"],
            activity_line  = activity_line,
            cid_line       = cid_line,
            ref_line       = ref_line,
            date           = d["date"],
        )

    def _on_generate(self):
        d = self._collect()
        warns = self._validate_data(d)
        if warns:
            ret = QMessageBox.warning(
                self, "Validation Warnings",
                "\\n".join(f"• {w}" for w in warns) + "\\n\\nContinue generation?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if ret == QMessageBox.StandardButton.No:
                return
        static = self._build_static_text(d)
        api_key = self._api_key_getter()
        self.btn_generate.setEnabled(False)
        self.btn_generate.setText("⏳  Generating…")
        self._gen_worker = _DbaGenWorker({"static_text": static}, api_key, self)
        self._gen_worker.finished.connect(self._on_gen_done)
        self._gen_worker.error.connect(self._on_gen_error)
        self._gen_worker.start()

    def _on_gen_done(self, text: str):
        self.btn_generate.setEnabled(True)
        self.btn_generate.setText("📄  Generate DBA Statement")
        # Strip markdown code fences GPT sometimes adds
        text = re.sub(r"^```[^\n]*\n?", "", text.strip())
        text = re.sub(r"\n?```$", "", text.strip())
        self.preview.setPlainText(text.strip())
        self.lbl_badge.setText("📋 DBA Statement — ready")

    def _on_gen_error(self, err: str):
        self.btn_generate.setEnabled(True)
        self.btn_generate.setText("📄  Generate DBA Statement")
        QMessageBox.critical(self, "Error", err)

    def _current_text(self) -> Optional[str]:
        t = self.preview.toPlainText().strip()
        if not t:
            QMessageBox.warning(self, "Nothing to export", "Generate a statement first.")
            return None
        return t

    def _default_filename(self, ext: str) -> str:
        legal = re.sub(r"[^\w]", "_", self.inp_legal_name.text().strip())[:20] or "dba"
        op    = re.sub(r"[^\w]", "_", self.inp_operating_name.text().strip())[:20] or "statement"
        return f"DBA_Statement_{legal}_{op}.{ext}"

    def _export_txt(self):
        text = self._current_text()
        if not text:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Save TXT", self._default_filename("txt"), "Text files (*.txt)")
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)

    def _export_docx(self):
        if not _DOCX_AVAILABLE:
            QMessageBox.warning(self, "DOCX unavailable", "pip install python-docx")
            return
        text = self._current_text()
        if not text:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Save DOCX", self._default_filename("docx"), "Word documents (*.docx)")
        if not path:
            return
        import tempfile as _tmp
        if not _DOCX_AVAILABLE:
            QMessageBox.warning(self, "DOCX unavailable", "pip install python-docx")
            return
        doc  = DocxDocument()  # type: ignore[possibly-undefined]
        _sig_tmp = None
        if self.sig_dba.has_signature():
            px = self.sig_dba.get_pixmap()
            if px is not None:
                _sig_tmp = _tmp.NamedTemporaryFile(suffix=".png", delete=False)
                px.save(_sig_tmp.name, "PNG")
                _sig_tmp.close()
        for line in text.splitlines():
            if line.strip() == "[SIG_DBA]":
                if _sig_tmp:
                    try:
                        from docx.shared import Inches as _Inches
                        doc.add_picture(_sig_tmp.name, width=_Inches(2.0))
                    except Exception:
                        doc.add_paragraph("[signature]")
                else:
                    doc.add_paragraph("_" * 30)
                continue
            p = doc.add_paragraph(line)
            if p.runs:
                from docx.shared import Pt as _Pt
                p.runs[0].font.size = _Pt(11)
            if line.strip().isupper() and len(line.strip()) > 4:
                if p.runs:
                    from docx.shared import Pt as _Pt
                    p.runs[0].bold = True; p.runs[0].font.size = _Pt(12)
            from docx.shared import Pt as _Pt
            p.paragraph_format.space_before = _Pt(2)
            p.paragraph_format.space_after  = _Pt(2)
        if _sig_tmp:
            try:
                os.unlink(_sig_tmp.name)
            except Exception:
                pass
        doc.save(path)

    def _export_pdf(self):
        if not _FPDF_AVAILABLE:
            QMessageBox.warning(self, "PDF unavailable", "pip install fpdf2")
            return
        text = self._current_text()
        if not text:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF", self._default_filename("pdf"), "PDF files (*.pdf)")
        if not path:
            return
        import tempfile as _tmp2
        if not _FPDF_AVAILABLE:
            QMessageBox.warning(self, "PDF unavailable", "pip install fpdf2")
            return
        pdf = _FPDF()  # type: ignore[possibly-undefined]
        pdf.set_margins(25, 20, 25)
        _WIN_FONTS = r"C:\Windows\Fonts"
        _arial  = os.path.join(_WIN_FONTS, "arial.ttf")
        _arialb = os.path.join(_WIN_FONTS, "arialbd.ttf")
        _use_arial = os.path.exists(_arial)
        if _use_arial:
            pdf.add_font("Arial", "",  _arial)
            pdf.add_font("Arial", "B", _arialb if os.path.exists(_arialb) else _arial)
            _FN = "Arial"
        else:
            _FN = "Helvetica"
        def _safe(s):
            return s if _use_arial else s.encode("latin-1", errors="replace").decode("latin-1")
        FONT_N = (_FN, "",  10); FONT_B = (_FN, "B", 10); FONT_T = (_FN, "B", 12)
        _sig_tmp2 = None
        if self.sig_dba.has_signature():
            px2 = self.sig_dba.get_pixmap()
            if px2 is not None:
                _sig_tmp2 = _tmp2.NamedTemporaryFile(suffix=".png", delete=False)
                px2.save(_sig_tmp2.name, "PNG")
                _sig_tmp2.close()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        W = pdf.epw
        def _mcell(txt, font=None, h=5, align="L"):
            pdf.set_x(pdf.l_margin)
            f = font or FONT_N
            pdf.set_font(str(f[0]), str(f[1]), float(f[2]))
            pdf.multi_cell(W, h, _safe(txt), align=align)
            pdf.set_x(pdf.l_margin)
        for i, line in enumerate(text.splitlines()):
            stripped = line.strip()
            if stripped == "[SIG_DBA]":
                if _sig_tmp2:
                    try:
                        pdf.image(_sig_tmp2.name, x=pdf.l_margin, w=50)
                        pdf.set_x(pdf.l_margin)
                    except Exception:
                        _mcell("_" * 30)
                else:
                    _mcell("_" * 30)
                continue
            if stripped == "":
                pdf.ln(3); pdf.set_x(pdf.l_margin); continue
            if i == 0 and stripped.isupper():
                _mcell(stripped, font=FONT_T, h=6, align="C"); continue
            if stripped.isupper() and len(stripped) > 4:
                pdf.ln(3); _mcell(stripped, font=FONT_B); continue
            if stripped.endswith(":") and len(stripped) < 80:
                pdf.ln(1); _mcell(stripped, font=FONT_B); continue
            _mcell(stripped)
        pdf.output(path)
        if _sig_tmp2:
            try:
                os.unlink(_sig_tmp2.name)
            except Exception:
                pass


# ── Octo Script Generator Widget ─────────────────────────────────────────────

_APPSCRIPT_CREATE_URL = "https://script.google.com/home/projects/create?"
_GS_TEMPLATE_NAME = "AdsMailImporter.gs"


def _load_gs_template() -> str:
    """Load AdsMailImporter.gs from the same directory as this file."""
    import os
    here = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(here, _GS_TEMPLATE_NAME)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""


def _strip_profile_name(window_title: str) -> str:
    """
    From '[Transferred] GERM1040 - Octium' extract 'GERM1040'.
    Strips any leading [...] block and trailing ' - <appname>' suffix.
    """
    import re as _re
    name = _re.sub(r'^\[.*?\]\s*', '', window_title).strip()   # strip [Transferred] etc.
    name = _re.sub(r'\s*-\s*[^-]+$', '', name).strip()         # strip ' - Octium Browser' etc.
    return name


def _scan_octo_windows() -> list:
    """
    Enumerate all top-level Windows windows.
    Only returns real Octo profile windows — those ending with ' - Octium'.
    Title format: '[Transferred] GERM1040 - Octium'  →  returns 'GERM1040'
    Works via ctypes — zero extra dependencies.
    """
    import ctypes
    import ctypes.wintypes

    titles: list = []
    WNDENUMPROC = ctypes.WINFUNCTYPE(ctypes.c_bool,
                                     ctypes.wintypes.HWND,
                                     ctypes.wintypes.LPARAM)

    def _cb(hwnd, _lp):
        length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
        if length > 0:
            buf = ctypes.create_unicode_buffer(length + 1)
            ctypes.windll.user32.GetWindowTextW(hwnd, buf, length + 1)
            t = buf.value
            # Match any window whose title contains 'octium' (profile windows)
            if 'octium' in t.lower():
                clean = _strip_profile_name(t)
                if clean:
                    titles.append(clean)
        return True

    try:
        ctypes.windll.user32.EnumWindows(WNDENUMPROC(_cb), 0)
    except Exception:
        pass

    return titles


class FastSiteParseWorker(QThread):
    """Quick site parser using Googlebot UA + BeautifulSoup, totheweb.com-style output."""
    finished = Signal(str, str, str)   # result_text, profile_name, final_url
    error = Signal(str)

    def __init__(self, url: str, profile_name: str = "", parent=None):
        super().__init__(parent)
        self._url = url
        self._profile_name = profile_name

    def run(self):
        try:
            import sys, os
            _root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            if _root not in sys.path:
                sys.path.insert(0, _root)

            from se_simulator.core.scanner import SEOScanner, format_result_text, normalize_url

            url = normalize_url(self._url)

            scanner = SEOScanner(progress_cb=lambda _: None)
            _, _, preferred = scanner.scan_pair(url)
            result_text = format_result_text(preferred)
            final_url = preferred.url
            self.finished.emit(result_text, self._profile_name, final_url)
        except Exception as exc:
            self.error.emit(str(exc))


# Persistent scan log file — saved next to doc_generator.py
_SCAN_LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "octo_scan_log.json")


class LogoFetchWorker(QThread):
    """Downloads the best logo/favicon for a given URL and saves to a temp PNG."""
    finished = Signal(str)   # path to saved PNG
    error    = Signal(str)

    def __init__(self, url: str, parent=None):
        super().__init__(parent)
        self._url = url

    def run(self):
        try:
            import requests as _req
            from urllib.parse import urlparse as _up, urljoin as _uj
            from bs4 import BeautifulSoup as _BS

            base_url = self._url
            if not base_url.startswith("http"):
                base_url = "https://" + base_url
            parsed = _up(base_url)
            origin = f"{parsed.scheme}://{parsed.netloc}"

            _HDR = {"User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            )}

            # Fetch the page to look for og:image and apple-touch-icon
            try:
                resp = _req.get(origin, headers=_HDR, timeout=12, allow_redirects=True)
                soup = _BS(resp.text, "html.parser")
            except Exception:
                soup = None

            candidates = []  # list of (kind, url, is_svg)

            if soup:
                # 1. <img> tags with "logo" in class/id/alt/src — best chance for real logo
                for tag in soup.find_all("img"):
                    combined = " ".join([
                        str(tag.get("class", "")),
                        str(tag.get("id", "")),
                        str(tag.get("alt", "")),
                        str(tag.get("src", "")),
                    ]).lower()
                    if "logo" in combined:
                        src = tag.get("src") or tag.get("data-src") or tag.get("data-lazy-src")
                        if src:
                            url_c = _uj(origin, str(src))
                            is_svg = url_c.lower().split("?")[0].endswith(".svg") or "image/svg" in str(tag.get("type",""))
                            candidates.append(("img-logo", url_c, is_svg))

                # 2. og:image — usually high-quality brand image
                og = soup.find("meta", property="og:image")
                if og and og.get("content"):
                    u = _uj(origin, str(og["content"]))
                    candidates.append(("og:image", u, u.lower().split("?")[0].endswith(".svg")))

                # 3. apple-touch-icon
                for rel in ("apple-touch-icon", "apple-touch-icon-precomposed"):
                    tag = soup.find("link", rel=rel)
                    if tag and tag.get("href"):
                        u = _uj(origin, str(tag["href"]))
                        candidates.append(("touch-icon", u, u.lower().split("?")[0].endswith(".svg")))

                # 4. Large favicon (icon link with size attribute)
                for tag in soup.find_all("link", rel=lambda r: r and "icon" in " ".join(r).lower()):
                    sizes = str(tag.get("sizes", ""))
                    href  = tag.get("href", "")
                    if href:
                        sz = 0
                        try:
                            sz = int(sizes.split("x")[0]) if "x" in sizes else 0
                        except Exception:
                            pass
                        u = _uj(origin, str(href))
                        candidates.append((f"icon-{sz}", u, u.lower().split("?")[0].endswith(".svg")))

            # 5. Standard favicon fallbacks
            candidates += [
                ("favicon.ico", f"{origin}/favicon.ico", False),
                ("favicon.png", f"{origin}/favicon.png", False),
            ]

            # Try each candidate, pick first one that loads
            import io as _io, os as _os
            from PIL import Image as _Img

            def _svg_to_png(svg_bytes: bytes, size: int = 300) -> "_Img.Image | None":
                """Convert SVG bytes → PIL Image using PySide6 QtSvg (no native DLL needed)."""
                try:
                    from PySide6.QtSvg import QSvgRenderer as _QSvgR
                    from PySide6.QtGui import QImage as _QImg, QPainter as _QP
                    from PySide6.QtCore import QByteArray as _QBA, Qt as _Qt
                    renderer = _QSvgR(_QBA(svg_bytes))
                    if not renderer.isValid():
                        return None
                    # Keep aspect ratio
                    s = renderer.defaultSize()
                    if s.width() > 0 and s.height() > 0:
                        ratio = s.width() / s.height()
                        w = size if ratio >= 1 else int(size * ratio)
                        h = int(size / ratio) if ratio >= 1 else size
                    else:
                        w = h = size
                    qimg = _QImg(w, h, _QImg.Format.Format_ARGB32)
                    qimg.fill(_Qt.GlobalColor.transparent)
                    painter = _QP(qimg)
                    renderer.render(painter)
                    painter.end()
                    buf = _io.BytesIO()
                    qimg.save_to_format = None  # just use PIL from raw bytes
                    # Convert QImage → bytes → PIL
                    ptr = qimg.bits()
                    arr = bytes(ptr)
                    pil_img = _Img.frombuffer("RGBA", (w, h), arr, "raw", "BGRA", 0, 1)
                    return pil_img.copy()
                except Exception:
                    pass
                # Fallback: cairosvg
                try:
                    import cairosvg as _cs
                    png_bytes = _cs.svg2png(bytestring=svg_bytes, output_width=size, output_height=size)
                    return _Img.open(_io.BytesIO(png_bytes)).convert("RGBA")
                except Exception:
                    pass
                return None

            tmp_dir = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "__pycache__")
            _os.makedirs(tmp_dir, exist_ok=True)
            out_path = _os.path.join(tmp_dir, "_logo_auto.png")

            for kind, img_url, is_svg in candidates:
                try:
                    r = _req.get(img_url, headers=_HDR, timeout=8, allow_redirects=True)
                    if r.status_code != 200 or not r.content:
                        continue

                    ct = r.headers.get("content-type", "")
                    is_svg = is_svg or "svg" in ct

                    if is_svg:
                        img = _svg_to_png(r.content)
                        if img is None:
                            continue
                    else:
                        img = _Img.open(_io.BytesIO(r.content))

                    # Skip tiny tracking pixels
                    if img.size[0] < 16 or img.size[1] < 16:
                        continue
                    # Convert to RGBA PNG
                    if img.mode not in ("RGB", "RGBA"):
                        img = img.convert("RGBA")
                    img.save(out_path, format="PNG")
                    self.finished.emit(out_path)
                    return
                except Exception:
                    continue

            self.error.emit("No usable logo found")
        except Exception as e:
            self.error.emit(str(e))


class ScanLogViewerDialog(QDialog):
    """Opens on double-click of a scan log row — shows full SEO scan result."""

    def __init__(self, entry: dict, parent=None):
        super().__init__(parent)
        self.entry = entry
        self.setWindowTitle("Scan Result Viewer")
        self.resize(780, 560)
        self.setStyleSheet("""
            QDialog { background: #1e1e1e; color: #e0e0e0; }
            QLabel  { color: #aaaaaa; font-size: 11px; }
            QLabel#val { color: #ffffff; font-size: 11px; font-weight: bold; }
            QTextEdit { background: #141414; color: #d4d4d4; border: 1px solid #3a3a3a;
                        font-family: Consolas, monospace; font-size: 11px; }
            QPushButton { background: #2d2d2d; color: #e0e0e0; border: 1px solid #444;
                          border-radius: 4px; padding: 5px 14px; font-size: 11px; }
            QPushButton:hover { background: #3c3c3c; }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        # Header grid: ts / profile / url / title
        grid = QGridLayout()
        grid.setHorizontalSpacing(10)
        grid.setVerticalSpacing(4)

        def lbl(text):
            w = QLabel(text)
            return w

        def val(text):
            w = QLabel(text)
            w.setObjectName("val")
            w.setTextInteractionFlags(Qt.TextSelectableByMouse)
            w.setWordWrap(True)
            return w

        grid.addWidget(lbl("Час:"),     0, 0)
        grid.addWidget(val(entry.get("ts", "")),      0, 1)
        grid.addWidget(lbl("Профіль:"), 1, 0)
        grid.addWidget(val(entry.get("profile", "")), 1, 1)
        grid.addWidget(lbl("URL:"),     2, 0)
        grid.addWidget(val(entry.get("url", "")),     2, 1)
        grid.addWidget(lbl("Title:"),   3, 0)
        grid.addWidget(val(entry.get("title", "")),   3, 1)
        grid.setColumnStretch(1, 1)
        layout.addLayout(grid)

        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet("color: #333;")
        layout.addWidget(sep)

        # Full result text
        self._text_edit = QTextEdit()
        self._text_edit.setReadOnly(True)
        result = entry.get("result", "")
        self._text_edit.setPlainText(result if result else "(no scan data stored for this entry)")
        layout.addWidget(self._text_edit, 1)

        # Buttons
        btn_row = QHBoxLayout()
        btn_copy_all = QPushButton("📋 Copy All")
        btn_copy_url = QPushButton("🔗 Copy URL")
        btn_close    = QPushButton("✕ Close")
        btn_copy_all.clicked.connect(self._copy_all)
        btn_copy_url.clicked.connect(self._copy_url)
        btn_close.clicked.connect(self.accept)
        btn_row.addWidget(btn_copy_all)
        btn_row.addWidget(btn_copy_url)
        btn_row.addStretch()
        btn_row.addWidget(btn_close)
        layout.addLayout(btn_row)

    def _copy_all(self):
        QApplication.clipboard().setText(self._text_edit.toPlainText())

    def _copy_url(self):
        QApplication.clipboard().setText(self.entry.get("url", ""))


class OctoScriptWidget(QWidget):
    """
    Tab: Auto-scans Windows window titles every 3 s for open Octo profiles.
    No button press needed — just open a profile in Octo Browser.
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self._generated_script: str = ""
        self._last_profiles: list = []
        self._scan_log: list = []   # list of (timestamp, profile, url, title)
        self._build_ui()
        self._load_scan_log()

        # Auto-scan timer — fires every 3 seconds
        from PySide6.QtCore import QTimer
        self._timer = QTimer(self)
        self._timer.setInterval(3000)
        self._timer.timeout.connect(self._scan)
        self._timer.start()
        self._scan()   # immediate first scan

        # Global hotkeys via Win32 RegisterHotKey — OFF by default, toggled by user
        self._hotkey_thread = None
        self._hotkeys_active = False
        # Do NOT auto-start: self._start_global_hotkeys()

    # Signal to safely call slot from hotkey thread
    _hotkey_signal = Signal(int)

    def _toggle_hotkeys(self):
        if self._hotkeys_active:
            # Stop: unregister and kill thread via PostQuitMessage
            try:
                import ctypes
                user32 = ctypes.windll.user32
                user32.UnregisterHotKey(None, 1)
                user32.UnregisterHotKey(None, 2)
                user32.UnregisterHotKey(None, 3)
                # Post WM_QUIT to the hotkey thread's message loop
                if self._hotkey_thread and self._hotkey_thread.is_alive():
                    user32.PostQuitMessage(0)
            except Exception:
                pass
            self._hotkeys_active = False
            self._btn_hotkey_toggle.setText("[R/F/E] Хоткеї  ○  Вимкнуто")
            self._btn_hotkey_toggle.setStyleSheet(
                "QPushButton { background: #45475a; color: #a6adc8; font-weight: bold;"
                " border-radius: 5px; padding: 5px 14px; font-size: 12px; }"
                "QPushButton:hover { background: #585b70; }"
            )
        else:
            self._start_global_hotkeys()
            self._hotkeys_active = True
            self._btn_hotkey_toggle.setText("[R/F/E] Хоткеї  ●  Ввімкнено")
            self._btn_hotkey_toggle.setStyleSheet(
                "QPushButton { background: #a6e3a1; color: #1e1e2e; font-weight: bold;"
                " border-radius: 5px; padding: 5px 14px; font-size: 12px; }"
                "QPushButton:hover { background: #b9f1b5; }"
            )

    def _start_global_hotkeys(self):
        import threading, ctypes, ctypes.wintypes

        self._hotkey_signal.connect(self._on_hotkey)

        def _is_octo_focused():
            user32 = ctypes.windll.user32
            hwnd = user32.GetForegroundWindow()
            length = user32.GetWindowTextLengthW(hwnd)
            if length <= 0:
                return False
            buf = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buf, length + 1)
            return 'octium' in buf.value.lower()

        def _hotkey_loop():
            user32 = ctypes.windll.user32
            MOD_NOREPEAT = 0x4000
            # Register only when we need them — but WM_HOTKEY is system-wide.
            # We gate execution in the handler via _is_octo_focused().
            user32.RegisterHotKey(None, 1, MOD_NOREPEAT, 0x52)  # R
            user32.RegisterHotKey(None, 2, MOD_NOREPEAT, 0x46)  # F
            user32.RegisterHotKey(None, 3, MOD_NOREPEAT, 0x45)  # E
            msg = ctypes.wintypes.MSG()
            while user32.GetMessageW(ctypes.byref(msg), None, 0, 0) != 0:
                if msg.message == 0x0312:  # WM_HOTKEY
                    # Strict check: only act when Octo Browser / Octium is the foreground app
                    if _is_octo_focused():
                        self._hotkey_signal.emit(int(msg.wParam))
                    # else: silently ignore — key press passes through to active app
                user32.TranslateMessage(ctypes.byref(msg))
                user32.DispatchMessageW(ctypes.byref(msg))
            user32.UnregisterHotKey(None, 1)
            user32.UnregisterHotKey(None, 2)
            user32.UnregisterHotKey(None, 3)

        self._hotkey_thread = threading.Thread(target=_hotkey_loop, daemon=True)
        self._hotkey_thread.start()

    def _on_hotkey(self, hid: int):
        if hid == 1:
            self._show_status(self.generate_and_copy(), 5000)
        elif hid == 2:
            self._do_global_f()
        elif hid == 3:
            raw = QApplication.clipboard().text().strip()
            if raw:
                # Normalize and set full URL
                if not re.match(r"^https?://", raw, re.I):
                    raw = "https://" + raw
                self._fast_url.setText(raw)
            self._start_fast_scan()

    def _do_global_r(self):
        self._show_status(self.generate_and_copy(), 5000)

    def _do_global_f(self):
        self.open_appscript()

    # ── UI ────────────────────────────────────────────────────────────────
    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(14, 14, 14, 10)
        root.setSpacing(10)

        title = QLabel("🤖  Appeal Script Generator — AdsMailImporter.gs")
        title.setStyleSheet("font-size: 16px; font-weight: bold; color: #89b4fa;")
        root.addWidget(title)

        sub = QLabel(
            "Відкрий профіль в Octo Browser — він з'явиться у списку автоматично (кожні 3 сек). "
            "[R] — генерує скрипт і копіює в буфер.  [F] — відкрити AppScript.  "
            "[E] — вставляє URL з буфера в Fast Site Parser і одразу сканує."
        )
        sub.setWordWrap(True)
        sub.setStyleSheet("color: #6c7086; font-size: 12px;")
        root.addWidget(sub)

        # Hotkey toggle button
        hk_row = QHBoxLayout()
        self._btn_hotkey_toggle = QPushButton("[R/F/E] Хоткеї  ○  Вимкнуто")
        self._btn_hotkey_toggle.setStyleSheet(
            "QPushButton { background: #45475a; color: #a6adc8; font-weight: bold;"
            " border-radius: 5px; padding: 5px 14px; font-size: 12px; }"
            "QPushButton:hover { background: #585b70; }"
        )
        self._btn_hotkey_toggle.setToolTip(
            "Ввімкнути/вимкнути глобальні хоткеї [R/F/E].\n"
            "Хоткеї спрацьовують тільки коли в фокусі Octo Browser."
        )
        self._btn_hotkey_toggle.clicked.connect(self._toggle_hotkeys)
        hk_row.addWidget(self._btn_hotkey_toggle)
        hk_row.addStretch()
        root.addLayout(hk_row)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        root.addWidget(splitter, 1)

        # ── Left: auto-detected profiles ──────────────────────────────────
        left_w = QWidget()
        ll = QVBoxLayout(left_w)
        ll.setContentsMargins(0, 0, 6, 0)
        ll.setSpacing(6)

        grp_find = QGroupBox("🖥️  Відкриті профілі Octo  (авто, кожні 3 сек)")
        fl = QVBoxLayout(grp_find)
        fl.setSpacing(4)

        self._lbl_status = QLabel("🔍  Сканую вікна…")
        self._lbl_status.setStyleSheet("color: #6c7086; font-size: 11px;")
        self._lbl_status.setWordWrap(True)
        fl.addWidget(self._lbl_status)

        from PySide6.QtWidgets import QListWidget
        self._list = QListWidget()
        self._list.setStyleSheet(
            "QListWidget { font-size: 14px; }"
            "QListWidget::item { padding: 6px 10px; }"
            "QListWidget::item:selected { background-color: #89b4fa; color: #1e1e2e; }"
        )
        self._list.itemClicked.connect(self._on_profile_clicked)
        fl.addWidget(self._list, 1)

        ll.addWidget(grp_find)

        # ── Scan Log table ─────────────────────────────────
        from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        grp_log = QGroupBox("📄  Лог сканів  (профіль → URL → Title)")
        log_lay = QVBoxLayout(grp_log)
        log_lay.setContentsMargins(4, 6, 4, 6)
        log_lay.setSpacing(4)

        self._log_table = QTableWidget(0, 4)
        self._log_table.setHorizontalHeaderLabels(["Час", "Профіль", "URL", "Title"])
        self._log_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self._log_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self._log_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self._log_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self._log_table.verticalHeader().setVisible(False)
        self._log_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self._log_table.setSelectionBehavior(QTableWidget.SelectRows)
        self._log_table.setAlternatingRowColors(True)
        self._log_table.setStyleSheet(
            "QTableWidget { font-size: 11px; alternate-background-color: #1e1e2e; background: #181825; }"
            "QTableWidget::item { padding: 2px 6px; }"
        )
        self._log_table.setMinimumHeight(180)
        self._log_table.doubleClicked.connect(self._open_log_entry)
        log_lay.addWidget(self._log_table)

        btn_clear_log = QPushButton("🗑  Очистити лог")
        btn_clear_log.setStyleSheet(
            "QPushButton { background: #45475a; color: #cdd6f4; padding: 3px 10px;"
            " border-radius: 4px; font-size: 11px; max-width: 120px; }"
            "QPushButton:hover { background: #585b70; }"
        )
        btn_clear_log.clicked.connect(self._clear_scan_log)
        log_lay.addWidget(btn_clear_log)
        ll.addWidget(grp_log)

        splitter.addWidget(left_w)

        # ── Right: generate + preview ─────────────────────────────────────
        right_w = QWidget()
        rl = QVBoxLayout(right_w)
        rl.setContentsMargins(6, 0, 0, 0)
        rl.setSpacing(8)

        # Hidden name field — auto-managed by scan/click
        self._inp_name = QLineEdit()
        self._inp_name.setVisible(False)

        btn_row = QHBoxLayout()

        self._btn_copy_script = QPushButton("📋  Скопіювати скрипт  [R]")
        self._btn_copy_script.setStyleSheet(
            "QPushButton { background-color: #89b4fa; color: #1e1e2e; font-weight: bold;"
            " padding: 7px 16px; border-radius: 6px; }"
            "QPushButton:hover { background-color: #b4befe; }"
        )
        self._btn_copy_script.clicked.connect(lambda: self._show_status(self.generate_and_copy(), 5000))
        btn_row.addWidget(self._btn_copy_script)

        self._btn_open_appscript = QPushButton("🔗  Відкрити AppScript  [F]")
        self._btn_open_appscript.setStyleSheet(
            "QPushButton { background-color: #cba6f7; color: #1e1e2e; font-weight: bold;"
            " padding: 7px 16px; border-radius: 6px; }"
            "QPushButton:hover { background-color: #f5c2e7; }"
        )
        self._btn_open_appscript.clicked.connect(self._on_open_appscript)
        btn_row.addWidget(self._btn_open_appscript)

        btn_row.addStretch()
        rl.addLayout(btn_row)

        preview_lbl = QLabel("📄  Згенерований скрипт (AdsMailImporter.gs):")
        preview_lbl.setStyleSheet("color: #a6e3a1; font-weight: bold;")
        rl.addWidget(preview_lbl)

        self._preview = QTextEdit()
        self._preview.setReadOnly(True)
        self._preview.setFont(QFont("Consolas", 10))
        self._preview.setPlaceholderText(
            "Тут з'явиться готовий скрипт AdsMailImporter.gs.\n\n"
            "1. Відкрий профіль в Octo Browser\n"
            "2. Він з'явиться у списку зліва автоматично\n"
            "3. Клікни на нього → назва заповниться\n"
            "4. Натисни «Згенерувати скрипт» — скрипт скопіюється сам\n"
            "5. [F] → Відкрий AppScript → Вставте → ▶ fullSetup()"
        )
        rl.addWidget(self._preview, 1)

        # ── Fast Site Parser ──────────────────────────────────────────────
        grp_fast = QGroupBox("⚡  Fast Site Parser  [E] — вставляє URL з буфера і сканує")
        grp_fast.setStyleSheet(
            "QGroupBox { font-weight: bold; color: #f9e2af; border: 1px solid #45475a;"
            " border-radius: 6px; margin-top: 8px; padding-top: 8px; }"
            "QGroupBox::title { subcontrol-origin: margin; left: 8px; }"
        )
        fl2 = QVBoxLayout(grp_fast)
        fl2.setContentsMargins(8, 8, 8, 8)
        fl2.setSpacing(5)

        url_row = QHBoxLayout()
        self._fast_url = QLineEdit()
        self._fast_url.setPlaceholderText("URL сайту…  (натисни [E] щоб вставити з буфера)")
        self._fast_url.returnPressed.connect(self._start_fast_scan)
        url_row.addWidget(self._fast_url, 1)
        self._fast_btn = QPushButton("⚡ Scan")
        self._fast_btn.setStyleSheet(
            "QPushButton { background-color: #f9e2af; color: #1e1e2e; font-weight: bold;"
            " padding: 5px 14px; border-radius: 5px; }"
            "QPushButton:hover { background-color: #fab387; }"
            "QPushButton:disabled { background-color: #45475a; color: #6c7086; }"
        )
        self._fast_btn.clicked.connect(self._start_fast_scan)
        url_row.addWidget(self._fast_btn)
        fl2.addLayout(url_row)

        self._fast_status = QLabel("Готово до сканування  •  [E] вставляє URL з буфера і починає скан")
        self._fast_status.setStyleSheet("color: #6c7086; font-size: 11px;")
        fl2.addWidget(self._fast_status)

        self._fast_out = QTextEdit()
        self._fast_out.setReadOnly(True)
        self._fast_out.setFixedHeight(130)
        self._fast_out.setFont(QFont("Consolas", 9))
        self._fast_out.setPlaceholderText("Результат парсингу з'явиться тут і буде скопійований в буфер…")
        fl2.addWidget(self._fast_out)

        rl.addWidget(grp_fast)

        splitter.addWidget(right_w)
        splitter.setSizes([380, 800])

    # ── Auto-scan (QTimer every 3 s) ──────────────────────────────────────
    def _scan(self):
        profiles = _scan_octo_windows()
        if profiles == self._last_profiles:
            return   # nothing changed — don't redraw

        # Detect newly added profiles
        new_profiles = [p for p in profiles if p not in self._last_profiles]
        self._last_profiles = profiles

        self._list.blockSignals(True)
        self._list.clear()
        for name in profiles:
            self._list.addItem(name)
        self._list.blockSignals(False)

        if profiles:
            # Auto-fill: if new profile appeared — use it; else if only one total — use it
            auto_name: str | None = None
            if new_profiles:
                auto_name = new_profiles[-1]   # most recently opened
            elif len(profiles) == 1:
                auto_name = profiles[0]

            if auto_name:
                self._inp_name.setText(auto_name)
                # Highlight in list
                for i in range(self._list.count()):
                    if self._list.item(i).text() == auto_name:
                        self._list.setCurrentRow(i)
                        break
                self._show_status(
                    f"✅  Профіль визначено: {auto_name}  •  клікни або [R]", 5000
                )
            else:
                self._show_status(
                    f"✅  {len(profiles)} профілів відкрито  •  клікни потрібний", 5000
                )
        else:
            self._inp_name.clear()
            self._show_status("⏳  Відкритих профілів Octo не знайдено…  сканую кожні 3 сек", 4000)

    # ── Profile click ─────────────────────────────────────────────────────
    def _on_profile_clicked(self, item):
        name = item.text()
        self._inp_name.setText(name)
        self._show_status(self.generate_and_copy(), 5000)

    # ── Generate ──────────────────────────────────────────────────────────
    def _on_generate(self):
        name = self._inp_name.text().strip()
        if not name:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Потрібна назва",
                                "Клікни на профіль у списку, або введи назву вручну.")
            return

        template = _load_gs_template()
        if not template:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Помилка",
                                 f"Файл {_GS_TEMPLATE_NAME} не знайдено поруч з doc_generator.py")
            return

        import re as _re
        script = _re.sub(
            r"(const OCTO_PROFILE_NAME\s*=\s*)'[^']*'",
            f"\\g<1>'{name}'",
            template,
        )
        self._generated_script = script
        self._preview.setPlainText(script)
        self._btn_copy_script.setEnabled(True)
        QApplication.clipboard().setText(script)
        self._show_status(
            f"✅  Скрипт згенеровано для «{name}» і скопійовано в буфер!  "
            f"[F] → Відкрий AppScript → Вставте → ▶ fullSetup()", 8000)

    # ── Actions ───────────────────────────────────────────────────────────
    def _on_copy_script(self):
        if self._generated_script:
            QApplication.clipboard().setText(self._generated_script)
            self._show_status("✅  Скрипт скопійовано в буфер  [R]", 3000)

    def _on_open_appscript(self):
        # Copy URL to clipboard then paste into whatever field is focused
        import ctypes, time
        QApplication.clipboard().setText(_APPSCRIPT_CREATE_URL)
        # Small delay so clipboard is ready, then send Ctrl+V to the OS
        time.sleep(0.05)
        ctypes.windll.user32.keybd_event(0x11, 0, 0, 0)        # Ctrl down
        ctypes.windll.user32.keybd_event(0x56, 0, 0, 0)        # V down
        ctypes.windll.user32.keybd_event(0x56, 0, 2, 0)        # V up
        ctypes.windll.user32.keybd_event(0x11, 0, 2, 0)        # Ctrl up
        self._show_status("📋  AppScript URL скопійовано і вставлено  [F]", 3000)

    def _show_status(self, msg: str, ms: int = 4000):
        w = self.parent()
        while w and not hasattr(w, "status_bar"):
            w = w.parent() if hasattr(w, "parent") else None
        if w:
            w.status_bar.showMessage(msg, ms)

    # ── Public helpers (called from DocGeneratorWindow hotkeys) ───────────
    def generate_and_copy(self) -> str:
        """Generate script from current profile name and copy. Returns status message."""
        name = self._inp_name.text().strip()
        if not name:
            return "⚠️  Відкрий профіль в Octo Browser  [R]"
        template = _load_gs_template()
        if not template:
            return f"❌  Файл {_GS_TEMPLATE_NAME} не знайдено"
        import re as _re
        script = _re.sub(
            r"(const OCTO_PROFILE_NAME\s*=\s*)'[^']*'",
            f"\\g<1>'{name}'",
            template,
        )
        self._generated_script = script
        self._preview.setPlainText(script)
        self._btn_copy_script.setEnabled(True)
        QApplication.clipboard().setText(script)
        return f"✅  Скрипт згенеровано і скопійовано — «{name}»  [R]"

    def copy_script_to_clipboard(self) -> bool:
        if self._generated_script:
            QApplication.clipboard().setText(self._generated_script)
            return True
        return False

    def open_appscript(self):
        self._on_open_appscript()

    # ── Fast Site Parser ─────────────────────────────────────────────────
    def _start_fast_scan(self):
        url = self._fast_url.text().strip()
        if not url:
            self._fast_status.setText("⚠️  URL пустий — вставте адресу сайту")
            return
        if not re.match(r"^https?://", url, re.I):
            url = "https://" + url
            self._fast_url.setText(url)
        self._fast_status.setText("⏳  Сканую (Googlebot)…")
        self._fast_btn.setEnabled(False)
        self._fast_out.clear()
        profile_name = self._inp_name.text().strip()
        self._fast_worker = FastSiteParseWorker(url, profile_name, self)
        self._fast_worker.finished.connect(self._on_fast_done)
        self._fast_worker.error.connect(self._on_fast_error)
        self._fast_worker.start()

    def _on_fast_done(self, result: str, profile_name: str, final_url: str):
        self._fast_out.setPlainText(result)
        self._fast_btn.setEnabled(True)
        self._fast_status.setText("✅  Готово — результат скопійовано в буфер  [E]")
        QApplication.clipboard().setText(result)
        # ── extract title from result text ──
        title = ""
        for line in result.splitlines():
            if line.startswith("TITLE"):
                title = line.split("\t", 1)[-1].strip()
                break
        # Only log if we got meaningful content (title present)
        if title:
            self._add_scan_log_row(profile_name, final_url, title, result)
        try:
            import winsound
            winsound.Beep(1000, 200)
        except Exception:
            pass

    def _add_scan_log_row(self, profile: str, url: str, title: str, result_text: str = ""):
        from PySide6.QtWidgets import QTableWidgetItem
        from datetime import datetime
        import json
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entry = {"ts": ts, "profile": profile, "url": url, "title": title, "result": result_text}
        self._scan_log.append(entry)
        try:
            with open(_SCAN_LOG_FILE, "w", encoding="utf-8") as f:
                json.dump(self._scan_log, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        row = self._log_table.rowCount()
        self._log_table.insertRow(row)
        for col, val in enumerate([ts, profile, url, title]):
            self._log_table.setItem(row, col, QTableWidgetItem(val))
        self._log_table.scrollToBottom()

    def _load_scan_log(self):
        import json
        from PySide6.QtWidgets import QTableWidgetItem
        if not os.path.exists(_SCAN_LOG_FILE):
            return
        try:
            with open(_SCAN_LOG_FILE, "r", encoding="utf-8") as f:
                self._scan_log = json.load(f)
        except Exception:
            self._scan_log = []
            return
        for entry in self._scan_log:
            row = self._log_table.rowCount()
            self._log_table.insertRow(row)
            for col, val in enumerate([
                entry.get("ts", ""),
                entry.get("profile", ""),
                entry.get("url", ""),
                entry.get("title", ""),
            ]):
                self._log_table.setItem(row, col, QTableWidgetItem(val))
        self._log_table.scrollToBottom()

    def _open_log_entry(self, index):
        row = index.row()
        if row < 0 or row >= len(self._scan_log):
            return
        entry = self._scan_log[row]
        dlg = ScanLogViewerDialog(entry, self)
        dlg.exec()

    def _clear_scan_log(self):
        import json
        self._scan_log = []
        self._log_table.setRowCount(0)
        try:
            with open(_SCAN_LOG_FILE, "w", encoding="utf-8") as f:
                json.dump([], f)
        except Exception:
            pass

    def _on_fast_error(self, msg: str):
        self._fast_status.setText(f"❌  {msg}")
        self._fast_btn.setEnabled(True)


# ── Main Window ───────────────────────────────────────────────────────────────
class DocGeneratorWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Document Generator — Advertising Agreement")
        self.setMinimumSize(1280, 820)
        self._worker: Optional[QThread] = None
        self._parse_worker: Optional[QThread] = None
        self._generated_text: str = ""
        self._setup_ui()
        self.setStyleSheet(DARK_STYLESHEET)
        self._status("Ready. Fill fields, parse a URL, or use Auto-fill.")
        # Sync name fields → signature widgets
        self.inp_contractor.textChanged.connect(self.sig_contractor.set_name)
        self.inp_contractor.textChanged.connect(self.stamp_contractor.set_name)
        self.inp_client.textChanged.connect(self.sig_client.set_name)

    # ── UI ────────────────────────────────────────────────────────────────────
    def _setup_ui(self):
        # ── Tab container ──────────────────────────────────────────────────
        self._tabs = QTabWidget()
        self.setCentralWidget(self._tabs)

        # ── Tab 1: Agreement Generator ─────────────────────────────────────
        agreement_tab = QWidget()
        root = QHBoxLayout(agreement_tab)
        root.setContentsMargins(10, 10, 10, 10)
        root.setSpacing(8)

        splitter = QSplitter(Qt.Horizontal)
        root.addWidget(splitter)

        # Scrollable left panel
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFixedWidth(375)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(4, 4, 8, 4)
        ll.setSpacing(8)
        scroll.setWidget(left)

        # Title
        title = QLabel("⚙️  Agreement Generator")
        title.setFont(QFont("Segoe UI", 14, QFont.Bold))
        title.setStyleSheet("color: #89b4fa; padding-bottom: 2px;")
        ll.addWidget(title)

        # ── 1. Site Parser ────────────────────────────────────────────────────
        grp_parse = QGroupBox("🌐  Site Parser")
        pl = QVBoxLayout(grp_parse)
        pl.setSpacing(5)
        row_url = QHBoxLayout()
        self.inp_url = QLineEdit()
        self.inp_url.setPlaceholderText("https://example.com")
        row_url.addWidget(self.inp_url)
        self.btn_parse = QPushButton("⚡ Parse")
        self.btn_parse.setObjectName("btn_parse")
        self.btn_parse.setFixedWidth(72)
        self.btn_parse.clicked.connect(self._on_parse)
        self.inp_url.returnPressed.connect(self._on_parse)
        row_url.addWidget(self.btn_parse)
        pl.addLayout(row_url)
        self.lbl_parse = QLabel("Parse URL → auto-fill Client + logo (Googlebot).")
        self.lbl_parse.setStyleSheet("color: #6c7086; font-size: 11px;")
        self.lbl_parse.setWordWrap(True)
        pl.addWidget(self.lbl_parse)
        ll.addWidget(grp_parse)

        # ── 2. Auto-fill ──────────────────────────────────────────────────────
        grp_auto = QGroupBox("🎲  Auto-fill")
        al = QVBoxLayout(grp_auto)
        al.setSpacing(5)
        self.chk_auto_num = QCheckBox("Random 5-digit agreement number (unique)")
        self.chk_auto_num.setChecked(True)
        al.addWidget(self.chk_auto_num)
        self.chk_auto_date = QCheckBox("Random past date (1 / 1.5 / 2 months before today)")
        self.chk_auto_date.setChecked(True)
        al.addWidget(self.chk_auto_date)
        self.chk_auto_term = QCheckBox("Random term from [3, 5, 6] months")
        self.chk_auto_term.setChecked(True)
        al.addWidget(self.chk_auto_term)
        row_range = QHBoxLayout()
        row_range.addWidget(self._lbl("Price range:"))
        self.cmb_range = QComboBox()
        for k in _PRICE_RANGES:
            self.cmb_range.addItem(k)
        self.cmb_range.setCurrentIndex(1)
        self.cmb_range.currentTextChanged.connect(self._on_range_changed)
        row_range.addWidget(self.cmb_range)
        al.addLayout(row_range)
        self.btn_autofill = QPushButton("🎲  Apply Auto-fill")
        self.btn_autofill.setObjectName("btn_autofill")
        self.btn_autofill.clicked.connect(self._on_autofill)
        al.addWidget(self.btn_autofill)
        ll.addWidget(grp_auto)

        # ── 3. Agreement Details ──────────────────────────────────────────────
        grp_det = QGroupBox("📄  Agreement Details")
        dl = QVBoxLayout(grp_det)
        dl.setSpacing(5)
        self.inp_number = self._field(dl, "Agreement Number", "e.g. 84421")
        self.inp_number.setText(_random_agreement_number())
        self.inp_date = self._date_field(dl, "Agreement Date")
        # Lock date field when auto-date checkbox is checked
        def _agr_toggle_date(checked: bool, _de=None):
            self.inp_date.setEnabled(not checked)
        self.chk_auto_date.toggled.connect(_agr_toggle_date)
        _agr_toggle_date(self.chk_auto_date.isChecked())  # initial lock
        self.inp_contractor = self._field(dl, "Contractor name", "Company name or individual full name")
        row_ctype = QHBoxLayout()
        row_ctype.addWidget(self._lbl("Contractor type"))
        self.cmb_contractor_type = NoScrollComboBox()
        self.cmb_contractor_type.addItems(["Organization", "Individual"])
        row_ctype.addWidget(self.cmb_contractor_type)
        dl.addLayout(row_ctype)
        self.inp_client = self._field(dl, "Client name", "Company name or individual full name")
        row_cltype = QHBoxLayout()
        row_cltype.addWidget(self._lbl("Client type"))
        self.cmb_client_type = NoScrollComboBox()
        self.cmb_client_type.addItems(["Organization", "Individual"])
        row_cltype.addWidget(self.cmb_client_type)
        dl.addLayout(row_cltype)

        row_price = QHBoxLayout()
        row_price.addWidget(self._lbl("Price"))
        self.inp_price = NoScrollDoubleSpinBox()
        self.inp_price.setRange(0, 9_999_999)
        self.inp_price.setDecimals(2)
        self.inp_price.setValue(5000.0)
        self.inp_price.setSingleStep(500)
        self.inp_price.setPrefix("$ ")
        row_price.addWidget(self.inp_price)
        self.cmb_currency = NoScrollComboBox()
        self.cmb_currency.addItems(["USD ($)", "EUR (€)"])
        self.cmb_currency.setFixedWidth(90)
        self.cmb_currency.currentTextChanged.connect(self._on_currency_changed)
        row_price.addWidget(self.cmb_currency)
        dl.addLayout(row_price)

        row_term = QHBoxLayout()
        row_term.addWidget(self._lbl("Term (months)"))
        self.inp_term = NoScrollSpinBox()
        self.inp_term.setRange(1, 60)
        self.inp_term.setValue(3)
        row_term.addWidget(self.inp_term)
        dl.addLayout(row_term)
        ll.addWidget(grp_det)

        # ── 4. Logo ───────────────────────────────────────────────────────────
        grp_logo = QGroupBox("🖼️  Client Logo  (DOCX / PDF)")
        logo_lay = QVBoxLayout(grp_logo)
        logo_lay.setContentsMargins(6, 6, 6, 6)
        logo_lay.setSpacing(4)
        logo_note = QLabel("Автоматично завантажується при парсингу URL або завантажте вручну.")
        logo_note.setStyleSheet("color: #6c7086; font-size: 11px;")
        logo_note.setWordWrap(True)
        logo_lay.addWidget(logo_note)
        self.logo_widget = LogoWidget()
        logo_lay.addWidget(self.logo_widget)
        self.lbl_logo_status = QLabel("")
        self.lbl_logo_status.setStyleSheet("color: #6c7086; font-size: 11px;")
        self.lbl_logo_status.setWordWrap(True)
        logo_lay.addWidget(self.lbl_logo_status)
        ll.addWidget(grp_logo)

        # ── 5. Signatures ─────────────────────────────────────────────────────
        grp_sig = QGroupBox("✍️  Signatures")
        sl = QVBoxLayout(grp_sig)
        sl.setSpacing(8)
        if not _PIL_AVAILABLE:
            warn = QLabel("⚠️ Pillow not installed — generation disabled.\npip install Pillow")
            warn.setStyleSheet("color: #f38ba8; font-size: 11px;")
            warn.setWordWrap(True)
            sl.addWidget(warn)
        self.sig_contractor = SignatureWidget("Contractor")
        sl.addWidget(self.sig_contractor)
        self.stamp_contractor = StampWidget("Contractor")
        sl.addWidget(self.stamp_contractor)
        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet("color: #45475a;")
        sl.addWidget(sep)
        self.sig_client = SignatureWidget("Client")
        sl.addWidget(self.sig_client)
        ll.addWidget(grp_sig)

        # ── 5. AI Rephrase ────────────────────────────────────────────────────
        grp_ai = QGroupBox("🤖  AI Rephrase (optional)")
        ail = QVBoxLayout(grp_ai)
        ail.setSpacing(5)
        note = QLabel("OpenAI API key → 10–20% variation, all values locked.")
        note.setWordWrap(True)
        note.setStyleSheet("color: #6c7086; font-size: 11px;")
        ail.addWidget(note)
        self.inp_api_key = QLineEdit()
        self.inp_api_key.setPlaceholderText("sk-...")
        self.inp_api_key.setEchoMode(QLineEdit.Password)
        # Load previously saved key
        _saved_key: str = str(QSettings("DocGenerator", "DocGenerator").value("openai_api_key", "") or "")
        if _saved_key:
            self.inp_api_key.setText(_saved_key)
            self.inp_api_key.setReadOnly(True)
        row_api = QHBoxLayout()
        row_api.addWidget(self.inp_api_key)
        self.btn_api_lock = QPushButton("🔒 Save & Lock" if not _saved_key else "🔓 Change")
        self.btn_api_lock.setFixedWidth(100)
        self.btn_api_lock.clicked.connect(self._on_api_lock_toggle)
        row_api.addWidget(self.btn_api_lock)
        ail.addLayout(row_api)
        ll.addWidget(grp_ai)

        # ── 6. Generate ───────────────────────────────────────────────────────
        self.btn_generate = QPushButton("⚡  Generate Agreement")
        self.btn_generate.setObjectName("btn_generate")
        self.btn_generate.clicked.connect(self._on_generate)
        ll.addWidget(self.btn_generate)

        # ── 7. Export ─────────────────────────────────────────────────────────
        grp_exp = QGroupBox("💾  Export")
        el = QVBoxLayout(grp_exp)
        el.setSpacing(5)

        self.btn_validate = QPushButton("🔍  Validate Document")
        self.btn_validate.setObjectName("btn_validate")
        self.btn_validate.clicked.connect(self._on_validate_btn)
        el.addWidget(self.btn_validate)

        self.btn_export_txt = QPushButton("💾  Save as .TXT")
        self.btn_export_txt.setObjectName("btn_export_txt")
        self.btn_export_txt.clicked.connect(self._export_txt)
        el.addWidget(self.btn_export_txt)

        self.btn_export_docx = QPushButton("📄  Save as .DOCX  (signatures embedded)")
        self.btn_export_docx.setObjectName("btn_export_docx")
        self.btn_export_docx.clicked.connect(self._export_docx)
        if not _DOCX_AVAILABLE:
            self.btn_export_docx.setEnabled(False)
            self.btn_export_docx.setToolTip("pip install python-docx")
        el.addWidget(self.btn_export_docx)

        self.btn_export_pdf = QPushButton("🖨️  Save as .PDF  (signatures embedded)")
        self.btn_export_pdf.setObjectName("btn_export_pdf")
        self.btn_export_pdf.clicked.connect(self._export_pdf)
        if not _FPDF_AVAILABLE:
            self.btn_export_pdf.setEnabled(False)
            self.btn_export_pdf.setToolTip("pip install fpdf2")
        el.addWidget(self.btn_export_pdf)

        self.btn_open_pdf_location = QPushButton("📂  Open PDF Location")
        self.btn_open_pdf_location.setObjectName("btn_open_pdf_location")
        self.btn_open_pdf_location.clicked.connect(self._open_pdf_location)
        self.btn_open_pdf_location.setVisible(False)
        el.addWidget(self.btn_open_pdf_location)
        ll.addWidget(grp_exp)

        ll.addStretch()

        # ── Right panel ───────────────────────────────────────────────────────
        rw = QWidget()
        rl = QVBoxLayout(rw)
        rl.setContentsMargins(0, 0, 0, 0)
        rl.setSpacing(5)

        hdr = QHBoxLayout()
        plbl = QLabel("📋  Document Preview  (editable)")
        plbl.setFont(QFont("Segoe UI", 11, QFont.Bold))
        plbl.setStyleSheet("color: #a6e3a1;")
        hdr.addWidget(plbl)
        hdr.addStretch()
        self.lbl_badge = QLabel("")
        self.lbl_badge.setStyleSheet("font-weight: bold; font-size: 12px;")
        hdr.addWidget(self.lbl_badge)
        rl.addLayout(hdr)

        self.preview = QTextEdit()
        self.preview.setReadOnly(False)
        self.preview.setPlaceholderText(
            "Generated agreement will appear here…\n\n"
            "[SIG_CONTRACTOR] / [SIG_CLIENT] = signature image placeholders.\n"
            "They are replaced by actual images in DOCX export."
        )
        self.preview.setFont(QFont("Consolas", 11))
        rl.addWidget(self.preview)

        splitter.addWidget(scroll)
        splitter.addWidget(rw)
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        splitter.setSizes([375, 905])

        self._tabs.addTab(agreement_tab, "📄  Agreement Generator")

        # ── Tab 2: Website Business Summary Parser ─────────────────────────
        self.summary_widget = WebsiteSummaryWidget(self)
        self.summary_widget.send_to_agreement.connect(self._on_send_to_agreement)
        self._tabs.addTab(self.summary_widget, "🌐  Business Summary Parser")

        # ── Tab 3: Reserve Docs / ODA ──────────────────────────────────────
        self.oda_widget = OdaWidget(self)
        self.oda_widget._api_key_getter = lambda: self.inp_api_key.text().strip()
        self._tabs.addTab(self.oda_widget, "🛡️  Reserve Docs")

        # ── Tab 4: DBA / Operating Business Name Statement ─────────────────
        self.dba_widget = DbaStatementWidget(self)
        self.dba_widget._api_key_getter = lambda: self.inp_api_key.text().strip()
        self._tabs.addTab(self.dba_widget, "📋  DBA Statement")

        # ── Tab 5: Octo Appeal Script Generator ─────────────────────────
        self.octo_script_widget = OctoScriptWidget(self)
        self._tabs.addTab(self.octo_script_widget, "🤖  Appeal Script")

        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)

    # ── Hotkeys ───────────────────────────────────────────────────────────────

    _APPSCRIPT_URL = (
        "https://script.google.com/macros/s/"
        "AKfycbyybHlqe3F9C9cL2LWRGJIbkxm6f0C0WCbBrruc2MKbOFZ7_ksRuJ_UCPjynVqA4CxK/exec"
    )

    def keyPressEvent(self, event):
        """F — copy AppScript URL;  R — copy current document preview."""
        focused = QApplication.focusWidget()
        if isinstance(focused, (QLineEdit, QTextEdit)):
            super().keyPressEvent(event)
            return
        key = event.key()
        if key == Qt.Key.Key_F:
            # Tab 4 (Appeal Script): open AppScript create URL
            if self._tabs.currentIndex() == 4:
                self.octo_script_widget.open_appscript()
            else:
                QApplication.clipboard().setText(self._APPSCRIPT_URL)
                self.status_bar.showMessage("✅  AppScript URL скопійовано в буфер  [F]", 3000)
        elif key == Qt.Key.Key_R:
            # Copy preview text from the currently active tab
            tab_idx = self._tabs.currentIndex()
            text = ""
            if tab_idx == 0:
                text = self.preview.toPlainText().strip()
            elif tab_idx == 2:
                text = self.oda_widget.preview.toPlainText().strip()
            elif tab_idx == 3:
                text = self.dba_widget.preview.toPlainText().strip()
            elif tab_idx == 4:
                msg = self.octo_script_widget.generate_and_copy()
                self.status_bar.showMessage(msg, 5000)
                return
            if text:
                QApplication.clipboard().setText(text)
                self.status_bar.showMessage("✅  Документ скопійовано в буфер  [R]", 3000)
            else:
                self.status_bar.showMessage("⚠️  Спочатку згенеруй документ  [R]", 3000)
        else:
            super().keyPressEvent(event)

    def _lbl(self, text: str) -> QLabel:
        l = QLabel(text)
        l.setObjectName("section_label")
        return l

    def _field(self, layout, label: str, placeholder: str = "") -> QLineEdit:
        layout.addWidget(self._lbl(label))
        inp = QLineEdit()
        inp.setPlaceholderText(placeholder)
        layout.addWidget(inp)
        return inp

    def _date_field(self, layout, label: str) -> QDateEdit:
        layout.addWidget(self._lbl(label))
        de = QDateEdit(_random_past_date())
        de.setCalendarPopup(True)
        de.setDisplayFormat("MMMM d, yyyy")
        layout.addWidget(de)
        return de

    def _status(self, msg: str):
        self.status_bar.showMessage(msg)

    def _open_pdf_location(self):
        path = getattr(self, "_last_pdf_path", None)
        if path and os.path.exists(path):
            import subprocess
            subprocess.Popen(["explorer", "/select,", os.path.normpath(path)])

    def _on_send_to_agreement(self, data: dict):
        """Receive brand/domain from Business Summary Parser, fill Client field."""
        brand = data.get("brand", "")
        if brand:
            self.inp_client.setText(brand)
        self._tabs.setCurrentIndex(0)
        ownership = data.get("ownership", "")
        suffix = f"  ({ownership} website)" if ownership else ""
        self._status(
            f"✅ Client filled from Business Summary: «{brand}»{suffix}"
        )

    # ── Auto-fill ─────────────────────────────────────────────────────────────
    def _on_currency_changed(self, text: str):
        if "EUR" in text:
            self.inp_price.setPrefix("€ ")
        else:
            self.inp_price.setPrefix("$ ")

    def _on_currency_changed(self, text: str):
        if "EUR" in text:
            self.inp_price.setPrefix("€ ")
        else:
            self.inp_price.setPrefix("$ ")

    def _on_range_changed(self, key: str):
        self.inp_price.setEnabled(key == "Custom")

    def _on_autofill(self):
        if self.chk_auto_num.isChecked():
            self.inp_number.setText(_random_agreement_number())
        if self.chk_auto_date.isChecked():
            self.inp_date.setDate(_random_past_date())
        if self.chk_auto_term.isChecked():
            self.inp_term.setValue(random.choice([3, 5, 6]))
        key = self.cmb_range.currentText()
        if key != "Custom":
            self.inp_price.setValue(_random_price(key))
        self._status("🎲 Auto-fill applied.")

    # ── Site parse ────────────────────────────────────────────────────────────
    def _on_parse(self):
        url = self.inp_url.text().strip()
        if not url:
            QMessageBox.warning(self, "URL Required", "Enter a URL to parse.")
            return
        if not url.startswith("http"):
            url = "https://" + url
        self.btn_parse.setEnabled(False)
        self.lbl_parse.setText("⏳ Scanning (Googlebot)…")
        self._parse_worker = FastSiteParseWorker(url, "", self)
        self._parse_worker.finished.connect(self._on_parse_fast_done)
        self._parse_worker.error.connect(self._on_parse_error)
        self._parse_worker.start()

    @staticmethod
    def _extract_brand_from_scan(result_text: str) -> str:
        """Extract brand name from FastSiteParseWorker result text."""
        title = ""
        h1    = ""
        for line in result_text.splitlines():
            if line.startswith("TITLE\t"):
                title = line[6:].strip()
            elif line.startswith("H1 HEADINGS\t"):
                raw = line[len("H1 HEADINGS\t"):].strip()
                if raw and raw != "No H1 Tags Found":
                    h1 = raw.split(" | ")[0].strip()
        # 1. Try title segment before separator
        if title:
            for sep in (" | ", " - ", " — ", " :: ", " · "):
                if sep in title:
                    part = title.split(sep)[0].strip()
                    if part and len(part.split()) <= 5 and len(part) <= 50:
                        return part
            if len(title.split()) <= 4 and len(title) <= 40:
                return title
        # 2. Short H1
        if h1 and len(h1.split()) <= 5 and len(h1) <= 50 and not h1.endswith("."):
            return h1
        # 3. Fallback: first title segment by any separator
        if title:
            return title.split("|")[0].split(" - ")[0].strip()
        return ""

    def _on_parse_fast_done(self, result_text: str, _profile: str, final_url: str):
        self.btn_parse.setEnabled(True)
        brand = self._extract_brand_from_scan(result_text)
        if brand:
            self.inp_client.setText(brand)
            self.lbl_parse.setText(f"✅ «{brand}» → Client filled.")
            self._status(f"✅ Parsed: {brand}")
        else:
            self.lbl_parse.setText("⚠️ No brand found. Fill manually.")
            self._status("⚠️ Parse OK but no brand detected.")
        # ── Auto-fetch logo ─────────────────────────────────────────────────
        self.lbl_logo_status.setText("🔄 Завантажую логотип…")
        self._logo_worker = LogoFetchWorker(final_url or self.inp_url.text().strip(), self)
        self._logo_worker.finished.connect(self._on_logo_fetched)
        self._logo_worker.error.connect(self._on_logo_error)
        self._logo_worker.start()

    def _on_logo_fetched(self, path: str):
        px = QPixmap(path)
        if not px.isNull():
            self.logo_widget._path    = path
            self.logo_widget._pixmap  = px
            self.logo_widget.preview.setPixmap(
                px.scaled(56, 36, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            )
            self.logo_widget.changed.emit()
            self.lbl_logo_status.setText("✅ Логотип завантажено.")
        else:
            self.lbl_logo_status.setText("⚠️ Не вдалось завантажити логотип.")

    def _on_logo_error(self, err: str):
        self.lbl_logo_status.setText(f"❌ {err}")

    def _on_parse_error(self, err: str):
        self.btn_parse.setEnabled(True)
        self.lbl_parse.setText(f"❌ {err}")
        self._status(f"❌ Parse error: {err}")

    # ── API key lock ──────────────────────────────────────────────────────────
    def _on_api_lock_toggle(self):
        settings = QSettings("DocGenerator", "DocGenerator")
        if self.inp_api_key.isReadOnly():
            # Currently locked → ask for confirmation to unlock
            reply = QMessageBox.question(
                self, "Change API Key",
                "Are you sure you want to clear the saved API key and enter a new one?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No,
            )
            if reply == QMessageBox.Yes:
                settings.remove("openai_api_key")
                self.inp_api_key.setReadOnly(False)
                self.inp_api_key.clear()
                self.inp_api_key.setFocus()
                self.btn_api_lock.setText("🔒 Save & Lock")
                self._status("🔓 API key cleared — enter a new key.")
        else:
            # Currently unlocked → save and lock
            key = self.inp_api_key.text().strip()
            if not key:
                QMessageBox.warning(self, "API Key", "Please enter an API key before saving.")
                return
            settings.setValue("openai_api_key", key)
            self.inp_api_key.setReadOnly(True)
            self.btn_api_lock.setText("🔓 Change")
            self._status("🔒 API key saved and locked.")

    # ── Generate ──────────────────────────────────────────────────────────────
    def _collect(self) -> Optional[dict]:
        number = self.inp_number.text().strip()
        date_str = self.inp_date.date().toString("MMMM d, yyyy")
        contractor = self.inp_contractor.text().strip()
        client = self.inp_client.text().strip()
        price = self.inp_price.value()
        term = self.inp_term.value()

        errs = []
        if not number:
            errs.append("Agreement Number is required.")
        if not contractor:
            errs.append("Contractor name is required.")
        if not client:
            errs.append("Client name is required.")
        if errs:
            QMessageBox.warning(self, "Validation Error", "\n".join(errs))
            return None

        currency_sym = "€" if "EUR" in self.cmb_currency.currentText() else "$"
        currency_word = "EUR" if "EUR" in self.cmb_currency.currentText() else "USD"
        return dict(
            number=number, date_str=date_str, contractor=contractor,
            client=client, price=price,
            price_str=f"{currency_word} {price:,.2f}",
            currency_sym=currency_sym, currency_word=currency_word,
            term=term, term_words=months_to_words(term),
        )

    def _on_generate(self):
        if self.chk_auto_num.isChecked(): self.inp_number.setText(_random_agreement_number())
        if self.chk_auto_date.isChecked(): self.inp_date.setDate(_random_past_date())
        if self.chk_auto_term.isChecked(): self.inp_term.setValue(random.choice([3, 5, 6]))

        data = AgreementData(
            mode=self.cmb_mode.currentText().lower(),
            agreement_number=self.inp_number.text(),
            agreement_date=self.inp_date.date().toString("yyyy-MM-dd"),
            service_fee=self.inp_price.value(),
            currency=self.cmb_currency.currentText().split("(")[0].strip(),
            contractor=PartyDetails(
                trade_name=self.inputs["contractor.trade_name"].text(),
                legal_name=self.inputs["contractor.legal_name"].text(),
                registration_number=self.inputs["contractor.registration_number"].text(),
                iban=self.inputs["contractor.iban"].text()
            ),
            client=PartyDetails(
                trade_name=self.inputs["client.trade_name"].text(),
                legal_name=self.inputs["client.legal_name"].text(),
                registration_number=self.inputs["client.registration_number"].text()
            )
        )
        verified = {path for path, chk in self.verifications.items() if chk.isChecked()}
        missing, warnings = validate_agreement(data, verified)
        self.val_text.setText("\n".join(warnings) if warnings else "Ready.")

        api_key = self.inp_api_key.text().strip()
        if not api_key:
            res = ContractGenerator().generate_full_agreement(data)
            self._set_preview(res["agreement_markdown"])
            self._status("✅ Local template used.")
            return

        self.btn_generate.setEnabled(False)
        self._status("🤖 AI Generating...")

        # Reuse existing OpenAIWorker for threading
        prompt = f"Generate JSON agreement: {json.dumps(data.to_dict())}"
        self._worker = OpenAIWorker(api_key, prompt, self)

        def _done(full):
            try:
                import json as _json
                # Find JSON block in AI response
                start = full.find('{')
                end = full.rfind('}') + 1
                if start != -1 and end > start:
                    ret = _json.loads(full[start:end])
                    text = ret.get("agreement_markdown", full)
                else: text = full
                text = protect_hallucinations(text, data)
                self._set_preview(text)
                self._status("✅ AI generation complete.")
            except:
                self._set_preview(full)
            self.btn_generate.setEnabled(True)

        self._worker.finished.connect(_done)
        self._worker.error.connect(lambda e: (self.btn_generate.setEnabled(True), self._status(f"❌ AI Error: {e}")))
        self._worker.start()

    def _set_preview(self, text):
        self._generated_text = text
        self.preview.setPlainText(text)

    def _on_save_project(self):
        path, _ = QFileDialog.getSaveFileName(self, "Save Project", "project.json", "JSON Files (*.json)")
        if path:
            data = AgreementData(mode=self.cmb_mode.currentText().lower())
            save_project(ProjectData(agreement_data=data, generated_sections={"full": self.preview.toPlainText()}), path)
            self._status(f"💾 Project saved: {path}")

    def _on_load_project(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open Project", "", "JSON Files (*.json)")
        if path:
            proj = load_project(path); self.preview.setPlainText(proj.generated_sections.get("full", ""))
            self._status(f"📂 Project loaded: {path}")

    def _export_txt(self):
        text = self.preview.toPlainText().strip()
        if not text: return
        num = self.inp_number.text().strip() or "agreement"
        path, _ = QFileDialog.getSaveFileName(self, "Save as TXT", f"Agreement_{num}.txt", "Text Files (*.txt)")
        if path:
            with open(path, "w", encoding="utf-8") as f: f.write(text.replace("[SIG_CONTRACTOR]", "_"*32).replace("[SIG_CLIENT]", "_"*32))
            self._status(f"💾 Saved: {path}")

    def _export_md(self):
        text = self.preview.toPlainText().strip()
        if not text: return
        num = self.inp_number.text().strip() or "agreement"
        path, _ = QFileDialog.getSaveFileName(self, "Save as Markdown", f"Agreement_{num}.md", "Markdown Files (*.md)")
        if path: export_markdown(text, path); self._status(f"📝 Saved: {path}")

    def _export_json_action(self):
        text = self.preview.toPlainText().strip()
        if not text: return
        data = AgreementData(mode=self.cmb_mode.currentText().lower())
        num = self.inp_number.text().strip() or "agreement"
        path, _ = QFileDialog.getSaveFileName(self, "Save as JSON", f"Agreement_{num}.json", "JSON Files (*.json)")
        if path: export_json(data, text, path); self._status(f"📦 Saved: {path}")

    def _export_pdf(self):
        text = self.preview.toPlainText().strip()
        if not text: return
        path, _ = QFileDialog.getSaveFileName(self, "Save as PDF", "Agreement.pdf", "PDF Files (*.pdf)")
        if not path: return
        data = AgreementData(mode=self.cmb_mode.currentText().lower())
        sig_c = sig_cl = None
        tmp = os.path.join(os.path.dirname(os.path.abspath(__file__)), "__pycache__")
        os.makedirs(tmp, exist_ok=True)
        if self.sig_contractor.has_signature():
            sig_c = os.path.join(tmp, "_sig_c_p.png"); self.sig_contractor.get_pixmap().save(sig_c, "PNG")
        if self.sig_client.has_signature():
            sig_cl = os.path.join(tmp, "_sig_cl_p.png"); self.sig_client.get_pixmap().save(sig_cl, "PNG")
        build_pdf(text, path, data, sig_c, sig_cl)
        self._status(f"🖨️ Saved: {path}")

    def _export_docx(self):
        text = self.preview.toPlainText().strip()
        if not text: return
        path, _ = QFileDialog.getSaveFileName(self, "Save as DOCX", "Agreement.docx", "DOCX Files (*.docx)")
        if not path: return
        data = AgreementData(mode=self.cmb_mode.currentText().lower())
        sig_c = sig_cl = stamp_c = None
        tmp = os.path.join(os.path.dirname(os.path.abspath(__file__)), "__pycache__")
        os.makedirs(tmp, exist_ok=True)
        if self.sig_contractor.has_signature():
            sig_c = os.path.join(tmp, "_sig_c_d.png"); self.sig_contractor.get_pixmap().save(sig_c, "PNG")
        if self.sig_client.has_signature():
            sig_cl = os.path.join(tmp, "_sig_cl_d.png"); self.sig_client.get_pixmap().save(sig_cl, "PNG")
        if self.stamp_contractor.has_stamp():
            stamp_c = os.path.join(tmp, "_stamp_c_d.png"); self.stamp_contractor.get_pixmap().save(stamp_c, "PNG")
        build_docx(text, path, data, sig_c, sig_cl, stamp_c, logo_path=self.logo_widget.get_path() if self.logo_widget.has_logo() else None)
        self._status(f"📄 Saved: {path}")

    def _on_regen_section(self):
        text = self.preview.toPlainText()
        if not text: return
        api_key = self.inp_api_key.text().strip()
        if not api_key: return
        self._status("♻️ AI Regenerating section...")
        data = AgreementData(mode=self.cmb_mode.currentText().lower())
        gen = ContractGenerator(openai_client=_OpenAI(api_key=api_key))
        self.preview.setPlainText(gen.regenerate_section(data, "Selected", text))

def main():
    app = QApplication(sys.argv)
    app.setApplicationName("Document Generator")
    win = DocGeneratorWindow()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
